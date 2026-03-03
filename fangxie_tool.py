# -*- coding: utf-8 -*-
"""
百家号文案仿写工具 - 可视化界面
功能：读取参考文案，根据选择的引流类型生成仿写文案
支持流式/非流式调用，主模型+备用模型切换
"""

import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext
import os
import json
import threading
from datetime import datetime
import requests
import re
import time
import random
from collections import Counter
from difflib import SequenceMatcher

# 爆款素材库路径
MATERIAL_LIBRARY_DIR = r"D:\A百家号带货视频\带货文案\爆款参考文案"
MATERIAL_LIBRARY_FILE = os.path.join(MATERIAL_LIBRARY_DIR, "爆款素材库.xlsx")
# 生成文案库（用于去模板比对与标题去重，放在爆款素材库目录）
GENERATED_LIBRARY_FILE = os.path.join(MATERIAL_LIBRARY_DIR, "生成文案库.xlsx")
GENERATED_LIBRARY_SHEET = "生成文案库"

# 开头钩子库 - 用于随机选择
HOOK_LIBRARY = {
    "A类-轻松幽默型": [
        "我发现一个规律：越是老实人，越容易被安排加班。你说气不气？",
        "昨天算了一笔账，这些年帮别人花的时间，够我学会三门外语了。",
        "我妈说我最大的优点是善良，我爸说这也是我最大的缺点。好家伙，亲爹。",
        "朋友说我是'便利店型人格'——24小时营业，随叫随到，还不涨价。",
        "有人说我脾气好，我笑了笑没说话。其实不是脾气好，是懒得计较。",
        "我终于明白，为什么'好人卡'发得最多——因为好人最好打发。",
        "同事问我：'你怎么从来不生气？'我说：'生气要花力气，我选择省着点用。'",
        "我的人生信条曾经是'吃亏是福'，直到我发现福没来，亏倒是吃了不少。",
        "别人都在研究怎么赚钱，我在研究怎么不被人当免费劳动力。",
        "我这人有个毛病，别人一说'就你能帮我'，我就跟中了蛊似的。",
        "我朋友说我是'人形充电宝'，谁没电了都来找我，充完就走。",
        "有人问我为什么总是笑，我说：哭也解决不了问题，还费纸巾。",
        "我发现一个真理：会哭的孩子有糖吃，不会哭的孩子只能自己买。",
        "我妈问我：'你怎么朋友那么多？'我说：'因为我从不麻烦别人。'",
        "别人的底线是不能触碰，我的底线是随时可以后退。",
        "我终于知道为什么我人缘好了——因为我从来不给别人添堵，只给自己添堵。",
        "有人说我是'海绵体质'，什么委屈都能吸收，还不反弹。",
    ],
    "B类-温暖治愈型": [
        "嘿，今天想跟你聊点轻松的，关于那些默默付出却不求回报的人。",
        "你知道吗，这世上有一种人，他们的好，是藏在细节里的。",
        "我一直觉得，善良的人身上有光，只是有时候这光被辜负了。",
        "如果你正在看这段话，我想告诉你：你的好，有人看得见。",
        "今天不讲大道理，就想跟你说说心里话。",
        "有些人，值得被这个世界温柔以待，比如正在看这段话的你。",
        "我见过很多人，但像你这样的，真的不多。",
        "你有没有被人夸过'你人真好'？今天我想认真聊聊这件事。",
        "这段话，送给每一个在生活里默默扛着的人。",
        "我相信，看到这里的你，一定是个心里有温度的人。",
        "你的温柔，不是软弱，是这个世界最稀缺的力量。",
        "那些你以为没人看见的付出，其实都被时间记住了。",
        "你不必向任何人证明你的好，懂你的人自然懂。",
        "愿你的善良，终有一天能被温柔以待。",
        "你已经很棒了，只是你自己不知道而已。",
        "这世上总有人在默默爱着你，只是你没发现。",
        "你的坚强，不是因为你不怕疼，是因为你扛得住。",
    ],
    "C类-反转惊喜型": [
        "我曾经以为自己是个'老好人'，后来发现，我是个'聪明的好人'。",
        "都说老实人吃亏，但我认识一个老实人，现在过得比谁都好。",
        "你以为善良是软弱？不，善良是一种选择，而且是强者的选择。",
        "有人说心软的人没出息，我偏不信这个邪。",
        "我见过最厉害的人，恰恰是最善良的那个。",
        "别人都说我太老实会吃亏，结果呢？我还真没亏。",
        "都说好人没好报，但我今天要讲一个好人有好报的故事。",
        "我以前觉得'人善被人欺'是真理，直到我遇见了一个人。",
        "谁说善良的人就要受委屈？我第一个不服。",
        "老实人的春天，其实一直都在，只是很多人没发现。",
        "你以为我是在忍？不，我是在等一个时机。",
        "那些笑我太傻的人，现在都在问我借钱。",
        "我不是不会反击，我只是在选择值不值得。",
        "都说我没脾气，其实我的脾气只留给值得的人。",
        "你以为我吃亏了？其实我赚的是人心。",
        "那些年我吃的亏，现在都变成了我的铠甲。",
        "我不争，不是因为我怕，是因为我知道时间会给答案。",
    ],
    "D类-对话引入型": [
        "他说完这句话，我愣在原地半天没回过神：'你就是太好说话了。'",
        "有人问我：'你这辈子最后悔的事是什么？'我想了想，说了两个字。",
        "算了，不争了——你是不是也经常这样跟自己说？",
        "我爷爷临终前拉着我的手说：'记住，吃亏的人，老天爷都记着账呢。'",
        "'你怎么这么傻？'这句话，你听过多少次了？",
        "我妈常说一句话：'人善被人欺，马善被人骑。'我以前不信，现在信了。",
        "'别太老实了，会吃亏的。'说这话的人，后来怎么样了？",
        "朋友跟我说：'你知道你最大的问题是什么吗？就是太把别人当回事。'",
        "'谢谢你'——这三个字，你等了多久才听到？",
        "有人当面问我：'你是不是傻？人家都欺负到头上了你还忍？'",
        "我问我爸：'为什么好人总是吃亏？'他沉默了很久才回答我。",
        "'你变了。'有人这样说我，我笑了笑没解释。",
        "领导拍着我的肩膀说：'小伙子，你太实在了。'我不知道这是夸还是骂。",
        "'你就不能为自己想想吗？'这句话，我听了无数遍。",
        "老婆问我：'你在外面受了委屈，为什么不跟我说？'",
        "'你这样的人，活该被人欺负。'这话扎心，但我没反驳。",
        "孩子问我：'爸爸，为什么你总是让着别人？'我不知道怎么回答。",
    ],
    "E类-生活观察型": [
        "菜市场大妈的一句话，让我愣了半天：'姑娘，你这么好说话，不怕被人欺负啊？'",
        "堵在路上，收音机里突然传来一句话，我眼眶一下就红了。",
        "排队的时候，前面两个人的对话让我心里一惊。",
        "超市结账时，前面那个人的一个举动，让我看清了人性。",
        "参加同学聚会，有个人的变化让我震惊了。",
        "小区门口，两个大妈聊天，聊着聊着说出了一个真相。",
        "去参加婚礼，新郎的一句话让全场安静了。",
        "早上买早餐，老板娘的一句话让我想了一整天。",
        "坐出租车，司机师傅跟我聊了一路，最后一句话让我沉默了。",
        "地铁上，旁边两个年轻人的对话，让我想起了自己。",
        "公司茶水间，无意中听到的一段话，让我彻底清醒了。",
        "深夜的便利店，店员的一句话，让我破防了。",
        "高铁上，邻座大哥讲的故事，让我一路都在想。",
        "咖啡店里，隔壁桌两个人的对话，让我若有所思。",
        "健身房里，教练的一句话，让我重新审视自己。",
        "书店角落，一个陌生人的举动，让我相信善良还在。",
        "电梯里，邻居的一句闲聊，让我想了一整天。",
    ],
    "F类-数字锚定型": [
        "认识老王二十年了，他教会我一件事：别对谁都掏心掏肺。",
        "被人欺负了三年，我终于想通了一个道理。",
        "用了十年时间，我才学会一个字：不。",
        "三次被人背叛之后，我悟了。",
        "五十岁之后，我才明白什么叫'人走茶凉'。",
        "帮了他七年，他一句谢谢都没说过。",
        "借出去的三万块，要了五年都没要回来。",
        "在这个单位干了八年，我终于明白了一个道理。",
        "三十五岁那年，我做了一个决定，从此人生开挂。",
        "被坑了五次之后，我终于学会了看人。",
        "工作十二年，我才明白什么叫'职场生存法则'。",
        "结婚十五年，我才懂得什么是真正的爱。",
        "创业失败三次，我终于找到了成功的秘诀。",
        "四十岁之前，我一直在为别人活；四十岁之后，我决定为自己活。",
        "被拒绝了一百次，我终于明白了一个道理。",
        "坚持了七年的习惯，彻底改变了我的人生。",
        "二十年的老友，一夜之间形同陌路，原因竟然是这个。",
    ],
    "G类-金句破题型": [
        "最傻的事，就是跟烂人讲道理。",
        "千万别做老好人，我吃过这个亏，现在告诉你。",
        "你越忍让，别人越得寸进尺，这是我用十年换来的教训。",
        "老实人不是没脾气，是把脾气都咽进了肚子里。",
        "你的善良，要带点锋芒。",
        "不是所有的忍让都叫大度，有时候那叫窝囊。",
        "这世上最傻的事，就是把真心给了不值得的人。",
        "有些人，你帮他一百次，他记不住；你拒绝他一次，他记你一辈子。",
        "成年人最大的清醒：不是所有人都值得你掏心掏肺。",
        "真正的高手，从不解释，只用结果说话。",
        "你的时间很贵，别浪费在不值得的人身上。",
        "人这辈子，最怕的不是对手，是那些表面和善的人。",
        "别把别人想得太好，也别把自己看得太低。",
        "你不必讨好所有人，做好自己就够了。",
        "有些路，注定要一个人走；有些苦，注定要一个人扛。",
        "真正的强大，是内心的平静，而不是外表的强硬。",
        "你的格局，决定了你的结局。",
    ],
    "H类-悬念钩子型": [
        "有件事我憋了很久，今天必须说出来。",
        "你可能不信，但接下来我说的都是真事。",
        "我要告诉你一个很多人不愿意承认的真相。",
        "接下来这段话，可能会让你不舒服，但我还是要说。",
        "有个规律，我观察了很多年才看透。",
        "今天说的这些话，可能会得罪人，但我不在乎。",
        "有些话，我本来不想说，但看到你，我忍不住了。",
        "接下来的话，你可能不爱听，但句句都是真的。",
        "有件事，我一直没跟任何人说过，今天破例。",
        "我发现了一个秘密，一直不敢说，今天豁出去了。",
        "接下来的内容，可能会颠覆你的认知。",
        "有个真相，很多人知道，但没人敢说。",
        "我要揭露一个潜规则，可能会得罪很多人。",
        "这件事我犹豫了很久，今天终于决定说出来。",
        "有些事，不说出来，我会后悔一辈子。",
        "接下来的话，听进去的人，人生会少走很多弯路。",
        "我要告诉你一个残酷的事实，准备好了吗？",
        "你知道为什么有些人总是被欺负吗？答案可能出乎你意料。",
        "你有没有想过，为什么你总是那个吃亏的人？",
        "你身上有一个特质，可能连你自己都没意识到。",
        "你以为的缺点，其实是你最大的优势，听我说完你就明白了。",
        "你一直在等的那个答案，今天我告诉你。",
        "你心里那个疑问，今天我帮你解开。",
        "你是不是一直在想，为什么自己总是遇人不淑？",
        "你可能从来没想过，你的善良其实是一种天赋。",
    ],
    "I类-共鸣代入型": [
        "那种被人当众下面子的感觉，你是不是也经历过？",
        "明明没做错什么，却总被人针对，你是不是也遇到过？",
        "有些委屈，说出来都没人信，你懂这种感觉吗？",
        "你是不是也有过这种感觉：付出最多的人，往往最不被珍惜。",
        "那种心寒的感觉，就像一盆冷水从头浇到脚，你体会过吗？",
        "有一种苦，叫做'打碎了牙往肚子里咽'，你尝过吗？",
        "那种被人利用完就扔掉的感觉，你是不是也经历过？",
        "那种明明很努力，却得不到认可的无力感，你是不是也有过？",
        "有些话，憋在心里太久，会发霉的，你是不是也这样？",
        "那种被人误解却无法解释的感觉，你是不是也经历过？",
        "你是不是也曾经，为了不让别人失望，把自己逼得很累？",
        "那种笑着说没事，转身却红了眼眶的感觉，你有过吗？",
        "有些伤，不是不疼，是疼久了就麻木了，你是不是也这样？",
        "你是不是也有过这种时刻：明明很委屈，却说不出口？",
        "那种被人当透明人的感觉，你是不是也体会过？",
        "有些苦，只有自己知道；有些泪，只能往心里流，你懂吗？",
        "你是不是也常常在想：为什么我对别人那么好，却换不来同样的真心？",
        "你是不是也有过这种感觉：越是在乎的人，越容易让你失望？",
        "你是不是也曾经，把别人的事当成自己的事，最后却没人领情？",
        "你是不是也有过这种时刻：明明很累，却不敢停下来？",
        "你是不是也常常觉得：自己付出了那么多，却总是被忽略？",
        "你是不是也有过这种感觉：越是善良，越容易被人当软柿子捏？",
        "你是不是也曾经，为了维护一段关系，委屈了自己很多次？",
        "你是不是也有过这种时刻：想发脾气，却又怕伤了和气？",
    ],
    "J类-人物故事型": [
        "我有个朋友，前两天跟我说了一件事，我听完沉默了很久。",
        "我爸这辈子只教过我一个道理，我到现在才真正理解。",
        "单位有个人，大家都不待见他，后来我才知道原因。",
        "我们小区有个大爷，天天在楼下坐着，有一天他跟我说了一番话。",
        "我表姐的经历，让我彻底看清了人心。",
        "我有个同学，当年是班里最老实的人，你猜他现在怎么样了？",
        "我舅舅年轻时吃过一个大亏，他把这个教训告诉了我。",
        "我认识一个人，他的经历让我相信：好人终有好报。",
        "我奶奶活了九十岁，临终前只跟我说了一句话。",
        "公司有个老员工，干了二十年，最后的结局让所有人唏嘘。",
        "我发小的故事，让我明白了什么叫'人心隔肚皮'。",
        "我师傅退休前，把他的看家本领全教给了我。",
        "我邻居家的孩子，从小被人看不起，现在却让所有人刮目相看。",
        "我大学室友的经历，让我懂得了什么叫'厚积薄发'。",
        "我前同事的遭遇，让我看清了职场的残酷。",
        "我姑姑的人生，是我见过最励志的故事。",
        "我老板的发家史，颠覆了我对成功的认知。",
    ],
    "K类-自我剖析型": [
        "说出来不怕你笑话，我以前也是个傻子。",
        "回头看看这些年，我最后悔的一件事是太心软。",
        "如果能重来，我绝对不会再做老好人。",
        "我吃过的亏，今天全告诉你，希望你别再走我的老路。",
        "我这辈子最大的毛病，就是太把别人当回事。",
        "我曾经也是个'老好人'，后来我学聪明了。",
        "我年轻时犯过一个错，现在想起来还后悔。",
        "我以前总觉得吃亏是福，现在不这么想了。",
        "我承认，我曾经是个讨好型人格，活得很累。",
        "我最大的弱点，就是不懂得拒绝。",
        "我花了很多年，才学会爱自己。",
        "我曾经把所有人都放在自己前面，结果呢？",
        "我以前总是委屈自己，成全别人，现在不了。",
        "我最后悔的事，就是太晚明白这个道理。",
        "我曾经以为忍让是美德，后来才知道那是懦弱。",
        "我用了半辈子，才学会一个字：不。",
        "我以前总怕得罪人，现在我只怕对不起自己。",
    ],
    "L类-转折反差型": [
        "以前我不信这个道理，直到自己栽了跟头。",
        "年轻的时候觉得这话是废话，现在才知道是真理。",
        "曾经有人跟我说过一句话，我没当回事，后来我后悔了。",
        "我一直以为自己做得对，直到那件事发生。",
        "以前别人说我太老实，我还不服气，现在我服了。",
        "我曾经以为善良是优点，后来才知道，善良过头就是缺点。",
        "年轻时我不懂，现在我懂了，可惜晚了。",
        "我以前总是忍，以为忍一忍就过去了，结果呢？",
        "我曾经以为努力就会有回报，后来发现不是这样的。",
        "以前我觉得朋友越多越好，现在我只想要几个真心的。",
        "我曾经追求完美，后来才明白，接受不完美才是真正的成熟。",
        "以前我总想证明自己，现在我只想做好自己。",
        "我曾经害怕孤独，后来发现，独处是一种能力。",
        "以前我总是向外求，现在我学会了向内看。",
        "我曾经以为成功就是有钱，后来才明白，内心的富足才是真正的成功。",
        "以前我总想改变别人，现在我只想改变自己。",
        "我曾经活在别人的眼光里，现在我只为自己而活。",
    ],
}

def get_random_hooks(count=3):
    """随机选择指定数量的不同类型开头，确保至少1个是轻松/温暖/反转类型"""
    light_types = ["A类-轻松幽默型", "B类-温暖治愈型", "C类-反转惊喜型"]
    all_types = list(HOOK_LIBRARY.keys())

    # 确保至少选1个轻松类型
    selected_types = [random.choice(light_types)]

    # 从剩余类型中再选(count-1)个
    remaining = [t for t in all_types if t != selected_types[0]]
    additional_count = min(count - 1, len(remaining))
    selected_types.extend(random.sample(remaining, additional_count))

    # 打乱顺序
    random.shuffle(selected_types)

    # 从每个类型中随机选一个示例
    result = []
    for t in selected_types:
        hook = random.choice(HOOK_LIBRARY[t])
        result.append({"type": t, "example": hook})

    return result

# 配置文件路径
CONFIG_FILE = os.path.join(os.path.dirname(__file__), "fangxie_config.json")

# 默认配置
DEFAULT_CONFIG = {
    "use_stream": True,  # 是否使用流式调用
    "similarity_threshold": 0.76,  # 生成文案相似度阈值（越低越严格）
    # 流式调用 - 主模型配置
    "stream_main_url": "https://api.aifuwu.icu/v1",
    "stream_main_key": "sk-hc6yUaXg89eK5UgUii10DPWmdaJZdPXqPbPcKSRbmWgxeeDK",
    "stream_main_model": "gemini-3-pro-preview",
    "stream_main_max_tokens": 16000,
    # 流式调用 - 备用模型配置
    "stream_backup_url": "https://yunyi.rdzhvip.com/v1",
    "stream_backup_key": "A68YGT53-41HW-8GT7-353U-MJVDNY8NE5KX",
    "stream_backup_model": "claude-opus-4-5-20251101",
    "stream_backup_max_tokens": 16000,
    # 非流式调用 - 主模型配置
    "non_stream_main_url": "https://yunyi.rdzhvip.com/v1",
    "non_stream_main_key": "A68YGT53-41HW-8GT7-353U-MJVDNY8NE5KX",
    "non_stream_main_model": "claude-opus-4-5-20251101",
    "non_stream_main_max_tokens": 16000,
    # 非流式调用 - 备用模型配置
    "non_stream_backup_url": "https://api.aifuwu.icu/v1",
    "non_stream_backup_key": "sk-hc6yUaXg89eK5UgUii10DPWmdaJZdPXqPbPcKSRbmWgxeeDK",
    "non_stream_backup_model": "gemini-3-pro-preview",
    "non_stream_backup_max_tokens": 16000,
    # 引流话术库（按类型分开存储）
    "yinliu_templates": {
        "置顶引流": [
            "关于怎么彻底走出这个困局，我在主页置顶视频里讲得很透。那里有一套方法，是我这些年摸爬滚打总结出来的，不适合在这里公开讲。你要是真想改变，点我头像，去看置顶第一条，看完你就明白该怎么做了。",
            "现在你就点开我的头像，进我主页，去看置顶的那个视频，我在那里给你留了一套破局的方法。听得懂那便是你的收获，听不懂说明时机还没到。准备好了吗？这一局该你赢了。",
            "从今天开始，学着为自己活一次... 答应我，现在点击我的头像进主页看看，置顶前两条视频，那里有我想对你说的心里话，这些话我只说给你听，因为你扛了太久，总该有人懂你的不易。",
            "如果你也想从这种困境里走出来，想活得通透一点、轻松一点，点我头像，去看主页置顶的视频。那里有你一直在找的答案，也有你需要的那份力量。我在那里等你。"
        ],
        "橱窗引流": [
            "现在你只需要做一件事，点开我的头像，进入主页橱窗，不要带着太多顾虑去挑拣，也不要问哪一件最好，你静静的看。若它让你心有所动，那就是适合你的，切莫错过。",
            "选一件，是给自己一份小小的犒赏；选两件，是为自己的生活增添一点温暖；若你愿意，选三件，就是给自己一个完整的礼物，让生活多一些美好。",
            "去吧，朋友，点开头像进橱窗，去找回那个本该发光的自己。我在那头等着，看你越来越好，自在如风。",
            "选一件，第一眼入心的就是你接下来全力冲锋的底气；选2件，便是对自己一路隐忍坚持的犒赏；选3件，更代表你不只顾着当下赶路，还在为长远未来布局筹谋。"
        ],
        "带货引流": [
            "轻轻一点左下角，它就能到你家。我知道你会犹豫，怕没用，怕白费钱，但我劝你大胆试这一次。这不是消费，是投资，投资你的安稳生活，投资你的人生底气。点击左下角把它请回家...",
            "好东西从不会长久停留，立刻点击左下角，让这份美好为你的生活增添一点温暖和力量...",
            "如今这几件东西我把它放到了左下角的通道里，别再犹豫了，去左下角把它请回家吧。这不是在帮别人，是在帮你自己... 对自己好一点，你值得。",
            "现在请顺应内心的指引，点击下方通道，把这款好物带回家。你为别人操心了大半辈子，是时候对自己好一点了。"
        ]
    }
}

def load_config():
    """加载配置"""
    if os.path.exists(CONFIG_FILE):
        try:
            with open(CONFIG_FILE, 'r', encoding='utf-8') as f:
                config = json.load(f)
                # 补充缺失的配置项
                for key in DEFAULT_CONFIG:
                    if key not in config:
                        config[key] = DEFAULT_CONFIG[key]
                return config
        except:
            pass
    return DEFAULT_CONFIG.copy()

def save_config(config):
    """保存配置"""
    with open(CONFIG_FILE, 'w', encoding='utf-8') as f:
        json.dump(config, f, ensure_ascii=False, indent=2)

class FangxieApp:
    def __init__(self, root):
        self.root = root
        self.root.title("百家号文案仿写工具")
        self.root.geometry("1100x950")
        self.root.minsize(900, 700)
        self.root.resizable(True, True)

        # 加载配置
        self.config = load_config()

        # 素材库路径（固定）
        self.material_path = r"D:\A百家号带货文案库"

        # 引流类型映射
        self.flow_types = {
            "置顶引流": "置顶视频引流素材.txt",
            "橱窗引流": "橱窗引流素材.txt",
            "带货引流": "带货引流素材.txt",
            "纯夸赞不引流": None
        }

        # 视频制作相关配置
        self.video_source_path = r"D:\BaiduNetdiskDownload\自然风景视频素材"
        self.max_videos_per_folder = 3
        self.ffmpeg_path = r"C:\Users\Administrator\AppData\Local\Microsoft\WinGet\Packages\Gyan.FFmpeg_Microsoft.Winget.Source_8wekyb3d8bbwe\ffmpeg-8.0.1-full_build\bin\ffmpeg.exe"
        self.ffprobe_path = r"C:\Users\Administrator\AppData\Local\Microsoft\WinGet\Packages\Gyan.FFmpeg_Microsoft.Winget.Source_8wekyb3d8bbwe\ffmpeg-8.0.1-full_build\bin\ffprobe.exe"
        self.whisper_model = None
        self.video_is_running = False

        self.create_widgets()

    def create_widgets(self):
        # 创建主Notebook（标签页）
        self.main_notebook = ttk.Notebook(self.root)
        self.main_notebook.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)

        # === 标签页1：文案生成 ===
        self.article_page = ttk.Frame(self.main_notebook)
        self.main_notebook.add(self.article_page, text="  文案生成  ")
        self.create_article_page()

        # === 标签页2：视频制作 ===
        self.video_page = ttk.Frame(self.main_notebook)
        self.main_notebook.add(self.video_page, text="  视频制作  ")
        self.create_video_page()

        # === 标签页3：API配置 ===
        self.api_page = ttk.Frame(self.main_notebook)
        self.main_notebook.add(self.api_page, text="  API配置  ")
        self.create_api_page()

        # 初始化运行状态
        self.is_running = False
        self.last_articles = []
        self.last_flow_type = ""
        self.last_yinliu_content = ""
        self.last_product_name = ""
        self.last_product_material = ""
        self.similarity_threshold = float(self.config.get("similarity_threshold", 0.76))

    def create_article_page(self):
        """创建文案生成页面"""
        # 创建滚动框架
        canvas = tk.Canvas(self.article_page)
        scrollbar = ttk.Scrollbar(self.article_page, orient="vertical", command=canvas.yview)
        scrollable_frame = ttk.Frame(canvas)

        scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )

        # 创建窗口并保存ID，用于后续调整宽度
        canvas_window = canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)

        # 让内容宽度和高度跟随Canvas变化
        def on_canvas_configure(event):
            canvas.itemconfig(canvas_window, width=event.width)
            # 让内容高度至少等于画布高度，这样日志区域可以扩展填满窗口
            canvas.itemconfig(canvas_window, height=max(event.height, scrollable_frame.winfo_reqheight()))
        canvas.bind("<Configure>", on_canvas_configure)

        def _on_mousewheel(event):
            canvas.yview_scroll(int(-1*(event.delta/120)), "units")
        canvas.bind_all("<MouseWheel>", _on_mousewheel)

        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")

        main_frame = ttk.Frame(scrollable_frame, padding="10")
        main_frame.pack(fill=tk.BOTH, expand=True)

        # 标题
        title_label = ttk.Label(main_frame, text="文案生成", font=("微软雅黑", 14, "bold"))
        title_label.pack(pady=(0, 10))

        # 流式/非流式变量（在API配置页面切换）
        self.use_stream = tk.BooleanVar(value=self.config.get("use_stream", True))

        # === 参考文案输入 ===
        input_frame = ttk.LabelFrame(main_frame, text="参考文案", padding="10")
        input_frame.pack(fill=tk.X, pady=5)

        # 输入方式选择
        self.input_mode = tk.StringVar(value="file")
        mode_frame = ttk.Frame(input_frame)
        mode_frame.pack(fill=tk.X, pady=(0, 5))
        ttk.Radiobutton(mode_frame, text="上传文件", variable=self.input_mode, value="file",
                       command=self.on_input_mode_change).pack(side=tk.LEFT, padx=10)
        ttk.Radiobutton(mode_frame, text="粘贴文本", variable=self.input_mode, value="paste",
                       command=self.on_input_mode_change).pack(side=tk.LEFT, padx=10)

        # 文件选择区域
        self.file_input_frame = ttk.Frame(input_frame)
        self.file_input_frame.pack(fill=tk.X)
        self.input_path = tk.StringVar(value=r"d:\A百家号带货文案库\仿写文案.txt")
        ttk.Entry(self.file_input_frame, textvariable=self.input_path, width=70).pack(side=tk.LEFT, fill=tk.X, expand=True)
        ttk.Button(self.file_input_frame, text="选择文件", command=self.select_input_file, width=10).pack(side=tk.LEFT, padx=5)
        ttk.Button(self.file_input_frame, text="选择文件夹", command=self.select_input_folder, width=10).pack(side=tk.LEFT)

        # 文本粘贴区域（默认隐藏）
        self.paste_input_frame = ttk.Frame(input_frame)
        self.paste_text = scrolledtext.ScrolledText(self.paste_input_frame, height=8, width=80)
        self.paste_text.pack(fill=tk.X)
        paste_btn_frame = ttk.Frame(self.paste_input_frame)
        paste_btn_frame.pack(fill=tk.X, pady=(5, 0))
        ttk.Button(paste_btn_frame, text="清空", command=lambda: self.paste_text.delete("1.0", tk.END), width=10).pack(side=tk.LEFT)

        # === 输出路径 ===
        output_frame = ttk.LabelFrame(main_frame, text="输出保存路径", padding="10")
        output_frame.pack(fill=tk.X, pady=5)

        self.output_path = tk.StringVar(value=r"D:\A百家号带货视频\带货文案")
        ttk.Entry(output_frame, textvariable=self.output_path, width=70).pack(side=tk.LEFT, fill=tk.X, expand=True)
        ttk.Button(output_frame, text="选择文件夹", command=self.select_output_folder, width=10).pack(side=tk.LEFT, padx=5)

        # === 流量文案保存路径 ===
        txt_output_frame = ttk.LabelFrame(main_frame, text="流量文案保存路径（生成流量文案按钮使用）", padding="10")
        txt_output_frame.pack(fill=tk.X, pady=5)

        self.txt_output_path = tk.StringVar(value=r"D:/AIDownloadFiles/国学json/百家号带货视频/baijiadaihuo/input/视频文案/流量文案")
        ttk.Entry(txt_output_frame, textvariable=self.txt_output_path, width=70).pack(side=tk.LEFT, fill=tk.X, expand=True)
        ttk.Button(txt_output_frame, text="选择文件夹", command=self.select_txt_output_folder, width=10).pack(side=tk.LEFT, padx=5)

        # 字数选择
        ttk.Label(txt_output_frame, text="  字数:").pack(side=tk.LEFT)
        self.word_count = tk.StringVar(value="1200")
        word_count_combo = ttk.Combobox(txt_output_frame, textvariable=self.word_count, width=8, state="readonly")
        word_count_combo['values'] = ["600", "700", "800", "900", "1000", "1200", "1500", "1600"]
        word_count_combo.pack(side=tk.LEFT, padx=5)

        # 生成篇数选择
        ttk.Label(txt_output_frame, text="  篇数:").pack(side=tk.LEFT)
        self.article_count = tk.StringVar(value="3")
        article_count_combo = ttk.Combobox(txt_output_frame, textvariable=self.article_count, width=5, state="readonly")
        article_count_combo['values'] = ["1", "2", "3", "5", "8"]
        article_count_combo.pack(side=tk.LEFT, padx=5)

        # === 语音合成配置 ===
        voice_config_frame = ttk.LabelFrame(main_frame, text="语音合成配置（合成语音按钮使用）", padding="10")
        voice_config_frame.pack(fill=tk.X, pady=5)

        # 音色选择
        voice_row1 = ttk.Frame(voice_config_frame)
        voice_row1.pack(fill=tk.X, pady=2)
        ttk.Label(voice_row1, text="音色选择:").pack(side=tk.LEFT)
        self.voice_type = tk.StringVar(value="智慧老者")
        self.voice_combo = ttk.Combobox(voice_row1, textvariable=self.voice_type, width=20, state="readonly")
        self.voice_combo['values'] = ["鸡汤女声", "智慧老者", "沉稳大气男声"]
        self.voice_combo.pack(side=tk.LEFT, padx=5)

        # 文案输入目录
        voice_row2 = ttk.Frame(voice_config_frame)
        voice_row2.pack(fill=tk.X, pady=2)
        ttk.Label(voice_row2, text="文案目录:").pack(side=tk.LEFT)
        self.voice_input_path = tk.StringVar(value=r"D:/AIDownloadFiles/国学json/百家号带货视频/baijiadaihuo/input/视频文案/流量文案")
        ttk.Entry(voice_row2, textvariable=self.voice_input_path, width=60).pack(side=tk.LEFT, fill=tk.X, expand=True, padx=5)
        ttk.Button(voice_row2, text="选择", command=self.select_voice_input_folder, width=6).pack(side=tk.LEFT)

        # 配音输出目录
        voice_row3 = ttk.Frame(voice_config_frame)
        voice_row3.pack(fill=tk.X, pady=2)
        ttk.Label(voice_row3, text="输出目录:").pack(side=tk.LEFT)
        self.voice_output_path = tk.StringVar(value=r"D:/AIDownloadFiles/国学json/百家号带货视频/baijiadaihuo/input/视频配音/流量语音")
        ttk.Entry(voice_row3, textvariable=self.voice_output_path, width=60).pack(side=tk.LEFT, fill=tk.X, expand=True, padx=5)
        ttk.Button(voice_row3, text="选择", command=self.select_voice_output_folder, width=6).pack(side=tk.LEFT)

        # === 引流类型选择 ===
        flow_frame = ttk.LabelFrame(main_frame, text="引流类型", padding="10")
        flow_frame.pack(fill=tk.X, pady=5)

        self.flow_type = tk.StringVar(value="置顶引流")
        for ft in self.flow_types.keys():
            ttk.Radiobutton(flow_frame, text=ft, variable=self.flow_type, value=ft,
                           command=self.on_flow_type_change).pack(side=tk.LEFT, padx=10)

        # === 引流话术区域 ===
        self.yinliu_frame = ttk.LabelFrame(main_frame, text="引流话术（可选）", padding="10")
        self.yinliu_frame.pack(fill=tk.X, pady=5)

        # 话术下拉框
        yinliu_combo_frame = ttk.Frame(self.yinliu_frame)
        yinliu_combo_frame.pack(fill=tk.X, pady=(0, 5))
        ttk.Label(yinliu_combo_frame, text="请选择结尾参考话术:").pack(side=tk.LEFT)
        self.yinliu_combo = ttk.Combobox(yinliu_combo_frame, width=50, state="readonly")
        self.yinliu_combo.pack(side=tk.LEFT, padx=5)
        self.yinliu_combo.bind("<<ComboboxSelected>>", self.on_yinliu_select)

        # 话术操作按钮
        yinliu_btn_frame = ttk.Frame(self.yinliu_frame)
        yinliu_btn_frame.pack(fill=tk.X, pady=(0, 5))
        ttk.Button(yinliu_btn_frame, text="保存当前话术", command=self.save_yinliu_template, width=12).pack(side=tk.LEFT, padx=2)
        ttk.Button(yinliu_btn_frame, text="删除选中", command=self.delete_yinliu_template, width=10).pack(side=tk.LEFT, padx=2)
        ttk.Button(yinliu_btn_frame, text="清空输入", command=self.clear_yinliu_text, width=10).pack(side=tk.LEFT, padx=2)

        # 话术输入框提示
        ttk.Label(self.yinliu_frame, text="↓ 在此输入结尾引流话术（AI会参考此风格，为每篇生成不同的结尾）：", foreground="gray").pack(anchor=tk.W)

        # 话术输入框
        self.yinliu_text = scrolledtext.ScrolledText(self.yinliu_frame, height=4, width=80)
        self.yinliu_text.pack(fill=tk.X)

        # 初始化话术下拉框
        self.update_yinliu_combo()

        # === 带货信息区域（默认隐藏） ===
        self.daihuo_frame = ttk.LabelFrame(main_frame, text="带货商品信息", padding="10")

        # 商品名称
        name_frame = ttk.Frame(self.daihuo_frame)
        name_frame.pack(fill=tk.X, pady=2)
        ttk.Label(name_frame, text="商品名称:", width=10).pack(side=tk.LEFT)
        self.product_name = tk.StringVar()
        ttk.Entry(name_frame, textvariable=self.product_name, width=50).pack(side=tk.LEFT, fill=tk.X, expand=True)

        # 产品素材
        ttk.Label(self.daihuo_frame, text="产品素材/介绍:").pack(anchor=tk.W, pady=(5, 2))
        self.product_material = scrolledtext.ScrolledText(self.daihuo_frame, height=4, width=80)
        self.product_material.pack(fill=tk.X)

        # === 操作按钮 ===
        self.btn_frame_container = ttk.Frame(main_frame)
        self.btn_frame_container.pack(fill=tk.X, pady=10)

        btn_frame = ttk.Frame(self.btn_frame_container)
        btn_frame.pack()

        self.start_btn = ttk.Button(btn_frame, text="开始生成", command=self.start_generate, width=15)
        self.start_btn.pack(side=tk.LEFT, padx=10)

        self.start_txt_btn = ttk.Button(btn_frame, text="生成流量文案", command=self.start_generate_txt, width=12)
        self.start_txt_btn.pack(side=tk.LEFT, padx=10)

        self.synth_voice_btn = ttk.Button(btn_frame, text="合成语音", command=self.start_synth_voice, width=12)
        self.synth_voice_btn.pack(side=tk.LEFT, padx=10)

        self.stop_btn = ttk.Button(btn_frame, text="停止", command=self.stop_generate, width=10, state=tk.DISABLED)
        self.stop_btn.pack(side=tk.LEFT, padx=10)

        self.regenerate_btn = ttk.Button(btn_frame, text="重新生成", command=self.regenerate, width=12, state=tk.DISABLED)
        self.regenerate_btn.pack(side=tk.LEFT, padx=10)

        ttk.Button(btn_frame, text="打开输出文件夹", command=self.open_output_folder, width=15).pack(side=tk.LEFT, padx=10)
        ttk.Button(btn_frame, text="去制作视频 →", command=self.go_to_video_page, width=12).pack(side=tk.LEFT, padx=10)

        # === 重新生成建议 ===
        self.suggestion_frame = ttk.LabelFrame(main_frame, text="修改建议", padding="10")
        self.suggestion_frame.pack(fill=tk.X, pady=5)

        ttk.Label(self.suggestion_frame, text="当您对生成的结果不满意时，请在下方输入意见，并点击重新生成：").pack(anchor=tk.W)
        self.suggestion_text = scrolledtext.ScrolledText(self.suggestion_frame, height=3, width=80)
        self.suggestion_text.pack(fill=tk.X)

        # === 进度条 ===
        progress_frame = ttk.Frame(main_frame)
        progress_frame.pack(fill=tk.X, pady=5)

        self.progress_var = tk.DoubleVar(value=0)
        self.progress_bar = ttk.Progressbar(progress_frame, variable=self.progress_var, maximum=100)
        self.progress_bar.pack(fill=tk.X)

        self.status_label = ttk.Label(progress_frame, text="就绪")
        self.status_label.pack(pady=5)

        # === 日志区域 ===
        log_frame = ttk.LabelFrame(main_frame, text="运行日志", padding="10")
        log_frame.pack(fill=tk.BOTH, expand=True, pady=5)

        self.log_text = scrolledtext.ScrolledText(log_frame, height=20, width=100)
        self.log_text.pack(fill=tk.BOTH, expand=True)

    def create_api_page(self):
        """创建API配置页面"""
        main_frame = ttk.Frame(self.api_page, padding="15")
        main_frame.pack(fill=tk.BOTH, expand=True)

        # 标题
        title_label = ttk.Label(main_frame, text="API配置", font=("微软雅黑", 14, "bold"))
        title_label.pack(pady=(0, 15))

        # 说明文字
        ttk.Label(main_frame, text="选择调用模式，配置对应的API参数。切换标签页即切换调用模式。", foreground="gray").pack(anchor=tk.W, pady=(0, 10))

        # 创建Notebook用于切换流式/非流式配置
        self.api_notebook = ttk.Notebook(main_frame)
        self.api_notebook.pack(fill=tk.BOTH, expand=True, pady=(0, 15))

        # === 流式调用配置页 ===
        stream_page = ttk.Frame(self.api_notebook, padding="10")
        self.api_notebook.add(stream_page, text="  流式调用（推荐）  ")

        # 流式-主模型
        stream_main_frame = ttk.LabelFrame(stream_page, text="主模型", padding="10")
        stream_main_frame.pack(fill=tk.X, pady=5)
        self._create_model_config(stream_main_frame, "stream_main")

        # 流式-备用模型
        stream_backup_frame = ttk.LabelFrame(stream_page, text="备用模型（主模型失败后自动切换）", padding="10")
        stream_backup_frame.pack(fill=tk.X, pady=5)
        self._create_model_config(stream_backup_frame, "stream_backup")

        # === 非流式调用配置页 ===
        non_stream_page = ttk.Frame(self.api_notebook, padding="10")
        self.api_notebook.add(non_stream_page, text="  非流式调用  ")

        # 非流式-主模型
        non_stream_main_frame = ttk.LabelFrame(non_stream_page, text="主模型", padding="10")
        non_stream_main_frame.pack(fill=tk.X, pady=5)
        self._create_model_config(non_stream_main_frame, "non_stream_main")

        # 非流式-备用模型
        non_stream_backup_frame = ttk.LabelFrame(non_stream_page, text="备用模型（主模型失败后自动切换）", padding="10")
        non_stream_backup_frame.pack(fill=tk.X, pady=5)
        self._create_model_config(non_stream_backup_frame, "non_stream_backup")

        # 绑定标签页切换事件，自动更新流式/非流式设置
        self.api_notebook.bind("<<NotebookTabChanged>>", self.on_api_tab_change)

        # 根据配置初始化选中的标签页
        if not self.config.get("use_stream", True):
            self.api_notebook.select(1)  # 选中非流式标签页

        # 去模板相似度阈值
        threshold_frame = ttk.LabelFrame(main_frame, text="去模板设置", padding="10")
        threshold_frame.pack(fill=tk.X, pady=(0, 10))
        ttk.Label(threshold_frame, text="相似度阈值:").pack(side=tk.LEFT)
        self.similarity_threshold_var = tk.StringVar(
            value=f"{float(self.config.get('similarity_threshold', 0.76)):.2f}"
        )
        threshold_combo = ttk.Combobox(
            threshold_frame,
            textvariable=self.similarity_threshold_var,
            width=8,
            state="readonly"
        )
        threshold_combo["values"] = [
            "0.50", "0.55", "0.60", "0.65", "0.68", "0.70",
            "0.72", "0.74", "0.76", "0.78", "0.80", "0.82", "0.85", "0.90"
        ]
        threshold_combo.pack(side=tk.LEFT, padx=8)
        ttk.Label(
            threshold_frame,
            text="越低越严格（重试更多）。建议 0.72~0.80",
            foreground="gray"
        ).pack(side=tk.LEFT, padx=8)

        # 底部按钮区域
        btn_frame = ttk.Frame(main_frame)
        btn_frame.pack(fill=tk.X, pady=10)

        ttk.Button(btn_frame, text="保存配置", command=self.save_api_config, width=15).pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_frame, text="重置默认", command=self.reset_api_config, width=15).pack(side=tk.LEFT, padx=5)

        # 当前模式提示
        self.api_mode_label = ttk.Label(main_frame, text="", font=("微软雅黑", 10))
        self.api_mode_label.pack(anchor=tk.W, pady=5)
        self._update_api_mode_label()

    def _update_api_mode_label(self):
        """更新当前模式提示"""
        mode = "流式调用" if self.use_stream.get() else "非流式调用"
        self.api_mode_label.config(text=f"当前使用：{mode}")

    def create_video_page(self):
        """创建视频制作页面"""
        # 创建滚动框架
        canvas = tk.Canvas(self.video_page)
        scrollbar = ttk.Scrollbar(self.video_page, orient="vertical", command=canvas.yview)
        scrollable_frame = ttk.Frame(canvas)

        scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )

        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)

        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")

        main_frame = ttk.Frame(scrollable_frame, padding="10")
        main_frame.pack(fill=tk.BOTH, expand=True)

        # 标题
        ttk.Label(main_frame, text="视频制作", font=("微软雅黑", 14, "bold")).pack(pady=(0, 10))

        # === 模式选择 ===
        mode_frame = ttk.LabelFrame(main_frame, text="1. 选择制作模式", padding="10")
        mode_frame.pack(fill=tk.X, pady=5)

        self.video_mode = tk.StringVar(value="audio")
        ttk.Radiobutton(mode_frame, text="使用已有音频（Whisper识别字幕）",
                       variable=self.video_mode, value="audio",
                       command=self.on_video_mode_change).pack(anchor=tk.W)
        ttk.Radiobutton(mode_frame, text="从生成的文案制作（TTS生成配音）",
                       variable=self.video_mode, value="tts",
                       command=self.on_video_mode_change).pack(anchor=tk.W)

        # === 音频选择区域 ===
        self.audio_frame = ttk.LabelFrame(main_frame, text="2. 选择配音文件", padding="10")
        self.audio_frame.pack(fill=tk.X, pady=5)

        audio_row = ttk.Frame(self.audio_frame)
        audio_row.pack(fill=tk.X)
        self.video_audio_path = tk.StringVar()
        ttk.Entry(audio_row, textvariable=self.video_audio_path, width=60).pack(side=tk.LEFT, fill=tk.X, expand=True)
        ttk.Button(audio_row, text="选择音频", command=self.select_video_audio, width=10).pack(side=tk.LEFT, padx=5)

        # === 文案选择区域（TTS模式） ===
        self.tts_frame = ttk.LabelFrame(main_frame, text="2. 选择文案", padding="10")

        tts_row1 = ttk.Frame(self.tts_frame)
        tts_row1.pack(fill=tk.X, pady=2)
        ttk.Label(tts_row1, text="选择文案:").pack(side=tk.LEFT)
        self.article_combo = ttk.Combobox(tts_row1, width=50, state="readonly")
        self.article_combo.pack(side=tk.LEFT, padx=5)
        self.article_combo.bind("<<ComboboxSelected>>", self.on_article_select)
        ttk.Button(tts_row1, text="刷新列表", command=self.refresh_article_list, width=10).pack(side=tk.LEFT)
        ttk.Button(tts_row1, text="随机选择", command=self.random_select_article, width=10).pack(side=tk.LEFT, padx=5)

        ttk.Label(self.tts_frame, text="文案预览:").pack(anchor=tk.W, pady=(5, 2))
        self.article_preview = scrolledtext.ScrolledText(self.tts_frame, height=4, width=80)
        self.article_preview.pack(fill=tk.X)

        # === 视频标题 ===
        title_frame = ttk.LabelFrame(main_frame, text="3. 视频标题（封面显示）", padding="10")
        title_frame.pack(fill=tk.X, pady=5)

        self.video_title = tk.StringVar()
        ttk.Entry(title_frame, textvariable=self.video_title, width=60).pack(fill=tk.X)

        # === 视频素材设置 ===
        video_frame = ttk.LabelFrame(main_frame, text="4. 视频素材设置", padding="10")
        video_frame.pack(fill=tk.X, pady=5)

        video_path_frame = ttk.Frame(video_frame)
        video_path_frame.pack(fill=tk.X, pady=2)
        ttk.Label(video_path_frame, text="素材文件夹:").pack(side=tk.LEFT)
        self.video_source_var = tk.StringVar(value=self.video_source_path)
        ttk.Entry(video_path_frame, textvariable=self.video_source_var, width=50).pack(side=tk.LEFT, padx=5)
        ttk.Button(video_path_frame, text="选择", command=self.select_video_source_folder, width=6).pack(side=tk.LEFT)

        video_param_frame = ttk.Frame(video_frame)
        video_param_frame.pack(fill=tk.X, pady=2)
        ttk.Label(video_param_frame, text="每文件夹最多取:").pack(side=tk.LEFT)
        self.max_per_folder = tk.StringVar(value="3")
        ttk.Combobox(video_param_frame, textvariable=self.max_per_folder, width=5,
                    values=["1", "2", "3", "4", "5"]).pack(side=tk.LEFT, padx=5)
        ttk.Label(video_param_frame, text="个视频").pack(side=tk.LEFT)

        # 视频素材静音选项
        video_mute_frame = ttk.Frame(video_frame)
        video_mute_frame.pack(fill=tk.X, pady=2)
        self.mute_video = tk.BooleanVar(value=True)
        ttk.Checkbutton(video_mute_frame, text="视频素材静音（只用配音，推荐）", variable=self.mute_video).pack(side=tk.LEFT)

        # === 背景音乐设置 ===
        bgm_frame = ttk.LabelFrame(main_frame, text="5. 背景音乐（可选）", padding="10")
        bgm_frame.pack(fill=tk.X, pady=5)

        bgm_enable_frame = ttk.Frame(bgm_frame)
        bgm_enable_frame.pack(fill=tk.X, pady=2)
        self.enable_bgm = tk.BooleanVar(value=False)
        ttk.Checkbutton(bgm_enable_frame, text="添加背景音乐", variable=self.enable_bgm,
                       command=self.on_bgm_toggle).pack(side=tk.LEFT)

        bgm_path_frame = ttk.Frame(bgm_frame)
        bgm_path_frame.pack(fill=tk.X, pady=2)
        ttk.Label(bgm_path_frame, text="BGM文件夹:").pack(side=tk.LEFT)
        self.bgm_folder = tk.StringVar(value=r"D:\A百家号带货视频\BGM")
        self.bgm_entry = ttk.Entry(bgm_path_frame, textvariable=self.bgm_folder, width=45)
        self.bgm_entry.pack(side=tk.LEFT, padx=5)
        self.bgm_btn = ttk.Button(bgm_path_frame, text="选择", command=self.select_bgm_folder, width=6)
        self.bgm_btn.pack(side=tk.LEFT)

        bgm_vol_frame = ttk.Frame(bgm_frame)
        bgm_vol_frame.pack(fill=tk.X, pady=2)
        ttk.Label(bgm_vol_frame, text="BGM音量:").pack(side=tk.LEFT)
        self.bgm_volume = tk.StringVar(value="15")
        ttk.Combobox(bgm_vol_frame, textvariable=self.bgm_volume, width=5,
                    values=["5", "10", "15", "20", "25", "30"]).pack(side=tk.LEFT, padx=5)
        ttk.Label(bgm_vol_frame, text="%").pack(side=tk.LEFT)

        ttk.Label(bgm_vol_frame, text="配音音量:").pack(side=tk.LEFT, padx=(20, 0))
        self.voice_volume = tk.StringVar(value="100")
        ttk.Combobox(bgm_vol_frame, textvariable=self.voice_volume, width=5,
                    values=["80", "90", "100", "110", "120"]).pack(side=tk.LEFT, padx=5)
        ttk.Label(bgm_vol_frame, text="%").pack(side=tk.LEFT)

        # === 字幕样式 ===
        style_frame = ttk.LabelFrame(main_frame, text="6. 字幕样式", padding="10")
        style_frame.pack(fill=tk.X, pady=5)

        style_row1 = ttk.Frame(style_frame)
        style_row1.pack(fill=tk.X, pady=2)

        ttk.Label(style_row1, text="字号:").pack(side=tk.LEFT)
        self.subtitle_size = tk.StringVar(value="36")
        subtitle_size_combo = ttk.Combobox(style_row1, textvariable=self.subtitle_size, width=6,
                     values=["24", "28", "32", "36", "40", "48"])
        subtitle_size_combo.pack(side=tk.LEFT, padx=5)
        subtitle_size_combo.bind("<<ComboboxSelected>>", self.update_subtitle_preview)

        ttk.Label(style_row1, text="颜色:").pack(side=tk.LEFT, padx=(15, 0))
        self.subtitle_color = tk.StringVar(value="金色")
        subtitle_color_combo = ttk.Combobox(style_row1, textvariable=self.subtitle_color, width=8,
                     values=["金色", "黄色", "橙色", "红色"])
        subtitle_color_combo.pack(side=tk.LEFT, padx=5)
        subtitle_color_combo.bind("<<ComboboxSelected>>", self.update_subtitle_preview)

        ttk.Label(style_row1, text="位置:").pack(side=tk.LEFT, padx=(15, 0))
        self.subtitle_position = tk.StringVar(value="底部")
        subtitle_pos_combo = ttk.Combobox(style_row1, textvariable=self.subtitle_position, width=8,
                     values=["底部", "中部", "顶部"])
        subtitle_pos_combo.pack(side=tk.LEFT, padx=5)
        subtitle_pos_combo.bind("<<ComboboxSelected>>", self.update_subtitle_preview)

        # 字幕预览区域
        preview_row = ttk.Frame(style_frame)
        preview_row.pack(fill=tk.X, pady=(10, 5))

        self.subtitle_preview_canvas = tk.Canvas(preview_row, width=400, height=80, bg="#1a1a2e", highlightthickness=1, highlightbackground="gray")
        self.subtitle_preview_canvas.pack(side=tk.LEFT, padx=5)
        self.update_subtitle_preview()  # 初始化预览

        # === 封面样式 ===
        cover_frame = ttk.LabelFrame(main_frame, text="7. 封面样式", padding="10")
        cover_frame.pack(fill=tk.X, pady=5)

        cover_row = ttk.Frame(cover_frame)
        cover_row.pack(fill=tk.X, pady=2)

        ttk.Label(cover_row, text="标题字号:").pack(side=tk.LEFT)
        self.cover_font_size = tk.StringVar(value="60")
        ttk.Combobox(cover_row, textvariable=self.cover_font_size, width=6,
                     values=["48", "54", "60", "72", "80"]).pack(side=tk.LEFT, padx=5)

        ttk.Label(cover_row, text="标题颜色:").pack(side=tk.LEFT, padx=(15, 0))
        self.cover_color = tk.StringVar(value="金色")
        ttk.Combobox(cover_row, textvariable=self.cover_color, width=8,
                     values=["金色", "黄色", "橙色", "红色"]).pack(side=tk.LEFT, padx=5)

        ttk.Label(cover_row, text="背景:").pack(side=tk.LEFT, padx=(15, 0))
        self.cover_bg = tk.StringVar(value="视频首帧")
        ttk.Combobox(cover_row, textvariable=self.cover_bg, width=10,
                     values=["视频首帧", "黑色背景"]).pack(side=tk.LEFT, padx=5)

        # === 输出设置 ===
        output_frame = ttk.LabelFrame(main_frame, text="8. 输出设置", padding="10")
        output_frame.pack(fill=tk.X, pady=5)

        self.video_output_path = tk.StringVar(value=r"D:\A百家号带货视频\成品视频")
        ttk.Entry(output_frame, textvariable=self.video_output_path, width=60).pack(side=tk.LEFT, fill=tk.X, expand=True)
        ttk.Button(output_frame, text="选择文件夹", command=self.select_video_output_folder, width=10).pack(side=tk.LEFT, padx=5)

        # === 操作按钮 ===
        btn_frame = ttk.Frame(main_frame)
        btn_frame.pack(pady=15)

        self.video_make_btn = ttk.Button(btn_frame, text="开始制作", command=self.start_make_video, width=15)
        self.video_make_btn.pack(side=tk.LEFT, padx=10)

        self.video_stop_btn = ttk.Button(btn_frame, text="停止", command=self.stop_make_video, width=10, state=tk.DISABLED)
        self.video_stop_btn.pack(side=tk.LEFT, padx=10)

        ttk.Button(btn_frame, text="打开输出文件夹", command=self.open_video_output, width=15).pack(side=tk.LEFT, padx=10)

        # === 进度 ===
        progress_frame = ttk.Frame(main_frame)
        progress_frame.pack(fill=tk.X, pady=5)

        self.video_progress_var = tk.DoubleVar(value=0)
        self.video_progress_bar = ttk.Progressbar(progress_frame, variable=self.video_progress_var, maximum=100)
        self.video_progress_bar.pack(fill=tk.X)

        self.video_status_label = ttk.Label(progress_frame, text="就绪")
        self.video_status_label.pack(pady=5)

        # === 日志 ===
        log_frame = ttk.LabelFrame(main_frame, text="制作日志", padding="10")
        log_frame.pack(fill=tk.BOTH, expand=True, pady=5)

        self.video_log_text = scrolledtext.ScrolledText(log_frame, height=10, width=80)
        self.video_log_text.pack(fill=tk.BOTH, expand=True)

    def go_to_video_page(self):
        """切换到视频制作页面"""
        self.main_notebook.select(1)
        self.refresh_article_list()

    def on_video_mode_change(self):
        """视频制作模式切换"""
        mode = self.video_mode.get()
        if mode == "audio":
            self.audio_frame.pack(fill=tk.X, pady=5, after=self.audio_frame.master.winfo_children()[1])
            self.tts_frame.pack_forget()
        else:
            self.tts_frame.pack(fill=tk.X, pady=5, after=self.tts_frame.master.winfo_children()[1])
            self.audio_frame.pack_forget()
            self.refresh_article_list()

    def select_video_audio(self):
        """选择视频配音文件"""
        file_path = filedialog.askopenfilename(
            title="选择配音文件",
            filetypes=[("音频文件", "*.mp3 *.wav *.m4a"), ("所有文件", "*.*")]
        )
        if file_path:
            self.video_audio_path.set(file_path)

    def select_video_source_folder(self):
        """选择视频素材文件夹"""
        folder_path = filedialog.askdirectory(title="选择视频素材文件夹")
        if folder_path:
            self.video_source_var.set(folder_path)

    def select_bgm_folder(self):
        """选择BGM文件夹"""
        folder_path = filedialog.askdirectory(title="选择背景音乐文件夹")
        if folder_path:
            self.bgm_folder.set(folder_path)

    def on_bgm_toggle(self):
        """BGM开关切换"""
        pass  # 界面会自动根据checkbox状态处理

    def get_random_bgm(self):
        """从BGM文件夹随机选择一首背景音乐"""
        bgm_folder = self.bgm_folder.get()
        if not os.path.exists(bgm_folder):
            return None

        bgm_files = []
        for f in os.listdir(bgm_folder):
            if f.lower().endswith(('.mp3', '.wav', '.m4a', '.aac', '.flac')):
                bgm_files.append(os.path.join(bgm_folder, f))

        if bgm_files:
            return random.choice(bgm_files)
        return None

    def select_video_output_folder(self):
        """选择视频输出文件夹"""
        folder_path = filedialog.askdirectory(title="选择输出文件夹")
        if folder_path:
            self.video_output_path.set(folder_path)

    def open_video_output(self):
        """打开视频输出文件夹"""
        output_path = self.video_output_path.get().replace('/', '\\')
        os.makedirs(output_path, exist_ok=True)
        os.startfile(output_path)

    def refresh_article_list(self):
        """刷新文案列表"""
        options = ["-- 请选择文案 --"]
        if self.last_articles:
            for i, article in enumerate(self.last_articles):
                # 提取标题（如果有）
                lines = article.strip().split('\n')
                title = ""
                for line in lines:
                    if line.startswith('【标题'):
                        title = line.replace('【标题1】', '').replace('【标题2】', '').replace('【标题3】', '').strip()[:20]
                        break
                if not title:
                    title = article[:25].replace('\n', ' ')
                options.append(f"文案{i+1}: {title}...")
        self.article_combo['values'] = options
        self.article_combo.current(0)

    def on_article_select(self, event=None):
        """选择文案时显示预览"""
        selection = self.article_combo.current()
        if selection > 0 and self.last_articles:
            article = self.last_articles[selection - 1]
            self.article_preview.delete("1.0", tk.END)
            self.article_preview.insert("1.0", article[:800] + "..." if len(article) > 800 else article)
            # 自动提取标题
            lines = article.strip().split('\n')
            for line in lines:
                if line.startswith('【标题'):
                    title = line.replace('【标题1】', '').replace('【标题2】', '').replace('【标题3】', '').strip()
                    self.video_title.set(title)
                    break

    def random_select_article(self):
        """随机选择一篇文案"""
        if self.last_articles:
            idx = random.randint(0, len(self.last_articles) - 1)
            self.article_combo.current(idx + 1)
            self.on_article_select()
        else:
            messagebox.showinfo("提示", "没有可用的文案，请先在「文案生成」页面生成文案")

    def video_log(self, message):
        """视频制作日志"""
        timestamp = datetime.now().strftime("%H:%M:%S")
        self.video_log_text.insert(tk.END, f"[{timestamp}] {message}\n")
        self.video_log_text.see(tk.END)
        self.root.update()

    def video_update_status(self, message):
        """更新视频制作状态"""
        self.video_status_label.config(text=message)
        self.root.update()

    def video_update_progress(self, value):
        """更新视频制作进度"""
        self.video_progress_var.set(value)
        self.root.update()

    def stop_make_video(self):
        """停止视频制作"""
        self.video_is_running = False
        self.video_log("用户停止了制作任务")
        self.video_update_status("已停止")

    def start_make_video(self):
        """开始制作视频"""
        mode = self.video_mode.get()

        if mode == "audio":
            # 音频模式：需要选择音频文件
            audio_path = self.video_audio_path.get()
            if not audio_path or not os.path.exists(audio_path):
                messagebox.showerror("错误", "请选择有效的配音文件")
                return
        else:
            # TTS模式：需要选择文案
            selection = self.article_combo.current()
            if selection <= 0 or not self.last_articles:
                messagebox.showerror("错误", "请选择一篇文案，或先在「文案生成」页面生成文案")
                return

        video_title = self.video_title.get().strip()
        if not video_title:
            messagebox.showerror("错误", "请输入视频标题")
            return

        # 更新配置
        self.video_source_path = self.video_source_var.get()
        self.max_videos_per_folder = int(self.max_per_folder.get())

        output_path = self.video_output_path.get()
        os.makedirs(output_path, exist_ok=True)

        self.video_make_btn.config(state=tk.DISABLED)
        self.video_stop_btn.config(state=tk.NORMAL)
        self.video_is_running = True
        self.video_log_text.delete(1.0, tk.END)

        # 在线程中执行
        thread = threading.Thread(target=self.make_video_task, args=(mode, video_title, output_path))
        thread.daemon = True
        thread.start()

    def make_video_task(self, mode, video_title, output_path):
        """视频制作主任务"""
        import subprocess

        try:
            temp_dir = os.path.join(output_path, "temp")
            os.makedirs(temp_dir, exist_ok=True)

            # 1. 获取音频
            if mode == "audio":
                audio_path = self.video_audio_path.get()
                self.video_log(f"使用已有音频: {os.path.basename(audio_path)}")
            else:
                # TTS生成配音
                self.video_update_status("正在生成配音...")
                self.video_update_progress(5)
                selection = self.article_combo.current()
                article = self.last_articles[selection - 1]
                # 提取正文（去掉标题部分）
                text_content = self.extract_article_text(article)
                audio_path = self.generate_tts(text_content, temp_dir)
                if not audio_path:
                    self.video_log("TTS配音生成失败")
                    self.finish_video_task()
                    return

            if not self.video_is_running:
                self.finish_video_task()
                return

            # 2. 获取音频时长
            self.video_update_status("正在分析音频...")
            self.video_update_progress(10)
            audio_duration = self.get_audio_duration(audio_path)
            self.video_log(f"音频时长: {audio_duration:.1f} 秒")

            if not self.video_is_running:
                self.finish_video_task()
                return

            # 3. 生成字幕（Whisper识别）
            self.video_update_status("正在识别字幕...")
            self.video_update_progress(15)
            srt_path = self.transcribe_audio_for_video(audio_path, temp_dir)

            if not self.video_is_running:
                self.finish_video_task()
                return

            # 4. 收集视频素材
            self.video_update_status("正在收集视频素材...")
            self.video_update_progress(30)
            video_files = self.collect_videos_for_video(audio_duration + 5)

            if not video_files:
                self.video_log("错误：没有收集到视频素材")
                self.finish_video_task()
                return

            if not self.video_is_running:
                self.finish_video_task()
                return

            # 5. 合并视频素材
            self.video_update_status("正在合并视频素材...")
            self.video_update_progress(45)

            list_file = os.path.join(temp_dir, "videos.txt")
            with open(list_file, 'w', encoding='utf-8') as f:
                for vf in video_files:
                    vf_fixed = vf.replace('\\', '/')
                    f.write(f"file '{vf_fixed}'\n")

            merged_video = os.path.join(temp_dir, "merged.mp4")
            cmd = f'"{self.ffmpeg_path}" -y -f concat -safe 0 -i "{list_file}" -c copy "{merged_video}"'
            self.video_log("合并视频素材...")
            subprocess.run(cmd, shell=True, capture_output=True, text=True)

            # 6. 裁剪到音频长度
            self.video_update_status("正在裁剪视频...")
            self.video_update_progress(55)
            trimmed_video = os.path.join(temp_dir, "trimmed.mp4")

            # 根据是否静音视频素材选择不同命令
            if self.mute_video.get():
                # 静音视频素材
                cmd = f'"{self.ffmpeg_path}" -y -i "{merged_video}" -t {audio_duration} -an -c:v copy "{trimmed_video}"'
            else:
                # 保留视频原声
                cmd = f'"{self.ffmpeg_path}" -y -i "{merged_video}" -t {audio_duration} -c copy "{trimmed_video}"'
            subprocess.run(cmd, shell=True, capture_output=True)

            # 7. 添加音频（配音 + 可选BGM）
            self.video_update_status("正在添加音频...")
            self.video_update_progress(60)
            with_audio = os.path.join(temp_dir, "with_audio.mp4")

            # 获取音量设置
            voice_vol = int(self.voice_volume.get()) / 100

            if self.enable_bgm.get():
                # 添加BGM
                bgm_path = self.get_random_bgm()
                if bgm_path:
                    bgm_vol = int(self.bgm_volume.get()) / 100
                    self.video_log(f"添加背景音乐: {os.path.basename(bgm_path)}")
                    # 混合配音和BGM
                    filter_complex = f"[1:a]volume={voice_vol}[voice];[2:a]volume={bgm_vol},aloop=loop=-1:size=2e+09[bgm];[voice][bgm]amix=inputs=2:duration=first:dropout_transition=2[aout]"
                    cmd = f'"{self.ffmpeg_path}" -y -i "{trimmed_video}" -i "{audio_path}" -i "{bgm_path}" -filter_complex "{filter_complex}" -map 0:v -map "[aout]" -c:v copy -c:a aac -b:a 192k -t {audio_duration} "{with_audio}"'
                else:
                    self.video_log("未找到BGM文件，跳过背景音乐")
                    cmd = f'"{self.ffmpeg_path}" -y -i "{trimmed_video}" -i "{audio_path}" -filter_complex "[1:a]volume={voice_vol}[aout]" -map 0:v -map "[aout]" -c:v copy -c:a aac -b:a 192k "{with_audio}"'
            else:
                # 只添加配音
                if voice_vol != 1.0:
                    cmd = f'"{self.ffmpeg_path}" -y -i "{trimmed_video}" -i "{audio_path}" -filter_complex "[1:a]volume={voice_vol}[aout]" -map 0:v -map "[aout]" -c:v copy -c:a aac -b:a 192k "{with_audio}"'
                else:
                    cmd = f'"{self.ffmpeg_path}" -y -i "{trimmed_video}" -i "{audio_path}" -map 0:v:0 -map 1:a:0 -c:v copy -c:a aac -b:a 192k "{with_audio}"'

            result = subprocess.run(cmd, shell=True, capture_output=True, text=True)
            if result.returncode != 0:
                self.video_log(f"音频处理警告: {result.stderr[:200] if result.stderr else ''}")

            if not self.video_is_running:
                self.finish_video_task()
                return

            # 8. 烧录字幕（偏移1秒，补偿封面时长）
            self.video_update_status("正在烧录字幕...")
            self.video_update_progress(70)
            with_subtitle = self.burn_subtitles_for_video(with_audio, srt_path, temp_dir, time_offset=1)

            if not self.video_is_running:
                self.finish_video_task()
                return

            # 9. 生成封面
            self.video_update_status("正在生成封面...")
            self.video_update_progress(85)
            final_video = self.add_cover_for_video(with_subtitle, video_title, temp_dir)

            # 10. 生成最终文件
            self.video_update_status("正在生成最终视频...")
            self.video_update_progress(95)

            timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
            safe_title_file = re.sub(r'[\\/:*?"<>|]', '', video_title)[:20]
            final_output = os.path.join(output_path, f"{timestamp}_{safe_title_file}.mp4")

            import shutil
            shutil.copy(final_video, final_output)

            # 11. 清理临时文件
            self.video_log("清理临时文件...")
            try:
                shutil.rmtree(temp_dir)
            except:
                pass

            self.video_update_progress(100)
            self.video_log(f"\n视频制作完成！")
            self.video_log(f"输出文件: {final_output}")

            self.finish_video_task()
            messagebox.showinfo("完成", f"视频制作完成！\n保存位置: {final_output}")

        except Exception as e:
            self.video_log(f"错误: {str(e)}")
            import traceback
            self.video_log(traceback.format_exc())
            self.finish_video_task()

    def extract_article_text(self, article):
        """从文案中提取正文（去掉标题）"""
        lines = article.strip().split('\n')
        text_lines = []
        in_content = False
        for line in lines:
            if line.startswith('---') or line.startswith('═'):
                in_content = True
                continue
            if in_content and not line.startswith('【标题'):
                text_lines.append(line.strip())
        return '\n'.join(text_lines) if text_lines else article

    def generate_tts(self, text, temp_dir):
        """使用TTS生成配音"""
        try:
            import edge_tts
            import asyncio

            self.video_log("正在使用Edge TTS生成配音...")
            audio_path = os.path.join(temp_dir, "tts_audio.mp3")

            async def generate():
                communicate = edge_tts.Communicate(text, "zh-CN-YunxiNeural")
                await communicate.save(audio_path)

            asyncio.run(generate())
            self.video_log("TTS配音生成成功")
            return audio_path

        except ImportError:
            self.video_log("错误：未安装edge-tts，请运行: pip install edge-tts")
            return None
        except Exception as e:
            self.video_log(f"TTS生成错误: {str(e)}")
            return None

    def get_audio_duration(self, audio_path):
        """获取音频时长"""
        import subprocess
        cmd = f'"{self.ffprobe_path}" -v error -show_entries format=duration -of default=noprint_wrappers=1:nokey=1 "{audio_path}"'
        result = subprocess.run(cmd, shell=True, capture_output=True, text=True)
        return float(result.stdout.strip())

    def get_video_duration(self, video_path):
        """获取视频时长"""
        import subprocess
        cmd = f'"{self.ffprobe_path}" -v error -show_entries format=duration -of default=noprint_wrappers=1:nokey=1 "{video_path}"'
        result = subprocess.run(cmd, shell=True, capture_output=True, text=True)
        try:
            return float(result.stdout.strip())
        except:
            return 0

    def transcribe_audio_for_video(self, audio_path, temp_dir):
        """使用Whisper识别音频生成字幕"""
        try:
            from faster_whisper import WhisperModel

            self.video_log("正在加载Whisper模型...")

            if self.whisper_model is None:
                self.whisper_model = WhisperModel("small", device="cpu", compute_type="int8")

            self.video_log("开始识别音频...")
            segments, info = self.whisper_model.transcribe(audio_path, language="zh")

            srt_path = os.path.join(temp_dir, "subtitle.srt")
            with open(srt_path, 'w', encoding='utf-8') as f:
                for i, segment in enumerate(segments):
                    start_str = self.seconds_to_srt_time(segment.start)
                    end_str = self.seconds_to_srt_time(segment.end)
                    text = segment.text.strip()

                    f.write(f"{i+1}\n")
                    f.write(f"{start_str} --> {end_str}\n")
                    f.write(f"{text}\n\n")

                    if not self.video_is_running:
                        return None

            self.video_log("Whisper识别完成")
            return srt_path

        except ImportError:
            self.video_log("警告：未安装faster-whisper，将不添加字幕")
            return None
        except Exception as e:
            self.video_log(f"Whisper识别错误: {str(e)}")
            return None

    def seconds_to_srt_time(self, seconds):
        """秒数转SRT时间格式"""
        hours = int(seconds // 3600)
        minutes = int((seconds % 3600) // 60)
        secs = int(seconds % 60)
        millis = int((seconds % 1) * 1000)
        return f"{hours:02d}:{minutes:02d}:{secs:02d},{millis:03d}"

    def collect_videos_for_video(self, needed_duration):
        """收集视频素材"""
        self.video_log(f"需要收集 {needed_duration:.1f} 秒的视频素材...")

        collected_videos = []
        total_duration = 0

        subfolders = []
        if os.path.exists(self.video_source_path):
            for item in os.listdir(self.video_source_path):
                item_path = os.path.join(self.video_source_path, item)
                if os.path.isdir(item_path):
                    subfolders.append(item_path)

        if not subfolders:
            # 如果没有子文件夹，直接从根目录取
            subfolders = [self.video_source_path]

        self.video_log(f"找到 {len(subfolders)} 个素材文件夹")

        folder_video_count = {folder: 0 for folder in subfolders}

        while total_duration < needed_duration:
            found_new = False

            for folder in subfolders:
                if folder_video_count[folder] >= self.max_videos_per_folder:
                    continue

                video_files = []
                try:
                    for f in os.listdir(folder):
                        if f.lower().endswith(('.mp4', '.mov', '.avi', '.mkv')):
                            video_files.append(os.path.join(folder, f))
                except:
                    continue

                random.shuffle(video_files)

                for vf in video_files:
                    if vf not in collected_videos:
                        duration = self.get_video_duration(vf)
                        if duration > 0:
                            collected_videos.append(vf)
                            total_duration += duration
                            folder_video_count[folder] += 1
                            self.video_log(f"  + {os.path.basename(vf)} ({duration:.1f}秒)")
                            found_new = True
                            break

                if total_duration >= needed_duration:
                    break

            if not found_new:
                self.video_log("警告：素材不足")
                break

        self.video_log(f"共收集 {len(collected_videos)} 个视频，总时长 {total_duration:.1f} 秒")
        return collected_videos

    def get_color_code(self, color_name):
        """颜色名称转ffmpeg颜色代码"""
        colors = {"金色": "gold", "黄色": "yellow", "橙色": "orange", "红色": "red"}
        return colors.get(color_name, "gold")

    def get_preview_color(self, color_name):
        """颜色名称转Tkinter颜色代码"""
        colors = {"金色": "#FFD700", "黄色": "#FFFF00", "橙色": "#FFA500", "红色": "#FF0000"}
        return colors.get(color_name, "#FFD700")

    def update_subtitle_preview(self, event=None):
        """更新字幕预览"""
        if not hasattr(self, 'subtitle_preview_canvas'):
            return

        canvas = self.subtitle_preview_canvas
        canvas.delete("all")

        # 获取当前设置
        font_size = int(self.subtitle_size.get())
        color = self.get_preview_color(self.subtitle_color.get())
        position = self.subtitle_position.get()

        # 预览文字
        preview_text = "这是字幕预览效果"

        # 计算预览字号（缩小比例显示）
        preview_font_size = max(12, font_size // 3)

        # 根据位置确定Y坐标
        canvas_height = 80
        if position == "底部":
            y = canvas_height - 15
        elif position == "中部":
            y = canvas_height // 2
        else:  # 顶部
            y = 15

        # 绘制文字（带黑色描边效果）
        x = 200  # 居中
        # 描边
        for dx, dy in [(-1,-1), (-1,1), (1,-1), (1,1), (-2,0), (2,0), (0,-2), (0,2)]:
            canvas.create_text(x+dx, y+dy, text=preview_text, fill="black",
                             font=("SimHei", preview_font_size, "bold"))
        # 主文字
        canvas.create_text(x, y, text=preview_text, fill=color,
                         font=("SimHei", preview_font_size, "bold"))

    def get_ass_color(self, color_name):
        """颜色名称转ASS字幕颜色代码（BGR格式）"""
        # ASS颜色格式: &HBBGGRR (蓝绿红)
        colors = {
            "金色": "&H0000D7FF",
            "黄色": "&H0000FFFF",
            "橙色": "&H0000A5FF",
            "红色": "&H000000FF"
        }
        return colors.get(color_name, "&H0000D7FF")

    def burn_subtitles_for_video(self, video_path, srt_path, temp_dir, time_offset=0):
        """烧录字幕

        Args:
            time_offset: 字幕时间偏移（秒），用于补偿封面时长
        """
        import subprocess
        import shutil

        if not srt_path or not os.path.exists(srt_path):
            self.video_log("无字幕文件，跳过字幕烧录")
            return video_path

        # 如果有时间偏移，调整字幕时间
        actual_srt_path = srt_path
        if time_offset > 0:
            actual_srt_path = os.path.join(temp_dir, "subtitle_offset.srt")
            self.offset_srt_time(srt_path, actual_srt_path, time_offset)
            self.video_log(f"字幕时间已偏移 {time_offset} 秒（补偿封面时长）")

        font_size = self.subtitle_size.get()
        font_color = self.get_ass_color(self.subtitle_color.get())
        position = self.subtitle_position.get()

        if position == "底部":
            margin_v = 30
            alignment = 2
        elif position == "中部":
            margin_v = 0
            alignment = 5
        else:
            margin_v = 30
            alignment = 8

        with_subtitle = os.path.join(temp_dir, "with_subtitle.mp4")

        # 复制字幕到临时目录，用简单文件名避免路径问题
        simple_srt = os.path.join(temp_dir, "sub.srt")
        shutil.copy(actual_srt_path, simple_srt)
        srt_escaped = simple_srt.replace('\\', '/').replace(':', '\\:')

        subtitle_filter = f"subtitles='{srt_escaped}':force_style='FontSize={font_size},PrimaryColour={font_color},OutlineColour=&H00000000,BorderStyle=1,Outline=2,Alignment={alignment},MarginV={margin_v}'"

        cmd = f'"{self.ffmpeg_path}" -y -i "{video_path}" -vf "{subtitle_filter}" -c:a aac -b:a 192k "{with_subtitle}"'
        self.video_log("正在烧录字幕...")
        result = subprocess.run(cmd, shell=True, capture_output=True, text=True)

        if result.returncode != 0:
            self.video_log(f"字幕烧录失败: {result.stderr[:200] if result.stderr else ''}")
            self.video_log("使用无字幕版本")
            return video_path

        self.video_log("字幕烧录成功")
        return with_subtitle

    def offset_srt_time(self, input_srt, output_srt, offset_seconds):
        """偏移SRT字幕时间"""
        def parse_time(time_str):
            parts = time_str.replace(',', ':').split(':')
            h, m, s, ms = int(parts[0]), int(parts[1]), int(parts[2]), int(parts[3])
            return h * 3600 + m * 60 + s + ms / 1000

        def format_time(seconds):
            h = int(seconds // 3600)
            m = int((seconds % 3600) // 60)
            s = int(seconds % 60)
            ms = int((seconds % 1) * 1000)
            return f"{h:02d}:{m:02d}:{s:02d},{ms:03d}"

        with open(input_srt, 'r', encoding='utf-8') as f:
            content = f.read()

        pattern = r'(\d{2}:\d{2}:\d{2},\d{3}) --> (\d{2}:\d{2}:\d{2},\d{3})'
        def replace_time(match):
            start = parse_time(match.group(1)) + offset_seconds
            end = parse_time(match.group(2)) + offset_seconds
            return f"{format_time(start)} --> {format_time(end)}"

        new_content = re.sub(pattern, replace_time, content)
        with open(output_srt, 'w', encoding='utf-8') as f:
            f.write(new_content)

    def add_cover_for_video(self, video_path, video_title, temp_dir):
        """添加封面"""
        import subprocess

        cover_font_size = self.cover_font_size.get()
        cover_color = self.get_color_code(self.cover_color.get())
        cover_bg = self.cover_bg.get()

        probe_cmd = f'"{self.ffprobe_path}" -v error -select_streams v:0 -show_entries stream=width,height -of csv=p=0 "{video_path}"'
        probe_result = subprocess.run(probe_cmd, shell=True, capture_output=True, text=True)
        try:
            width, height = map(int, probe_result.stdout.strip().split(','))
        except:
            width, height = 1920, 1080

        cover_video = os.path.join(temp_dir, "cover.mp4")
        cover_duration = 1  # 封面时长1秒

        safe_title = video_title.replace("'", "").replace(":", "\\:").replace("\\", "/")
        if len(safe_title) > 12:
            mid = len(safe_title) // 2
            safe_title = safe_title[:mid] + "\\n" + safe_title[mid:]

        title_filter = f"drawtext=text='{safe_title}':fontfile='C\\:/Windows/Fonts/msyh.ttc':fontsize={cover_font_size}:fontcolor={cover_color}:borderw=4:bordercolor=black:x=(w-text_w)/2:y=(h-text_h)/2"

        if cover_bg == "视频首帧":
            first_frame = os.path.join(temp_dir, "first_frame.jpg")
            cmd = f'"{self.ffmpeg_path}" -y -i "{video_path}" -vframes 1 -q:v 2 "{first_frame}"'
            subprocess.run(cmd, shell=True, capture_output=True)
            # 生成带静音音轨的封面视频
            cmd = f'"{self.ffmpeg_path}" -y -loop 1 -i "{first_frame}" -f lavfi -i anullsrc=r=44100:cl=stereo -vf "{title_filter}" -t {cover_duration} -r 30 -c:v libx264 -pix_fmt yuv420p -c:a aac -shortest "{cover_video}"'
        else:
            # 黑色背景 + 静音音轨
            cmd = f'"{self.ffmpeg_path}" -y -f lavfi -i color=c=black:s={width}x{height}:d={cover_duration}:r=30 -f lavfi -i anullsrc=r=44100:cl=stereo -vf "{title_filter}" -c:v libx264 -pix_fmt yuv420p -c:a aac -t {cover_duration} -shortest "{cover_video}"'

        self.video_log("生成封面视频...")
        result = subprocess.run(cmd, shell=True, capture_output=True, text=True)

        if result.returncode != 0:
            self.video_log(f"封面生成失败: {result.stderr[:200] if result.stderr else ''}")
            self.video_log("使用无封面版本")
            return video_path

        # 重新编码主视频确保格式一致（保留音频）
        main_encoded = os.path.join(temp_dir, "main_encoded.mp4")
        cmd = f'"{self.ffmpeg_path}" -y -i "{video_path}" -c:v libx264 -preset fast -crf 23 -c:a aac -ar 44100 -b:a 192k "{main_encoded}"'
        self.video_log("重新编码主视频...")
        subprocess.run(cmd, shell=True, capture_output=True)

        concat_list = os.path.join(temp_dir, "concat_list.txt")
        with open(concat_list, 'w', encoding='utf-8') as f:
            f.write(f"file '{cover_video.replace(chr(92), '/')}'\n")
            f.write(f"file '{main_encoded.replace(chr(92), '/')}'\n")

        final_video = os.path.join(temp_dir, "final.mp4")
        cmd = f'"{self.ffmpeg_path}" -y -f concat -safe 0 -i "{concat_list}" -c copy "{final_video}"'
        result = subprocess.run(cmd, shell=True, capture_output=True, text=True)

        if result.returncode != 0:
            self.video_log("封面合并失败，使用无封面版本")
            return video_path

        self.video_log("封面合并成功")
        return final_video

    def finish_video_task(self):
        """完成视频制作任务"""
        self.video_is_running = False
        self.video_make_btn.config(state=tk.NORMAL)
        self.video_stop_btn.config(state=tk.DISABLED)
        self.video_update_status("完成")

    def _create_model_config(self, parent, prefix):
        """创建模型配置UI组件"""
        # URL
        url_frame = ttk.Frame(parent)
        url_frame.pack(fill=tk.X, pady=2)
        ttk.Label(url_frame, text="URL:", width=12).pack(side=tk.LEFT)
        url_var = tk.StringVar(value=self.config.get(f"{prefix}_url", DEFAULT_CONFIG.get(f"{prefix}_url", "")))
        setattr(self, f"{prefix}_url", url_var)
        ttk.Entry(url_frame, textvariable=url_var, width=60).pack(side=tk.LEFT, fill=tk.X, expand=True)

        # Key
        key_frame = ttk.Frame(parent)
        key_frame.pack(fill=tk.X, pady=2)
        ttk.Label(key_frame, text="API Key:", width=12).pack(side=tk.LEFT)
        key_var = tk.StringVar(value=self.config.get(f"{prefix}_key", DEFAULT_CONFIG.get(f"{prefix}_key", "")))
        setattr(self, f"{prefix}_key", key_var)
        ttk.Entry(key_frame, textvariable=key_var, width=60).pack(side=tk.LEFT, fill=tk.X, expand=True)

        # Model + Max Tokens + 测试按钮
        model_frame = ttk.Frame(parent)
        model_frame.pack(fill=tk.X, pady=2)
        ttk.Label(model_frame, text="模型:", width=12).pack(side=tk.LEFT)
        model_var = tk.StringVar(value=self.config.get(f"{prefix}_model", DEFAULT_CONFIG.get(f"{prefix}_model", "")))
        setattr(self, f"{prefix}_model", model_var)
        ttk.Entry(model_frame, textvariable=model_var, width=35).pack(side=tk.LEFT)
        ttk.Label(model_frame, text="  Max Tokens:").pack(side=tk.LEFT)
        tokens_var = tk.StringVar(value=str(self.config.get(f"{prefix}_max_tokens", DEFAULT_CONFIG.get(f"{prefix}_max_tokens", 16000))))
        setattr(self, f"{prefix}_max_tokens", tokens_var)
        ttk.Entry(model_frame, textvariable=tokens_var, width=10).pack(side=tk.LEFT)

        # 测试按钮
        ttk.Button(model_frame, text="测试连接", width=10,
                   command=lambda p=prefix: self.test_api_connection(p)).pack(side=tk.LEFT, padx=10)

    def test_api_connection(self, prefix):
        """测试API连接，自动尝试不同格式"""
        url = getattr(self, f"{prefix}_url").get().strip()
        key = getattr(self, f"{prefix}_key").get().strip()
        model = getattr(self, f"{prefix}_model").get().strip()

        if not url or not key or not model:
            messagebox.showwarning("提示", "请先填写完整的URL、API Key和模型名称")
            return

        # 在新线程中测试，避免界面卡住
        def do_test():
            self.log(f"开始测试 {prefix} 连接...")

            # 尝试不同的API格式
            formats_to_try = [
                ("OpenAI格式", self._test_openai_format),
                ("OpenAI格式(无v1)", self._test_openai_format_no_v1),
                ("Anthropic格式", self._test_anthropic_format),
            ]

            for format_name, test_func in formats_to_try:
                self.log(f"  尝试 {format_name}...")
                success, message, working_url = test_func(url, key, model)
                if success:
                    self.log(f"  ✓ {format_name} 成功！")
                    # 如果URL被修正了，更新配置
                    if working_url and working_url != url:
                        getattr(self, f"{prefix}_url").set(working_url)
                        self.log(f"  已自动修正URL为: {working_url}")
                    self.root.after(0, lambda: messagebox.showinfo("测试成功",
                        f"连接成功！\n\n格式: {format_name}\nURL: {working_url or url}\n模型: {model}\n\n响应: {message}"))
                    return
                else:
                    self.log(f"  ✗ {format_name} 失败: {message[:50]}...")

            self.log("所有格式均测试失败")
            self.root.after(0, lambda: messagebox.showerror("测试失败",
                f"所有API格式均无法连接\n\n请检查:\n1. URL是否正确\n2. API Key是否有效\n3. 模型名称是否正确"))

        import threading
        threading.Thread(target=do_test, daemon=True).start()

    def _test_openai_format(self, base_url, api_key, model):
        """测试OpenAI兼容格式"""
        try:
            url = f"{base_url.rstrip('/')}/chat/completions"
            headers = {
                "Content-Type": "application/json",
                "Authorization": f"Bearer {api_key}"
            }
            data = {
                "model": model,
                "messages": [{"role": "user", "content": "你好，请回复：测试成功"}],
                "max_tokens": 50
            }
            response = requests.post(url, headers=headers, json=data, timeout=30)
            if response.status_code == 200:
                result = response.json()
                content = result.get("choices", [{}])[0].get("message", {}).get("content", "")
                return True, content, base_url
            else:
                return False, f"状态码{response.status_code}: {response.text[:100]}", None
        except Exception as e:
            return False, str(e), None

    def _test_openai_format_no_v1(self, base_url, api_key, model):
        """测试OpenAI格式，自动添加/v1"""
        # 如果URL已经包含/v1，跳过
        if "/v1" in base_url:
            return False, "URL已包含/v1，跳过此测试", None

        new_url = f"{base_url.rstrip('/')}/v1"
        success, message, _ = self._test_openai_format(new_url, api_key, model)
        if success:
            return True, message, new_url
        return False, message, None

    def _test_anthropic_format(self, base_url, api_key, model):
        """测试Anthropic原生格式"""
        try:
            # 尝试Anthropic格式的endpoint
            url = f"{base_url.rstrip('/')}/messages"
            if "/v1" not in url:
                url = f"{base_url.rstrip('/')}/v1/messages"

            headers = {
                "Content-Type": "application/json",
                "x-api-key": api_key,
                "anthropic-version": "2023-06-01"
            }
            data = {
                "model": model,
                "messages": [{"role": "user", "content": "你好，请回复：测试成功"}],
                "max_tokens": 50
            }
            response = requests.post(url, headers=headers, json=data, timeout=30)
            if response.status_code == 200:
                result = response.json()
                content = result.get("content", [{}])[0].get("text", "")
                return True, content, base_url
            else:
                return False, f"状态码{response.status_code}: {response.text[:100]}", None
        except Exception as e:
            return False, str(e), None

    def on_api_tab_change(self, event=None):
        """API标签页切换 - 自动更新流式/非流式设置"""
        selected_tab = self.api_notebook.index(self.api_notebook.select())
        # 0 = 流式, 1 = 非流式
        use_stream = (selected_tab == 0)
        self.use_stream.set(use_stream)
        self.config["use_stream"] = use_stream
        save_config(self.config)
        mode = "流式" if use_stream else "非流式"
        self.log(f"已切换到{mode}调用模式")
        # 更新模式提示
        if hasattr(self, 'api_mode_label'):
            self._update_api_mode_label()

    def on_stream_change(self):
        """流式开关改变 - 自动保存（保留兼容）"""
        self.config["use_stream"] = self.use_stream.get()
        save_config(self.config)
        mode = "流式" if self.use_stream.get() else "非流式"
        self.log(f"已切换到{mode}调用模式")

    def _parse_similarity_threshold(self):
        """解析并校验相似度阈值"""
        raw_value = self.similarity_threshold_var.get().strip()
        value = float(raw_value)
        if value < 0.50 or value > 0.90:
            raise ValueError("相似度阈值必须在 0.50 ~ 0.90 之间")
        return round(value, 2)

    def save_api_config(self):
        """保存API配置"""
        try:
            self.config["use_stream"] = self.use_stream.get()
            threshold_value = self._parse_similarity_threshold()
            self.config["similarity_threshold"] = threshold_value
            self.similarity_threshold = threshold_value
            # 流式配置
            self.config["stream_main_url"] = self.stream_main_url.get().strip()
            self.config["stream_main_key"] = self.stream_main_key.get().strip()
            self.config["stream_main_model"] = self.stream_main_model.get().strip()
            self.config["stream_main_max_tokens"] = int(self.stream_main_max_tokens.get())
            self.config["stream_backup_url"] = self.stream_backup_url.get().strip()
            self.config["stream_backup_key"] = self.stream_backup_key.get().strip()
            self.config["stream_backup_model"] = self.stream_backup_model.get().strip()
            self.config["stream_backup_max_tokens"] = int(self.stream_backup_max_tokens.get())
            # 非流式配置
            self.config["non_stream_main_url"] = self.non_stream_main_url.get().strip()
            self.config["non_stream_main_key"] = self.non_stream_main_key.get().strip()
            self.config["non_stream_main_model"] = self.non_stream_main_model.get().strip()
            self.config["non_stream_main_max_tokens"] = int(self.non_stream_main_max_tokens.get())
            self.config["non_stream_backup_url"] = self.non_stream_backup_url.get().strip()
            self.config["non_stream_backup_key"] = self.non_stream_backup_key.get().strip()
            self.config["non_stream_backup_model"] = self.non_stream_backup_model.get().strip()
            self.config["non_stream_backup_max_tokens"] = int(self.non_stream_backup_max_tokens.get())
            save_config(self.config)
            messagebox.showinfo("成功", "API配置已保存")
        except Exception as e:
            messagebox.showerror("错误", f"保存配置失败：{e}")

    def reset_api_config(self):
        """重置为默认配置"""
        if messagebox.askyesno("确认", "确定要重置为默认配置吗？"):
            self.config = DEFAULT_CONFIG.copy()
            self.use_stream.set(self.config["use_stream"])
            # 流式配置
            self.stream_main_url.set(self.config["stream_main_url"])
            self.stream_main_key.set(self.config["stream_main_key"])
            self.stream_main_model.set(self.config["stream_main_model"])
            self.stream_main_max_tokens.set(str(self.config["stream_main_max_tokens"]))
            self.stream_backup_url.set(self.config["stream_backup_url"])
            self.stream_backup_key.set(self.config["stream_backup_key"])
            self.stream_backup_model.set(self.config["stream_backup_model"])
            self.stream_backup_max_tokens.set(str(self.config["stream_backup_max_tokens"]))
            # 非流式配置
            self.non_stream_main_url.set(self.config["non_stream_main_url"])
            self.non_stream_main_key.set(self.config["non_stream_main_key"])
            self.non_stream_main_model.set(self.config["non_stream_main_model"])
            self.non_stream_main_max_tokens.set(str(self.config["non_stream_main_max_tokens"]))
            self.non_stream_backup_url.set(self.config["non_stream_backup_url"])
            self.non_stream_backup_key.set(self.config["non_stream_backup_key"])
            self.non_stream_backup_model.set(self.config["non_stream_backup_model"])
            self.non_stream_backup_max_tokens.set(str(self.config["non_stream_backup_max_tokens"]))
            self.similarity_threshold_var.set(f"{float(self.config.get('similarity_threshold', 0.76)):.2f}")
            self.similarity_threshold = float(self.config.get("similarity_threshold", 0.76))
            save_config(self.config)
            messagebox.showinfo("成功", "已重置为默认配置")

    def on_flow_type_change(self, event=None):
        """引流类型改变时的处理"""
        flow_type = self.flow_type.get()

        if flow_type == "带货引流":
            self.daihuo_frame.pack(fill=tk.X, pady=5, before=self.btn_frame_container)
        else:
            self.daihuo_frame.pack_forget()

        if flow_type == "纯夸赞不引流":
            self.yinliu_frame.pack_forget()
        else:
            self.yinliu_frame.pack(fill=tk.X, pady=5, before=self.daihuo_frame if flow_type == "带货引流" else self.btn_frame_container)

        # 切换引流类型时更新话术下拉框
        self.update_yinliu_combo()

    def update_yinliu_combo(self):
        """更新话术下拉框"""
        flow_type = self.flow_type.get()
        if flow_type == "纯夸赞不引流":
            return

        # 获取当前类型的话术列表
        templates = self.config.get("yinliu_templates", {}).get(flow_type, [])

        # 构建下拉框选项（显示前20个字）
        options = ["-- 自定义（在下方输入框填写）--"]
        for i, tpl in enumerate(templates):
            preview = tpl[:20].replace('\n', ' ') + "..." if len(tpl) > 20 else tpl.replace('\n', ' ')
            options.append(f"{i+1}. {preview}")

        self.yinliu_combo['values'] = options
        self.yinliu_combo.current(0)

    def on_yinliu_select(self, event=None):
        """选择话术时填充到文本框"""
        flow_type = self.flow_type.get()
        selection = self.yinliu_combo.current()

        if selection <= 0:  # 选择了"不使用"
            return

        # 获取对应的话术内容
        templates = self.config.get("yinliu_templates", {}).get(flow_type, [])
        if selection - 1 < len(templates):
            template = templates[selection - 1]
            # 填充到文本框
            self.yinliu_text.delete("1.0", tk.END)
            self.yinliu_text.insert("1.0", template)

    def save_yinliu_template(self):
        """保存当前话术到列表"""
        flow_type = self.flow_type.get()
        if flow_type == "纯夸赞不引流":
            messagebox.showwarning("提示", "纯夸赞模式不需要引流话术")
            return

        content = self.yinliu_text.get("1.0", tk.END).strip()
        if not content:
            messagebox.showwarning("提示", "请先输入话术内容")
            return

        # 确保配置中有话术库
        if "yinliu_templates" not in self.config:
            self.config["yinliu_templates"] = {"置顶引流": [], "橱窗引流": [], "带货引流": []}
        if flow_type not in self.config["yinliu_templates"]:
            self.config["yinliu_templates"][flow_type] = []

        # 检查是否已存在
        templates = self.config["yinliu_templates"][flow_type]
        if content in templates:
            messagebox.showinfo("提示", "该话术已存在")
            return

        # 添加到列表
        templates.append(content)
        save_config(self.config)

        # 更新下拉框
        self.update_yinliu_combo()
        messagebox.showinfo("成功", f"话术已保存到【{flow_type}】列表")

    def delete_yinliu_template(self):
        """删除选中的话术"""
        flow_type = self.flow_type.get()
        selection = self.yinliu_combo.current()

        if selection <= 0:
            messagebox.showwarning("提示", "请先选择要删除的话术")
            return

        templates = self.config.get("yinliu_templates", {}).get(flow_type, [])
        if selection - 1 < len(templates):
            if messagebox.askyesno("确认", "确定要删除选中的话术吗？"):
                del templates[selection - 1]
                save_config(self.config)
                self.update_yinliu_combo()
                self.yinliu_text.delete("1.0", tk.END)
                messagebox.showinfo("成功", "话术已删除")

    def clear_yinliu_text(self):
        """清空话术输入框"""
        self.yinliu_text.delete("1.0", tk.END)
        self.yinliu_combo.current(0)

    def on_input_mode_change(self):
        """切换参考文案输入方式"""
        mode = self.input_mode.get()
        if mode == "file":
            self.paste_input_frame.pack_forget()
            self.file_input_frame.pack(fill=tk.X)
        else:
            self.file_input_frame.pack_forget()
            self.paste_input_frame.pack(fill=tk.X)

    def select_input_file(self):
        file_path = filedialog.askopenfilename(
            title="选择参考文案文件",
            filetypes=[("文本文件", "*.txt"), ("所有文件", "*.*")]
        )
        if file_path:
            self.input_path.set(file_path)

    def select_input_folder(self):
        folder_path = filedialog.askdirectory(title="选择参考文案文件夹")
        if folder_path:
            self.input_path.set(folder_path)

    def select_output_folder(self):
        folder_path = filedialog.askdirectory(title="选择输出保存文件夹")
        if folder_path:
            self.output_path.set(folder_path)

    def select_txt_output_folder(self):
        """选择TXT保存文件夹"""
        folder_path = filedialog.askdirectory(title="选择TXT保存文件夹")
        if folder_path:
            self.txt_output_path.set(folder_path)

    def select_voice_input_folder(self):
        """选择语音合成的文案输入目录"""
        folder_path = filedialog.askdirectory(title="选择文案目录")
        if folder_path:
            self.voice_input_path.set(folder_path)

    def select_voice_output_folder(self):
        """选择语音合成的输出目录"""
        folder_path = filedialog.askdirectory(title="选择配音输出目录")
        if folder_path:
            self.voice_output_path.set(folder_path)

    def log(self, message):
        """添加日志"""
        timestamp = datetime.now().strftime("%H:%M:%S")
        log_line = f"[{timestamp}] {message}\n"
        self.log_text.insert(tk.END, log_line)
        self.log_text.see(tk.END)
        self.root.update()

        try:
            log_file = os.path.join(self.output_path.get() or ".", "debug_log.txt")
            with open(log_file, 'a', encoding='utf-8') as f:
                f.write(log_line)
        except:
            pass

    def update_status(self, message):
        """更新状态"""
        self.status_label.config(text=message)
        self.root.update()

    def update_progress(self, value):
        """更新进度条"""
        self.progress_var.set(value)
        self.root.update()

    def start_generate(self):
        """开始生成"""
        input_mode = self.input_mode.get()
        output_path = self.output_path.get()

        # 根据输入模式验证
        if input_mode == "paste":
            paste_content = self.paste_text.get("1.0", tk.END).strip()
            if not paste_content:
                messagebox.showerror("错误", "请粘贴参考文案内容")
                return
        else:
            input_path = self.input_path.get()
            if not input_path:
                messagebox.showerror("错误", "请选择参考文案路径")
                return
            if not os.path.exists(input_path):
                messagebox.showerror("错误", f"参考文案路径不存在：{input_path}")
                return

        if not output_path:
            messagebox.showerror("错误", "请选择输出保存路径")
            return

        if self.flow_type.get() == "带货引流":
            if not self.product_name.get().strip():
                messagebox.showerror("错误", "请填写带货商品名称")
                return
            if not self.product_material.get("1.0", tk.END).strip():
                messagebox.showerror("错误", "请填写产品素材/介绍")
                return

        os.makedirs(output_path, exist_ok=True)

        self.start_btn.config(state=tk.DISABLED)
        self.stop_btn.config(state=tk.NORMAL)
        self.is_running = True

        self.log_text.delete(1.0, tk.END)

        thread = threading.Thread(target=self.generate_task)
        thread.daemon = True
        thread.start()

    def start_generate_txt(self):
        """开始生成并保存为TXT（每篇单独保存，用标题命名）"""
        input_mode = self.input_mode.get()
        txt_output_path = self.txt_output_path.get()

        # 根据输入模式验证
        if input_mode == "paste":
            paste_content = self.paste_text.get("1.0", tk.END).strip()
            if not paste_content:
                messagebox.showerror("错误", "请粘贴参考文案内容")
                return
        else:
            input_path = self.input_path.get()
            if not input_path:
                messagebox.showerror("错误", "请选择参考文案路径")
                return
            if not os.path.exists(input_path):
                messagebox.showerror("错误", f"参考文案路径不存在：{input_path}")
                return

        if not txt_output_path:
            messagebox.showerror("错误", "请选择TXT保存路径")
            return

        if self.flow_type.get() == "带货引流":
            if not self.product_name.get().strip():
                messagebox.showerror("错误", "请填写带货商品名称")
                return
            if not self.product_material.get("1.0", tk.END).strip():
                messagebox.showerror("错误", "请填写产品素材/介绍")
                return

        os.makedirs(txt_output_path, exist_ok=True)

        self.start_btn.config(state=tk.DISABLED)
        self.start_txt_btn.config(state=tk.DISABLED)
        self.stop_btn.config(state=tk.NORMAL)
        self.is_running = True
        self.txt_generate_success = False  # 初始化为失败状态

        self.log_text.delete(1.0, tk.END)

        thread = threading.Thread(target=self.generate_txt_task)
        thread.daemon = True
        thread.start()

    def generate_txt_task(self):
        """生成TXT任务主函数"""
        try:
            txt_output_path = self.txt_output_path.get()
            flow_type = self.flow_type.get()

            # 读取参考文案
            self.log("正在读取参考文案...")
            input_mode = self.input_mode.get()

            if input_mode == "paste":
                paste_content = self.paste_text.get("1.0", tk.END).strip()
                if not paste_content:
                    self.log("错误：请粘贴参考文案内容")
                    self.finish_txt_task()
                    return
                files_content = [("粘贴文案", paste_content)]
            else:
                input_path = self.input_path.get()
                if os.path.isfile(input_path):
                    with open(input_path, 'r', encoding='utf-8') as f:
                        content = f.read()
                    files_content = [(os.path.basename(input_path), content)]
                else:
                    files_content = []
                    for fname in os.listdir(input_path):
                        if fname.endswith('.txt'):
                            fpath = os.path.join(input_path, fname)
                            with open(fpath, 'r', encoding='utf-8') as f:
                                files_content.append((fname, f.read()))

            if not files_content:
                self.log("错误：没有找到任何文案文件")
                self.finish_txt_task()
                return

            # 读取引流素材
            yinliu_content = self.yinliu_text.get("1.0", tk.END).strip()

            # 带货信息
            product_name = ""
            product_material = ""
            if flow_type == "带货引流":
                product_name = self.product_name.get().strip()
                product_material = self.product_material.get("1.0", tk.END).strip()

            total_files = len(files_content)
            for idx, (fname, content) in enumerate(files_content):
                if not self.is_running:
                    break

                self.log(f"\n{'='*50}")
                self.log(f"处理文件 [{idx+1}/{total_files}]: {fname}")
                self.update_progress((idx / total_files) * 100)

                # 分割参考文案
                articles = self.parse_articles(content)
                self.log(f"识别到 {len(articles)} 篇参考文案")

                # 把参考文案追加到素材库
                for article in articles:
                    self.append_reference_to_library(article, flow_type)

                # 保存最后一次生成的信息
                self.last_articles = articles
                self.last_flow_type = flow_type
                self.last_yinliu_content = yinliu_content
                self.last_product_name = product_name
                self.last_product_material = product_material

                # 用于记录已使用的标题，确保3篇不重复
                used_titles = []
                # 记录成功和失败的篇数
                success_count = 0
                failed_list = []

                for art_idx, article in enumerate(articles):
                    if not self.is_running:
                        break

                    self.log(f"\n--- 处理第 {art_idx+1} 篇参考文案 ---")
                    self.update_status(f"正在生成第 {art_idx+1} 篇...")

                    # 生成仿写文案
                    result = self.generate_document(
                        article, flow_type, yinliu_content,
                        product_name, product_material
                    )

                    if result:
                        # 保存为TXT，用标题命名
                        self.save_as_txt(result, txt_output_path, used_titles)
                        self.txt_generate_success = True
                        success_count += 1
                    else:
                        self.log(f"第 {art_idx+1} 篇生成失败")
                        failed_list.append(art_idx + 1)

                # 输出汇总结果
                self.log(f"\n{'='*50}")
                self.log(f"【生成结果汇总】")
                self.log(f"成功: {success_count} 篇")
                if failed_list:
                    self.log(f"失败: {len(failed_list)} 篇，分别是第 {', '.join(map(str, failed_list))} 篇")
                else:
                    self.log(f"失败: 0 篇")

            self.update_progress(100)
            self.log("\n" + "="*50)
            self.log("全部处理完成！")
            self.finish_txt_task()

        except Exception as e:
            self.log(f"任务出错: {str(e)}")
            import traceback
            self.log(traceback.format_exc())
            self.txt_generate_success = False
            self.finish_txt_task()

    def split_articles_from_result(self, content):
        """把生成的大文案拆分成多篇，返回列表 [(纯正文, [标题列表]), ...]"""
        articles = []
        # 按【第X篇】分割，支持多种格式
        pattern = r'═+\s*\n*【第[一二三四五六七八九十\d]+篇】\s*\n*═*'
        parts = re.split(pattern, content)

        for part in parts:
            part = part.strip()
            if not part or len(part) < 100:
                continue

            # 提取该篇的标题和正文
            titles = []
            lines = part.split('\n')
            in_title_section = False
            content_start_idx = 0

            for i, line in enumerate(lines):
                line_stripped = line.strip()

                # 方式1: 【标题】后面跟着多行标题
                if line_stripped.startswith('【标题】') or line_stripped == '【标题】':
                    in_title_section = True
                    after_tag = line_stripped.replace('【标题】', '').strip()
                    if after_tag:
                        titles.append(after_tag)
                    continue

                # 方式2: 【标题1】【标题2】等格式（备用模型常用）
                title_match = re.match(r'【标题[1-5一二三四五]】[：:]?\s*(.+)', line_stripped)
                if title_match:
                    title_content = title_match.group(1).strip()
                    if title_content and len(title_content) > 2:
                        titles.append(title_content)
                    if not in_title_section:
                        in_title_section = True
                    continue

                # 方式3: 标题1：xxx 或 标题一：xxx 格式
                title_match2 = re.match(r'标题[1-5一二三四五][：:]\s*(.+)', line_stripped)
                if title_match2:
                    title_content = title_match2.group(1).strip()
                    if title_content and len(title_content) > 2:
                        titles.append(title_content)
                    if not in_title_section:
                        in_title_section = True
                    continue

                if in_title_section:
                    if line_stripped == '---' or line_stripped.startswith('---'):
                        # 找到分隔线，正文从下一行开始
                        in_title_section = False
                        content_start_idx = i + 1
                    elif line_stripped and not line_stripped.startswith('═') and len(titles) < 5:
                        # 普通标题行（无前缀）
                        clean_title = re.sub(r'^[\d]+[.、\s]*', '', line_stripped).strip()
                        if clean_title and len(clean_title) > 2 and not clean_title.startswith('【'):
                            titles.append(clean_title)

            # 提取纯正文（---之后的内容）
            if content_start_idx > 0:
                body_lines = lines[content_start_idx:]
                # 去掉开头的空行
                while body_lines and not body_lines[0].strip():
                    body_lines.pop(0)
                body_content = '\n'.join(body_lines).strip()
            else:
                # 如果没找到---，尝试去掉标题部分，找第一个非标题行
                body_lines = []
                found_body = False
                for line in lines:
                    line_stripped = line.strip()
                    # 跳过标题相关的行
                    if re.match(r'【标题[1-5一二三四五]?】', line_stripped):
                        continue
                    if re.match(r'标题[1-5一二三四五][：:]', line_stripped):
                        continue
                    if line_stripped == '---':
                        found_body = True
                        continue
                    if found_body or (line_stripped and len(line_stripped) > 30):
                        found_body = True
                        body_lines.append(line)
                body_content = '\n'.join(body_lines).strip() if body_lines else part

            articles.append((body_content, titles))

        return articles

    def save_as_txt(self, content, output_path, used_titles):
        """把生成的文案拆分成多篇，每篇单独保存为TXT，用标题命名"""
        try:
            # 拆分成多篇
            articles = self.split_articles_from_result(content)

            if not articles:
                self.log("警告：未能拆分出文章，保存整个内容")
                timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
                filename = f"{timestamp}.txt"
                filepath = os.path.join(output_path, filename)
                with open(filepath, 'w', encoding='utf-8') as f:
                    f.write(content)
                self.log(f"TXT已保存: {filename}")
                return [filepath]

            saved_files = []
            self.log(f"识别到 {len(articles)} 篇文章，开始分别保存...")
            existing_titles = self.get_generated_title_set()

            for idx, (article_content, titles) in enumerate(articles):
                selected_title = ""
                if not titles:
                    self.log(f"第{idx+1}篇未提取到标题，使用时间戳命名")
                    timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
                    filename = f"{timestamp}_{idx+1}.txt"
                    selected_title = f"未提取标题_{timestamp}_{idx+1}"
                else:
                    # 输出该篇所有标题
                    self.log(f"第{idx+1}篇提取到 {len(titles)} 个标题：")
                    for t_idx, t in enumerate(titles):
                        self.log(f"  标题{t_idx+1}: {t}")

                    # 从标题中随机选一个，避免与本次和历史已用标题重复
                    available_titles = [t for t in titles if t not in used_titles and t not in existing_titles]
                    if not available_titles:
                        available_titles = [t for t in titles if t not in used_titles]
                    if not available_titles:
                        available_titles = titles

                    selected_title = random.choice(available_titles)
                    self.log(f"  → 随机选中: {selected_title}")
                    used_titles.append(selected_title)
                    existing_titles.add(selected_title)

                    # 清理文件名中的非法字符
                    safe_title = re.sub(r'[\\/:*?"<>|]', '', selected_title)
                    safe_title = safe_title[:50]
                    filename = f"{safe_title}.txt"

                filepath = os.path.join(output_path, filename)

                # 如果文件已存在，添加时间戳
                if os.path.exists(filepath):
                    timestamp = datetime.now().strftime("%H%M%S")
                    name, ext = os.path.splitext(filename)
                    filename = f"{name}_{timestamp}{ext}"
                    filepath = os.path.join(output_path, filename)

                with open(filepath, 'w', encoding='utf-8') as f:
                    f.write(article_content)

                # 生成文案入库Excel，便于长期去重（即使txt被移动）
                flow_type = self.flow_type.get() if hasattr(self, "flow_type") else ""
                self.append_generated_to_library(flow_type, selected_title or filename, article_content)

                self.log(f"第{idx+1}篇TXT已保存: {filename}")
                saved_files.append(filepath)

            return saved_files

        except Exception as e:
            self.log(f"保存TXT失败: {str(e)}")
            import traceback
            self.log(traceback.format_exc())
            return None

    def finish_txt_task(self):
        """完成TXT生成任务，恢复UI状态"""
        self.is_running = False
        self.start_btn.config(state=tk.NORMAL)
        self.start_txt_btn.config(state=tk.NORMAL)
        self.stop_btn.config(state=tk.DISABLED)
        if self.last_articles:
            self.regenerate_btn.config(state=tk.NORMAL)
        self.update_status("处理完成")

        # 只有生成成功时才询问是否打开输出文件夹
        if self.txt_generate_success:
            if messagebox.askyesno("完成", "TXT生成完成！是否打开TXT保存文件夹？"):
                txt_path = self.txt_output_path.get().replace('/', '\\')
                if os.path.exists(txt_path):
                    os.startfile(txt_path)
        else:
            messagebox.showwarning("提示", "生成失败，请查看日志了解详情")

    def stop_generate(self):
        """停止生成"""
        self.is_running = False
        self.log("用户停止了生成任务")
        self.update_status("已停止")
        self.start_btn.config(state=tk.NORMAL)
        self.stop_btn.config(state=tk.DISABLED)

    def regenerate(self):
        """重新生成 - 使用上次的参考文案和当前设置"""
        if not self.last_articles:
            messagebox.showwarning("提示", "没有可重新生成的文案，请先执行一次生成")
            return

        output_path = self.output_path.get()
        if not output_path:
            messagebox.showerror("错误", "请选择输出保存路径")
            return

        os.makedirs(output_path, exist_ok=True)

        self.start_btn.config(state=tk.DISABLED)
        self.stop_btn.config(state=tk.NORMAL)
        self.regenerate_btn.config(state=tk.DISABLED)
        self.is_running = True

        self.log_text.delete(1.0, tk.END)
        self.log("开始重新生成...")

        # 获取用户建议
        suggestion = self.suggestion_text.get("1.0", tk.END).strip()
        # 过滤掉默认提示文字
        if suggestion.startswith("例如："):
            suggestion = ""

        thread = threading.Thread(target=self.regenerate_task, args=(suggestion,))
        thread.daemon = True
        thread.start()

    def regenerate_task(self, user_suggestion):
        """重新生成任务"""
        try:
            output_path = self.output_path.get()

            if user_suggestion:
                self.log(f"用户建议：{user_suggestion}")

            total = len(self.last_articles)
            for art_idx, article in enumerate(self.last_articles):
                if not self.is_running:
                    break

                self.log(f"\n--- 重新生成第 {art_idx+1} 篇 ---")
                self.update_status(f"正在重新生成第 {art_idx+1} 篇...")
                self.update_progress((art_idx / total) * 100)

                # 生成仿写文案（带用户建议）
                result = self.generate_document(
                    article, self.last_flow_type, self.last_yinliu_content,
                    self.last_product_name, self.last_product_material,
                    user_suggestion
                )

                if result:
                    self.save_document(result, output_path, art_idx + 1)
                else:
                    self.log(f"第 {art_idx+1} 篇重新生成失败")

            self.update_progress(100)
            self.log("\n" + "="*50)
            self.log("重新生成完成！")
            self.finish_task()

        except Exception as e:
            self.log(f"重新生成出错: {str(e)}")
            import traceback
            self.log(traceback.format_exc())
            self.finish_task()

    def generate_task(self):
        """生成任务主函数"""
        try:
            output_path = self.output_path.get()
            flow_type = self.flow_type.get()

            # 读取参考文案
            self.log("正在读取参考文案...")
            input_mode = self.input_mode.get()

            if input_mode == "paste":
                # 粘贴文本模式
                paste_content = self.paste_text.get("1.0", tk.END).strip()
                if not paste_content:
                    self.log("错误：请粘贴参考文案内容")
                    self.finish_task()
                    return
                files_content = [("粘贴文案", paste_content)]
            else:
                # 文件模式
                input_path = self.input_path.get()
                if os.path.isfile(input_path):
                    with open(input_path, 'r', encoding='utf-8') as f:
                        content = f.read()
                    files_content = [(os.path.basename(input_path), content)]
                else:
                    files_content = []
                    for fname in os.listdir(input_path):
                        if fname.endswith('.txt'):
                            fpath = os.path.join(input_path, fname)
                            with open(fpath, 'r', encoding='utf-8') as f:
                                files_content.append((fname, f.read()))

            if not files_content:
                self.log("错误：没有找到任何文案文件")
                self.finish_task()
                return

            # 读取引流素材
            yinliu_content = self.yinliu_text.get("1.0", tk.END).strip()

            # 带货信息
            product_name = ""
            product_material = ""
            if flow_type == "带货引流":
                product_name = self.product_name.get().strip()
                product_material = self.product_material.get("1.0", tk.END).strip()

            total_files = len(files_content)
            for idx, (fname, content) in enumerate(files_content):
                if not self.is_running:
                    break

                self.log(f"\n{'='*50}")
                self.log(f"处理文件 [{idx+1}/{total_files}]: {fname}")
                self.update_progress((idx / total_files) * 100)

                # 分割参考文案
                articles = self.parse_articles(content)
                self.log(f"识别到 {len(articles)} 篇参考文案")

                # 把参考文案追加到素材库
                for article in articles:
                    self.append_reference_to_library(article, flow_type)

                # 保存最后一次生成的信息（用于重新生成）
                self.last_articles = articles
                self.last_flow_type = flow_type
                self.last_yinliu_content = yinliu_content
                self.last_product_name = product_name
                self.last_product_material = product_material

                for art_idx, article in enumerate(articles):
                    if not self.is_running:
                        break

                    self.log(f"\n--- 处理第 {art_idx+1} 篇参考文案 ---")
                    self.update_status(f"正在生成第 {art_idx+1} 篇...")

                    # 生成仿写文案
                    result = self.generate_document(
                        article, flow_type, yinliu_content,
                        product_name, product_material
                    )

                    if result:
                        # 保存文档
                        self.save_document(result, output_path, art_idx + 1)
                    else:
                        self.log(f"第 {art_idx+1} 篇生成失败")

            self.update_progress(100)
            self.log("\n" + "="*50)
            self.log("全部处理完成！")
            self.finish_task()

        except Exception as e:
            self.log(f"任务出错: {str(e)}")
            import traceback
            self.log(traceback.format_exc())
            self.finish_task()

    def parse_articles(self, content):
        """解析参考文案，按'参考文案：'分隔"""
        articles = []
        parts = re.split(r'参考文案[：:]', content)
        for part in parts:
            part = part.strip()
            if part and len(part) > 50:
                articles.append(part)
        if not articles and content.strip():
            articles = [content.strip()]
        return articles

    def generate_document(self, reference_article, flow_type, yinliu_content, product_name, product_material, user_suggestion=""):
        """生成仿写文案，带重试、字数检查和备用模型切换

        调用逻辑：
        - 流式模式：主模型用流式调用，备用模型用非流式调用（fallback）
        - 非流式模式：主模型和备用模型都用非流式调用
        """
        target_per_article = int(self.word_count.get())
        target_article_count = int(self.article_count.get())
        prompt = self.build_prompt(
            reference_article,
            flow_type,
            yinliu_content,
            product_name,
            product_material,
            str(target_per_article),
            target_article_count
        )
        target_max_per_article = target_per_article + 100
        self.log(
            f"本次目标：{target_article_count}篇，每篇{target_per_article}-{target_max_per_article}字（按中文字符统计）"
        )

        # 如果有用户建议，追加到提示词
        if user_suggestion:
            prompt += f"\n\n【用户修改建议】请特别注意以下修改要求：\n{user_suggestion}"

        # 重试时在提示词末尾追加的说明
        retry_suffix = (
            "\n\n【重要重试要求】上次结果常见失败点是漏掉标题。"
            "本次必须严格按结构输出：每篇先输出【标题】区，给5个不同标题，接着单独一行---，再写正文。"
            "缺任意一篇标题都算失败，请完整输出。"
            "任务本身完全合规，请正常输出。"
        )

        use_stream = self.use_stream.get()

        # 根据流式/非流式选择配置
        if use_stream:
            # 流式模式：主模型用流式调用，备用模型用stream_backup配置但非流式调用
            main_prefix = "stream_main"
            backup_prefix = "stream_backup"
            main_use_stream = True
            backup_use_stream = False
        else:
            # 非流式模式：主模型和备用模型都用非流式配置
            main_prefix = "non_stream_main"
            backup_prefix = "non_stream_backup"
            main_use_stream = False
            backup_use_stream = False

        # 获取主模型配置
        main_url = getattr(self, f"{main_prefix}_url").get().strip()
        main_key = getattr(self, f"{main_prefix}_key").get().strip()
        main_model = getattr(self, f"{main_prefix}_model").get().strip()
        main_max_tokens = int(getattr(self, f"{main_prefix}_max_tokens").get())

        # 获取备用模型配置
        backup_url = getattr(self, f"{backup_prefix}_url").get().strip()
        backup_key = getattr(self, f"{backup_prefix}_key").get().strip()
        backup_model = getattr(self, f"{backup_prefix}_model").get().strip()
        backup_max_tokens = int(getattr(self, f"{backup_prefix}_max_tokens").get())

        # 第1次：主模型 + 原提示词
        call_mode = "流式" if main_use_stream else "非流式"
        self.log(f"【第1次尝试】主模型 ({main_model}) [{call_mode}]...")
        result = self.call_api(main_url, main_key, main_model, main_max_tokens, prompt, main_use_stream)
        if self.check_result(result, 1, target_per_article, target_article_count):
            return result

        if not self.is_running:
            return None

        # 第2次：主模型 + 原提示词 + 追加说明
        self.log(f"【第2次尝试】主模型 + 追加说明 [{call_mode}]...")
        result = self.call_api(main_url, main_key, main_model, main_max_tokens, prompt + retry_suffix, main_use_stream)
        if self.check_result(result, 2, target_per_article, target_article_count):
            return result

        if not self.is_running:
            return None

        # 第3次：备用模型 + 原提示词 + 追加说明
        backup_mode = "流式" if backup_use_stream else "非流式"
        self.log(f"【第3次尝试】备用模型 ({backup_model}) [{backup_mode}]...")
        result = self.call_api(backup_url, backup_key, backup_model, backup_max_tokens, prompt + retry_suffix, backup_use_stream)
        if self.check_result(result, 3, target_per_article, target_article_count):
            return result

        if not self.is_running:
            return None

        # 第4次：备用模型 + 原提示词 + 追加说明
        self.log(f"【第4次尝试】备用模型（最后一次）[{backup_mode}]...")
        result = self.call_api(backup_url, backup_key, backup_model, backup_max_tokens, prompt + retry_suffix, backup_use_stream)
        if self.check_result(result, 4, target_per_article, target_article_count):
            return result

        self.log("4次尝试均失败")
        return None

    def check_result(self, result, attempt_num, per_article, article_count):
        """检查结果：有效性 + 字数（根据页面设置的篇数和字数要求）"""
        if not result or not self.is_valid_result(result):
            self.log(f"第{attempt_num}次：API返回无效或被拒绝")
            return False

        char_count = self.count_chinese_chars(result)
        self.log(f"第{attempt_num}次：生成总字数 {char_count}")

        max_per_article = per_article + 100

        # 先做总字数下限校验（超出上限不拦截）
        min_total = per_article * article_count
        max_total = max_per_article * article_count
        if char_count < min_total:
            self.log(f"总字数不足{min_total}（{article_count}篇×{per_article}字），需要重试...")
            return False
        if char_count > max_total:
            self.log(f"总字数超出建议区间上限{max_total}，但不拦截")

        # 开头反故事化校验：命中故事触发词则重试
        if self.has_story_opening(result, article_count):
            self.log("检测到故事化开头（如“那天/有一次/有人问我”等），仅提示不拦截")

        # 不因标题缺失重试：标题通过提示词强约束，但校验阶段不拦截
        split_articles = self.split_articles_from_result(result)
        if len(split_articles) < article_count:
            self.log(f"篇数识别不足（需要{article_count}篇，识别到{len(split_articles)}篇），继续按结果保存")

        # 仅按总字数下限判断；不校验单篇范围
        avg_per_article = char_count // article_count
        similar, matched_file, sim_score = self.is_too_similar_to_recent(result)
        if similar:
            self.log(f"与近期文案相似度过高（{sim_score:.2f}，参考：{matched_file}），需要重试...")
            return False
        self.log(f"字数合格（仅校验总字数下限{min_total}；实际{char_count}，平均每篇{avg_per_article}字），生成成功！")
        return True

    def normalize_for_similarity(self, text):
        """清洗文本用于相似度比较"""
        if not text:
            return ""
        return "".join(re.findall(r"[\u4e00-\u9fffA-Za-z0-9]", text))

    def has_story_opening(self, result, article_count):
        """检测是否出现故事叙述式开场（仅检查每篇开头）"""
        story_markers = [
            "那天", "有一次", "后来我", "小时候", "当年",
            "我有个", "我认识", "我朋友", "有人问我", "有人跟我说"
        ]

        segments = self.split_articles_from_result(result)
        if segments and len(segments) >= 1:
            targets = [content[:180] for content, _ in segments[:article_count]]
        else:
            # 无法拆分时，检查整体前部
            targets = [result[:320]]

        for start_text in targets:
            cleaned = re.sub(r"\s+", "", start_text)
            for marker in story_markers:
                if marker in cleaned:
                    return True
        return False

    def get_generated_title_set(self):
        """读取已入库的生成标题，用于跨批次去重"""
        try:
            import openpyxl
            if not os.path.exists(GENERATED_LIBRARY_FILE):
                return set()
            wb = openpyxl.load_workbook(GENERATED_LIBRARY_FILE, read_only=True, data_only=True)
            ws = wb[GENERATED_LIBRARY_SHEET] if GENERATED_LIBRARY_SHEET in wb.sheetnames else wb.active
            titles = set()
            for row in ws.iter_rows(min_row=2, values_only=True):
                # 列结构：日期, 引流类型, 标题, 正文, 正文字数
                title = (row[2] or "") if len(row) >= 3 else ""
                title = str(title).strip()
                if title:
                    titles.add(title)
            wb.close()
            return titles
        except Exception:
            return set()

    def append_generated_to_library(self, flow_type, title, article_content):
        """将本次生成文案（标题+正文）入库Excel"""
        try:
            import openpyxl
            from openpyxl import Workbook

            if not os.path.exists(MATERIAL_LIBRARY_DIR):
                os.makedirs(MATERIAL_LIBRARY_DIR)

            if os.path.exists(GENERATED_LIBRARY_FILE):
                wb = openpyxl.load_workbook(GENERATED_LIBRARY_FILE)
                ws = wb[GENERATED_LIBRARY_SHEET] if GENERATED_LIBRARY_SHEET in wb.sheetnames else wb.active
            else:
                wb = Workbook()
                ws = wb.active
                ws.title = GENERATED_LIBRARY_SHEET
                ws.append(["日期", "引流类型", "标题", "正文", "正文字数"])

            today = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            char_count = self.count_chinese_chars(article_content)
            ws.append([today, flow_type, title, article_content, char_count])
            wb.save(GENERATED_LIBRARY_FILE)
        except Exception as e:
            self.log(f"生成文案入库失败: {str(e)}")

    def get_generated_library_corpus(self, max_rows=120):
        """从生成文案Excel读取近期语料，避免txt搬走后失去去重能力"""
        try:
            import openpyxl
            if not os.path.exists(GENERATED_LIBRARY_FILE):
                return []
            wb = openpyxl.load_workbook(GENERATED_LIBRARY_FILE, read_only=True, data_only=True)
            ws = wb[GENERATED_LIBRARY_SHEET] if GENERATED_LIBRARY_SHEET in wb.sheetnames else wb.active
            rows = list(ws.iter_rows(min_row=2, values_only=True))
            wb.close()
            corpus = []
            for row in rows[-max_rows:]:
                if not row:
                    continue
                title = str(row[2] or "").strip() if len(row) >= 3 else ""
                body = str(row[3] or "").strip() if len(row) >= 4 else ""
                if body:
                    label = f"EXCEL:{title}" if title else "EXCEL:未命名标题"
                    corpus.append((label, body))
            return corpus
        except Exception:
            return []

    def get_recent_corpus_texts(self, max_files=60):
        """读取近期输出文案 + 生成文案Excel，作为去模板化比对语料"""
        dirs = []
        for attr in ("txt_output_path", "output_path"):
            if hasattr(self, attr):
                try:
                    path = getattr(self, attr).get().strip()
                    if path and os.path.exists(path):
                        dirs.append(path)
                except Exception:
                    continue

        all_files = []
        seen = set()
        for folder in dirs:
            try:
                for name in os.listdir(folder):
                    if not name.lower().endswith(".txt"):
                        continue
                    full_path = os.path.join(folder, name)
                    if full_path in seen or not os.path.isfile(full_path):
                        continue
                    seen.add(full_path)
                    all_files.append((os.path.getmtime(full_path), full_path, name))
            except Exception:
                continue

        all_files.sort(key=lambda x: x[0], reverse=True)
        results = []
        for _, full_path, name in all_files[:max_files]:
            try:
                with open(full_path, "r", encoding="utf-8", errors="ignore") as f:
                    text = f.read().strip()
                if text:
                    results.append((name, text))
            except Exception:
                continue

        # 追加Excel语料作为长期记忆，避免用户移动txt后失效
        excel_corpus = self.get_generated_library_corpus(max_rows=max_files * 2)
        results.extend(excel_corpus)
        return results

    def text_similarity_score(self, text_a, text_b):
        """混合相似度：编辑序列相似度 + 6字片段重合度"""
        a = self.normalize_for_similarity(text_a)
        b = self.normalize_for_similarity(text_b)
        if not a or not b:
            return 0.0

        seq_ratio = SequenceMatcher(None, a, b).ratio()

        n = 6
        if len(a) < n or len(b) < n:
            return seq_ratio
        set_a = {a[i:i + n] for i in range(len(a) - n + 1)}
        set_b = {b[i:i + n] for i in range(len(b) - n + 1)}
        if not set_a or not set_b:
            return seq_ratio
        overlap = len(set_a & set_b) / max(1, min(len(set_a), len(set_b)))
        return max(seq_ratio, overlap)

    def is_too_similar_to_recent(self, text):
        """检查与近期文案是否过于相似"""
        corpus = self.get_recent_corpus_texts(max_files=60)
        if not corpus:
            return False, "", 0.0

        best_score = 0.0
        best_name = ""
        for name, old_text in corpus:
            score = self.text_similarity_score(text, old_text)
            if score > best_score:
                best_score = score
                best_name = name
        return best_score >= self.similarity_threshold, best_name, best_score

    def collect_overused_phrases(self, min_files=4, max_phrases=12):
        """提取近期文案里高频复用短语，加入黑名单避免模板化"""
        corpus = self.get_recent_corpus_texts(max_files=80)
        if not corpus:
            return []

        counter = Counter()
        n = 6
        for _, text in corpus:
            cleaned = "".join(re.findall(r"[\u4e00-\u9fff]", text))
            if len(cleaned) < n:
                continue
            seen = set()
            for i in range(len(cleaned) - n + 1):
                gram = cleaned[i:i + n]
                if gram in seen:
                    continue
                seen.add(gram)
                counter[gram] += 1

        phrases = []
        for gram, cnt in counter.most_common(80):
            if cnt < min_files:
                break
            if gram.startswith("第") and "篇" in gram:
                continue
            if gram.startswith("标题") or gram.startswith("正文"):
                continue
            phrases.append(gram)
            if len(phrases) >= max_phrases:
                break
        return phrases

    def build_dynamic_strategy_instruction(self, flow_type, article_count):
        """构建动态改写策略，降低固定模板感"""
        role_pool = [
            "严师点醒型：语气克制有力，先点破误区，再给出方向",
            "长者托底型：语气沉稳温厚，强调你并不差、只是太久被忽视",
            "朋友交心型：像深夜长谈，情绪递进明显但不煽情过头",
            "高位洞察型：站在更高认知视角，拆穿现象背后的规律",
            "同行共情型：先精准描述你的不容易，再给可执行的转变路径",
            "导师鼓劲型：先认可价值，再给行动指令，带出冲劲"
        ]
        title_intents = [
            "被看见", "身份抬升", "反常识点醒", "损失规避", "结果预告", "尊重与认可"
        ]
        emotion_arcs = [
            "压抑现实→被理解→被抬高→行动冲动",
            "先刺痛→再托底→再升维→再引导",
            "先共情→再鼓劲→再点醒→再召唤",
            "先肯定→再反转→再加压→再给出口"
        ]
        citation_buckets = [
            "古诗词", "史书典故", "格言俗语", "现代名句", "生活观察"
        ]
        cta_pool = {
            "置顶引流": ["答案预告型", "方法揭晓型", "错过损失型", "专属提醒型"],
            "橱窗引流": ["场景匹配型", "价值放大型", "犒赏自己型", "立即行动型"],
            "带货引流": ["痛点解决型", "结果对比型", "稀缺时效型", "信任托付型"],
            "纯夸赞不引流": ["鼓励收束型", "祝福落点型", "方向建议型", "认同陪伴型"]
        }
        flow_cta = cta_pool.get(flow_type, ["行动引导型", "价值提示型", "陪伴提醒型"])

        roles = random.sample(role_pool, min(article_count, len(role_pool)))
        if len(roles) < article_count:
            roles.extend(random.choices(role_pool, k=article_count - len(roles)))
        intents = random.sample(title_intents, min(article_count, len(title_intents)))
        if len(intents) < article_count:
            intents.extend(random.choices(title_intents, k=article_count - len(intents)))
        arcs = random.sample(emotion_arcs, min(article_count, len(emotion_arcs)))
        if len(arcs) < article_count:
            arcs.extend(random.choices(emotion_arcs, k=article_count - len(arcs)))
        citations = random.sample(citation_buckets, min(article_count, len(citation_buckets)))
        if len(citations) < article_count:
            citations.extend(random.choices(citation_buckets, k=article_count - len(citations)))
        ctas = random.sample(flow_cta, min(article_count, len(flow_cta)))
        if len(ctas) < article_count:
            ctas.extend(random.choices(flow_cta, k=article_count - len(ctas)))

        article_plans = []
        for idx in range(article_count):
            article_plans.append(
                f"- 第{idx + 1}篇：角色={roles[idx]}；标题意图={intents[idx]}；情绪曲线={arcs[idx]}；"
                f"引经类别={citations[idx]}；收尾钩子={ctas[idx]}"
            )

        forbidden_phrases = self.collect_overused_phrases(min_files=4, max_phrases=10)
        forbidden_text = "、".join(forbidden_phrases) if forbidden_phrases else "无（仍需避免套话）"

        return f"""## 动态改写策略（必须执行，禁止模板化）
### 先内部写蓝图（不输出蓝图）
- 在正式写作前，先在内部完成每篇蓝图：受众处境、痛点、情绪转折、身份抬升句、行动引导句。
- 蓝图只用于思考，不要在最终答案里输出“蓝图/分析/说明”等字样。

### 本次分篇策略
{chr(10).join(article_plans)}

### 去模板硬约束
- 每篇开头句式必须不同，禁止复用“你有没有发现/我想告诉你/这世上”等高频模板句。
- 每篇中段的论证方式必须不同：至少一篇“拆解误区”，至少一篇“价值抬升”，至少一篇“行动催化”。
- 引经据典必须按上面的“引经类别”分配，不允许多篇重复同一类。
- 标题必须与正文强绑定，不能只换同义词。
- 下面这些是近期高频短语，禁止原样复用：{forbidden_text}
"""

    def count_chinese_chars(self, text):
        """统计中文字符数量"""
        if not text:
            return 0
        chinese_chars = re.findall(r'[\u4e00-\u9fff]', text)
        return len(chinese_chars)

    def is_valid_result(self, result):
        """检查结果是否有效"""
        if not result:
            return False
        # 检查是否是拒绝回答
        reject_keywords = ["抱歉", "无法", "不能", "拒绝", "违反", "政策", "sorry", "cannot", "can't"]
        result_lower = result.lower()
        for kw in reject_keywords:
            if kw in result_lower and len(result) < 500:
                return False
        return len(result) > 300

    def call_api(self, base_url, api_key, model, max_tokens, prompt, use_stream):
        """调用API（流式或非流式）"""
        if use_stream:
            return self.call_llm_stream(base_url, api_key, model, max_tokens, prompt)
        else:
            return self.call_llm_non_stream(base_url, api_key, model, max_tokens, prompt)

    def call_llm_stream(self, base_url, api_key, model, max_tokens, prompt):
        """流式API调用"""
        url = f"{base_url.rstrip('/')}/chat/completions"
        headers = {
            "Content-Type": "application/json",
            "Authorization": f"Bearer {api_key}"
        }
        data = {
            "model": model,
            "messages": [
                {"role": "system", "content": "你是一个专业的文案写作专家，擅长仿写百家号引流文案。"},
                {"role": "user", "content": prompt}
            ],
            "max_tokens": max_tokens,
            "stream": True
        }
        if "thinking" not in model.lower():
            data["temperature"] = 0.7

        self.log(f"[流式API] 请求: {url}, 模型: {model}")

        try:
            response = requests.post(url, headers=headers, json=data, timeout=300, stream=True)

            if response.status_code != 200:
                error_text = response.text[:200] if response.text else "空响应"
                self.log(f"[流式API] 失败: HTTP {response.status_code}: {error_text}")
                return None

            full_content = ""
            for line in response.iter_lines():
                if not self.is_running:
                    return None
                if line:
                    line_text = line.decode('utf-8')
                    if line_text.startswith('data: '):
                        json_str = line_text[6:]
                        if json_str.strip() == '[DONE]':
                            break
                        try:
                            chunk = json.loads(json_str)
                            if 'choices' in chunk and len(chunk['choices']) > 0:
                                delta = chunk['choices'][0].get('delta', )
                                content = delta.get('content', '')
                                if content:
                                    full_content += content
                        except json.JSONDecodeError:
                            continue

            if full_content:
                self.log(f"[流式API] 成功获取 {len(full_content)} 字符")
                return full_content
            else:
                self.log("[流式API] 返回空内容")
                return None

        except Exception as e:
            self.log(f"[流式API] 异常: {type(e).__name__}: {e}")
            return None

    def call_llm_non_stream(self, base_url, api_key, model, max_tokens, prompt):
        """非流式API调用"""
        url = f"{base_url.rstrip('/')}/chat/completions"
        headers = {
            "Content-Type": "application/json",
            "Authorization": f"Bearer {api_key}"
        }
        data = {
            "model": model,
            "messages": [
                {"role": "system", "content": "你是一个专业的文案写作专家，擅长仿写百家号引流文案。"},
                {"role": "user", "content": prompt}
            ],
            "max_tokens": max_tokens
        }
        if "thinking" not in model.lower():
            data["temperature"] = 0.7

        self.log(f"[非流式API] 请求: {url}, 模型: {model}")

        try:
            response = requests.post(url, headers=headers, json=data, timeout=300)

            if response.status_code != 200:
                error_text = response.text[:200] if response.text else "空响应"
                self.log(f"[非流式API] 失败: HTTP {response.status_code}: {error_text}")
                return None

            result = response.json()
            content = result["choices"][0]["message"]["content"]
            self.log(f"[非流式API] 成功获取 {len(content)} 字符")
            return content

        except Exception as e:
            self.log(f"[非流式API] 异常: {type(e).__name__}: {e}")
            return None

    def build_prompt(self, reference_article, flow_type, yinliu_content, product_name, product_material, word_count, article_count=3):
        """构建生成提示词 - 完整详细版（整合自Skill）"""
        flow_instruction = self.get_flow_instruction(flow_type, yinliu_content, product_name, product_material, article_count)
        dynamic_instruction = self.build_dynamic_strategy_instruction(flow_type, article_count)
        strict_output_protocol = f"""## 硬性输出协议（必须严格遵守）
- 先在内部完成每篇正文，再基于该篇正文提炼5个标题（内部步骤，不要输出“草稿”）。
- 必须输出{article_count}篇；每篇先输出【标题】区，再输出正文。
- 每篇【标题】区必须有且仅有5行标题，每行1个，不要写“标题1：”前缀编号。
- 标题区后必须输出单独一行 `---`，再开始正文。
- 任意一篇缺少标题区/标题不足5个/没有`---`分隔，均视为任务失败。
"""
        title_positive_protocol = """## 标题风格硬规则（最高优先级）
- 每篇的5个标题都必须是积极、正向、夸赞、认可、鼓励的语气。
- 标题必须让读者感到“被理解、被尊重、被肯定、被看见”。
- 标题可以有“点醒感”和“压迫感”，但压迫感只能针对困局，不得打压读者本人。
- 禁止负面贬低、羞辱、恐吓、诅咒、讽刺读者的表达。
- 标题必须和该篇正文内容强绑定，不能空泛套话。
- 同一篇的5个标题必须明显不同，不得同义改写。
"""

        # 根据篇数随机选择开头
        random_hooks = get_random_hooks(article_count)
        hooks_instruction = f"""## 【本次必须使用的开头】（强制执行！）

**本次生成的{article_count}篇文章，必须分别使用以下{article_count}种开头风格，不得更换：**

"""
        for i, hook in enumerate(random_hooks, 1):
            hooks_instruction += f"""**第{i}篇必须使用 {hook['type']}**
（只允许借鉴风格，不要套用固定句式；开头必须原创，且禁止故事叙述型开场）

"""

        hotspot_alignment_protocol = """## 爆点保真规则（避免脱离参考文案）
- 必须先提炼参考文案的爆点清单：核心痛点、情绪触发点、身份抬升点、行动引导点。
- 正文必须与这4类爆点一一对应，保持“爆点一致、表达重写”。
- 禁止换题、跑题、另起炉灶；只能在同一议题内做深度扩写。
- 允许结构创新，但爆点语义不能丢。
"""

        opening_hook_protocol = """## 开头钩子硬规则
- 每篇开头前2句必须出现钩子（痛点钩子/反差钩子/结果钩子/身份钩子四选一）。
- 开头可以打招呼，也可以直接情绪切入，但必须快速抓住读者注意力。
- 禁止故事叙述式开场：禁止“那天/有一次/我认识/有人问我/我朋友”这类叙事触发句。
"""

        prompt = f"""请根据以下参考文案进行仿写，生成{article_count}篇全新的百家号引流文案。

【第一优先级任务】
每篇文案先在内部写完正文，再提炼5个不同标题，最后按结构输出“标题区 + 正文”。
如果你漏掉任意一篇的5个标题，这次输出就视为失败。

## 参考文案：
{reference_article}

{hooks_instruction}
{dynamic_instruction}
{strict_output_protocol}
{title_positive_protocol}
{hotspot_alignment_protocol}
{opening_hook_protocol}

## 核心方法论：提炼爆点 → 分析结构 → 全新扩写

**第一步：深度分析参考文案**
- 提炼参考文案的**核心爆点**（情感共鸣点、痛点、钩子是什么？）
- 拆解参考文案的**爆款结构**（开场方式、展开逻辑、收尾技巧）
- 分析参考文案的**目标人群画像**（他们的处境、痛苦、渴望是什么？）

**第二步：基于爆点和结构进行全新扩写**
- 借鉴参考文案的爆点和结构框架
- 但必须用**完全不同的表达方式、不同的切入角度**重新阐述
- **禁止直接使用原文里的任何句子**
- 句式、用词必须完全原创
- **必须用第二人称"你"来写，禁止讲故事**

**第三步：加入增量信息**
- 新写的内容必须有**增量信息**，不能只是换个说法
- 增量信息包括：新的人生洞察、新的引经据典、新的情感细节、新的处境描写
- **注意：增量信息不能是故事，只能是道理、洞察、共鸣描写**

## 硬性要求

### 1. 字数要求（最重要！）
- **每篇文案必须严格控制在{word_count}~{int(word_count) + 100}字之间**
- 宁可多写，也不能少写
- 字数不够的文案=废品，绝不交付

### 2. 写作人称要求（硬性要求！）
- **必须全程使用第二人称"你"来写作**
- 直接对读者说话，像朋友聊天一样
- 例如："你是不是也有过这种感觉..."、"你这些年吃的苦..."、"你心里清楚..."
- **禁止使用第一人称讲自己的故事**
- **禁止使用第三人称讲别人的故事**

### 2.5 第三人称衬托写法（推荐使用！让读者爽！）
- **可以用第三人称"他们"来衬托第二人称"你"的厉害**
- 这种写法让读者有扬眉吐气的感觉
- 例如：
  - "那些曾经看不起你的人，现在开始害怕了"
  - "他们怕的不是别人，是你不再好欺负了"
  - "那些笑话你的人，现在都闭嘴了"
  - "曾经否定你的人，开始后悔了"
  - "他们以为你会认输，没想到你越来越强"
- **"他们"指那些曾经否定/看不起/欺负读者的人**
- **这种写法的效果：让读者觉得自己很厉害，那些人都怕了/后悔了**
- 每篇文案中可以适当穿插1-2处这种写法，增加爽感

### 2.6 多种内容切入角度（{article_count}篇文案必须使用不同的切入角度！）

**【A类：缘分/命定型】（很受欢迎！）**
- 从宏观视角开头，让读者觉得自己是"被选中的"
- 例如：
  - "大千世界，芸芸众生，能看到这段话的人，都不简单"
  - "茫茫人海，你能刷到这条，就说明你跟别人不一样"
  - "万千人中，你停下来看这段文字，这就是缘分"
  - "这世上那么多人，偏偏是你看到了，说明这段话就是说给你听的"
- 效果：让读者觉得自己特别、被选中、命中注定

**【B类：好消息/喜讯型】**
- 用好消息开头，让读者心情愉悦
- 例如："告诉你一个好消息..."、"有件喜事想跟你说..."、"你等的好事要来了..."
- 效果：让读者期待、开心

**【C类：送礼/赠予型】**
- 像送礼物一样送给读者一句话、一个道理
- 例如："今天送你一句话..."、"把这份祝福送给你..."、"这段话送给正在看的你..."
- 效果：让读者感到被重视、被关爱

**【D类：引经据典型】**
- 用名言名句、古语经典自然引入
- **注意：不要指定具体哪本书，让内容自然流畅，不要生搬硬套**
- 例如："古人说得好，..."、"有句老话讲得透，..."、"老祖宗留下一句话，..."
- 效果：增加文案的厚重感和说服力
- **禁止生硬引用，必须与内容自然融合**

**【E类：夸孩子型】**
- 夸读者的孩子，间接夸读者教育得好
- 例如："你的孩子有你这样的父母，是他的福气"、"能把孩子教得这么懂事，你真的很了不起"
- 效果：让读者感到骄傲、被认可

**【F类：夸家庭型】**
- 夸读者对家庭的付出和贡献
- 例如："有你这样的人撑着，这个家散不了"、"这个家能有今天，全靠你在扛"
- 效果：让读者感到自己的价值被看见

**【G类：夸品质型】**
- 直接夸读者的某种品质
- 例如："你这种人，心太善了"、"像你这么实在的人，现在真不多了"
- 效果：让读者感到被欣赏

**【H类：懂你型】**
- 表达对读者的理解和共情
- 例如："我知道你这些年有多不容易"、"你心里的苦，不说我也懂"
- 效果：让读者感到被理解、被看见

**【I类：预言/好结果型】**
- 预言读者会有好结果
- 例如："像你这样的人，以后一定会越来越好"、"你的好日子在后头呢"
- 效果：给读者希望和信心

**切入角度使用规则：**
- {article_count}篇文案必须使用{article_count}种不同的切入角度
- 每篇文案可以混合使用多种角度，但要有一个主要角度
- 切入角度要与文案内容自然融合，不能生硬
- **缘分/命定型（A类）很受欢迎，建议至少1篇使用**
- **禁止直接照搬上面的例子！例子只是说明方向，必须根据文案内容原创全新的表达**
- **每篇文案的切入方式必须完全不同，禁止雷同**

### 3. 禁止讲故事（硬性要求！）
- **绝对禁止讲任何具体故事**
- 不能讲"我有个朋友..."、"我认识一个人..."、"有一次我..."
- 不能讲"张三怎么样..."、"李四怎么样..."等第三人称故事
- 不能讲任何有具体人物、具体情节的故事
- **只能用第二人称直接描述读者的处境、感受、经历**
- 用"你是不是..."、"你有没有..."、"你心里..."来引发共鸣
- 用概括性的描述代替具体故事

### 4. 相似度控制（硬性要求！）
- **与参考文案的相似度必须低于10%**
- 不能直接改写原文句子
- 不能只是替换同义词
- 必须用全新的表达方式
- 结构可以借鉴，但内容必须完全原创

## 开头要求（{article_count}篇必须完全不同类型！像开盲盒一样有惊喜！）

**【最重要】强制多样化机制：**
- **每次生成必须使用全新的开头，绝对禁止重复**
- **{article_count}篇文章的开头必须来自{article_count}个不同的大类**
- **同一个开头句式只能用一次**
- **必须包含至少1篇轻松/温暖/有趣风格的开头**
- **禁止全部都是沉重压抑的风格**

**永久禁止的开头（已经用烂了！）：**
- ❌ "我劝你别太善良"
- ❌ "有一种人，你越对他好，他越瞧不起你"
- ❌ "凭什么受伤的总是老实人？"
- ❌ "四十岁之后，我才明白一个道理"
- ❌ "那天饭局上，有人说了一句话"
- ❌ "你累了"、"深夜睡不着"、"夜深人静"
- ❌ "我知道你是什么样的人"
- ❌ "我问你一个问题"
- ❌ "你有没有发现"、"你有没有想过"
- ❌ "人这辈子"、"人啊"开头的句式
- ❌ 任何以"有一种..."开头的句式

**开头钩子库（150+种，{article_count}篇必须从不同类型中选择，必须有1篇轻松风格）：**

**A类-轻松幽默型（必选！{article_count}篇中至少用1个）：**
- "我发现一个规律：越是老实人，越容易被安排加班。你说气不气？"
- "昨天算了一笔账，这些年帮别人花的时间，够我学会三门外语了。"
- "我妈说我最大的优点是善良，我爸说这也是我最大的缺点。好家伙，亲爹。"
- "朋友说我是'便利店型人格'——24小时营业，随叫随到，还不涨价。"
- "有人说我脾气好，我笑了笑没说话。其实不是脾气好，是懒得计较。"
- "我终于明白，为什么'好人卡'发得最多——因为好人最好打发。"
- "同事问我：'你怎么从来不生气？'我说：'生气要花力气，我选择省着点用。'"
- "我的人生信条曾经是'吃亏是福'，直到我发现福没来，亏倒是吃了不少。"
- "别人都在研究怎么赚钱，我在研究怎么不被人当免费劳动力。"
- "我这人有个毛病，别人一说'就你能帮我'，我就跟中了蛊似的。"

**B类-温暖治愈型（推荐！让读者感到被理解）：**
- "嘿，今天想跟你聊点轻松的，关于那些默默付出却不求回报的人。"
- "你知道吗，这世上有一种人，他们的好，是藏在细节里的。"
- "我一直觉得，善良的人身上有光，只是有时候这光被辜负了。"
- "如果你正在看这段话，我想告诉你：你的好，有人看得见。"
- "今天不讲大道理，就想跟你说说心里话。"
- "有些人，值得被这个世界温柔以待，比如正在看这段话的你。"
- "我见过很多人，但像你这样的，真的不多。"
- "你有没有被人夸过'你人真好'？今天我想认真聊聊这件事。"
- "这段话，送给每一个在生活里默默扛着的人。"
- "我相信，看到这里的你，一定是个心里有温度的人。"

**C类-反转惊喜型（有趣！先抑后扬）：**
- "我曾经以为自己是个'老好人'，后来发现，我是个'聪明的好人'。"
- "都说老实人吃亏，但我认识一个老实人，现在过得比谁都好。"
- "你以为善良是软弱？不，善良是一种选择，而且是强者的选择。"
- "有人说心软的人没出息，我偏不信这个邪。"
- "我见过最厉害的人，恰恰是最善良的那个。"
- "别人都说我太老实会吃亏，结果呢？我还真没亏。"
- "都说好人没好报，但我今天要讲一个好人有好报的故事。"
- "我以前觉得'人善被人欺'是真理，直到我遇见了一个人。"
- "谁说善良的人就要受委屈？我第一个不服。"
- "老实人的春天，其实一直都在，只是很多人没发现。"

**D类-对话引入型：**
- "他说完这句话，我愣在原地半天没回过神：'你就是太好说话了。'"
- "有人问我：'你这辈子最后悔的事是什么？'我想了想，说了两个字。"
- "算了，不争了——你是不是也经常这样跟自己说？"
- "我爷爷临终前拉着我的手说：'记住，吃亏的人，老天爷都记着账呢。'"
- "'你怎么这么傻？'这句话，你听过多少次了？"
- "我妈常说一句话：'人善被人欺，马善被人骑。'我以前不信，现在信了。"
- "'别太老实了，会吃亏的。'说这话的人，后来怎么样了？"
- "朋友跟我说：'你知道你最大的问题是什么吗？就是太把别人当回事。'"
- "'谢谢你'——这三个字，你等了多久才听到？"
- "有人当面问我：'你是不是傻？人家都欺负到头上了你还忍？'"

**E类-生活观察型（接地气！）：**
- "菜市场大妈的一句话，让我愣了半天：'姑娘，你这么好说话，不怕被人欺负啊？'"
- "堵在路上，收音机里突然传来一句话，我眼眶一下就红了。"
- "排队的时候，前面两个人的对话让我心里一惊。"
- "超市结账时，前面那个人的一个举动，让我看清了人性。"
- "参加同学聚会，有个人的变化让我震惊了。"
- "小区门口，两个大妈聊天，聊着聊着说出了一个真相。"
- "去参加婚礼，新郎的一句话让全场安静了。"
- "年夜饭桌上，我爸突然放下筷子，说了一句话。"
- "早上买早餐，老板娘的一句话让我想了一整天。"
- "坐出租车，司机师傅跟我聊了一路，最后一句话让我沉默了。"

**F类-数字锚定型：**
- "认识老王二十年了，他教会我一件事：别对谁都掏心掏肺。"
- "被人欺负了三年，我终于想通了一个道理。"
- "用了十年时间，我才学会一个字：不。"
- "三次被人背叛之后，我悟了。"
- "五十岁之后，我才明白什么叫'人走茶凉'。"
- "帮了他七年，他一句谢谢都没说过。"
- "借出去的三万块，要了五年都没要回来。"
- "在这个单位干了八年，我终于明白了一个道理。"

**G类-金句破题型：**
- "最傻的事，就是跟烂人讲道理。"
- "你知道什么人最可怕吗？不是坏人，是那些笑着捅你刀子的人。"
- "千万别做老好人，我吃过这个亏，现在告诉你。"
- "你越忍让，别人越得寸进尺，这是我用十年换来的教训。"
- "老实人不是没脾气，是把脾气都咽进了肚子里。"
- "你的善良，要带点锋芒。"
- "不是所有的忍让都叫大度，有时候那叫窝囊。"
- "这世上最傻的事，就是把真心给了不值得的人。"
- "有些人，你帮他一百次，他记不住；你拒绝他一次，他记你一辈子。"

**H类-悬念钩子型：**
- "有件事我憋了很久，今天必须说出来。"
- "你可能不信，但接下来我说的都是真事。"
- "我要告诉你一个很多人不愿意承认的真相。"
- "接下来这段话，可能会让你不舒服，但我还是要说。"
- "有个规律，我观察了很多年才看透。"
- "今天说的这些话，可能会得罪人，但我不在乎。"
- "有些话，我本来不想说，但看到你，我忍不住了。"
- "接下来的话，你可能不爱听，但句句都是真的。"
- "有件事，我一直没跟任何人说过，今天破例。"

**I类-共鸣代入型：**
- "那种被人当众下面子的感觉，我太懂了。"
- "被最信任的人捅刀子，那种滋味，经历过的人都懂。"
- "明明没做错什么，却总被人针对，这种事我也遇到过。"
- "有些委屈，说出来都没人信。"
- "你是不是也有过这种感觉：付出最多的人，往往最不被珍惜。"
- "那种心寒的感觉，我懂。就像一盆冷水从头浇到脚。"
- "有一种苦，叫做'打碎了牙往肚子里咽'。"
- "那种被人利用完就扔掉的感觉，我经历过。"

**J类-人物故事型：**
- "我有个朋友，前两天跟我说了一件事，我听完沉默了很久。"
- "我爸这辈子只教过我一个道理，我到现在才真正理解。"
- "单位有个人，大家都不待见他，后来我才知道原因。"
- "我们小区有个大爷，天天在楼下坐着，有一天他跟我说了一番话。"
- "我表姐的经历，让我彻底看清了人心。"
- "我有个同学，当年是班里最老实的人，你猜他现在怎么样了？"
- "我舅舅年轻时吃过一个大亏，他把这个教训告诉了我。"
- "我认识一个人，他的经历让我相信：好人终有好报。"

**K类-自我剖析型：**
- "说出来不怕你笑话，我以前也是个傻子。"
- "回头看看这些年，我最后悔的一件事是太心软。"
- "如果能重来，我绝对不会再做老好人。"
- "我吃过的亏，今天全告诉你，希望你别再走我的老路。"
- "我这辈子最大的毛病，就是太把别人当回事。"
- "我曾经也是个'老好人'，后来我学聪明了。"
- "我年轻时犯过一个错，现在想起来还后悔。"
- "我以前总觉得吃亏是福，现在不这么想了。"

**L类-转折反差型：**
- "以前我不信这个道理，直到自己栽了跟头。"
- "年轻的时候觉得这话是废话，现在才知道是真理。"
- "曾经有人跟我说过一句话，我没当回事，后来我后悔了。"
- "我一直以为自己做得对，直到那件事发生。"
- "以前别人说我太老实，我还不服气，现在我服了。"
- "我曾经以为善良是优点，后来才知道，善良过头就是缺点。"
- "年轻时我不懂，现在我懂了，可惜晚了。"
- "我以前总是忍，以为忍一忍就过去了，结果呢？"

## 引经据典要求（{article_count}篇必须从不同类别中选择！）

**禁止高频使用的引用（太常见了，用了就重写！）：**
- ❌ 《道德经》"上善若水"
- ❌ 《增广贤文》"谁人背后无人说"
- ❌ 曾国藩的任何名言（用太多了）
- ❌ 杨绛"人生最曼妙的风景"
- ❌ 王阳明"此心光明"
- ❌ 莫言的任何名言（用太多了）
- ❌ 螃蟹效应、破窗效应（用太多了）

**引用类型库（10大类，每篇必须从不同类别中选择！）：**

**类型1-儒家经典：**
- 《论语》："君子坦荡荡，小人长戚戚"
- 《论语》："己所不欲，勿施于人"
- 《孟子》："穷则独善其身，达则兼济天下"
- 《孟子》："生于忧患，死于安乐"
- 《大学》："知止而后有定，定而后能静"
- 《中庸》："君子素其位而行，不愿乎其外"

**类型2-道家智慧：**
- 《道德经》："知人者智，自知者明"
- 《道德经》："大直若屈，大巧若拙"
- 《道德经》："祸兮福之所倚，福兮祸之所伏"
- 《庄子》："井蛙不可以语于海者，拘于虚也"
- 《庄子》："相濡以沫，不如相忘于江湖"
- 《列子》："天地无全功，圣人无全能"

**类型3-处世格言：**
- 《菜根谭》："宠辱不惊，闲看庭前花开花落"
- 《菜根谭》："路径窄处，留一步与人行"
- 《围炉夜话》："十分不耐烦，乃为人大病"
- 《小窗幽记》："宠辱不惊，看庭前花开花落"
- 《幽梦影》："少年人须有老成之识见，老成人须有少年之襟怀"
- 《呻吟语》："轻信轻发，听言之大戒也"

**类型4-史书智慧：**
- 《史记》："桃李不言，下自成蹊"
- 《史记》："燕雀安知鸿鹄之志"
- 《资治通鉴》："兼听则明，偏信则暗"
- 《战国策》："狡兔三窟，仅得免其死耳"
- 《左传》："多行不义必自毙"
- 《汉书》："水至清则无鱼，人至察则无徒"

**类型5-诗词名句：**
- 苏轼："竹杖芒鞋轻胜马，谁怕？一蓑烟雨任平生"
- 苏轼："人有悲欢离合，月有阴晴圆缺"
- 辛弃疾："众里寻他千百度，蓦然回首，那人却在灯火阑珊处"
- 陆游："山重水复疑无路，柳暗花明又一村"
- 李白："长风破浪会有时，直挂云帆济沧海"
- 杜甫："会当凌绝顶，一览众山小"
- 王维："行到水穷处，坐看云起时"
- 白居易："试玉要烧三日满，辨材须待七年期"

**类型6-明清名人：**
- 王阳明《传习录》："破山中贼易，破心中贼难"
- 王阳明："你未看此花时，此花与汝同归于寂"
- 张居正："谋之在众，断之在独"
- 于谦："粉身碎骨浑不怕，要留清白在人间"
- 左宗棠："好便宜者，不可与之交财"
- 林则徐："海纳百川，有容乃大"
- 纪晓岚："事能知足心常惬，人到无求品自高"

**类型7-民国大家：**
- 鲁迅："真的猛士，敢于直面惨淡的人生"
- 鲁迅："不在沉默中爆发，就在沉默中灭亡"
- 胡适："做学问要在不疑处有疑，待人要在有疑处不疑"
- 梁启超："患难困苦，是磨炼人格之最高学校"
- 林语堂："人生不过如此，且行且珍惜"
- 丰子恺："不乱于心，不困于情，不畏将来，不念过往"

**类型8-当代作家：**
- 杨绛《走到人生边上》："一个人经过不同程度的锻炼，就获得不同程度的修养"
- 季羡林："人生在世，不如意事十之八九，常想一二，不思八九"
- 余华《活着》："人是为活着本身而活着，而不是为了活着之外的任何事物所活着"
- 史铁生："命定的局限尽可永在，不屈的挑战却不可须臾或缺"
- 路遥《平凡的世界》："生活不能等待别人来安排，要自己去争取和奋斗"
- 汪曾祺："人间送小温，不要太多，不要太少"
- 贾平凹："人的一生，苦也罢，乐也罢，最重要的是心间的一泓清泉里不能没有月辉"
- 三毛："一个人至少拥有一个梦想，有一个理由去坚强"

**类型9-当代名人金句：**
- 白岩松："一个人的价值，不是看他拥有多少，而是看他能给予多少"
- 罗翔："一个知识越贫乏的人，越是拥有一种莫名奇怪的勇气和自豪感"
- 董宇辉："人生就是一场修行，修的是一颗心"
- 任正非："烧不死的鸟就是凤凰"
- 曹德旺："人这一辈子，吃亏就是福"
- 褚时健："人生总有起落，精神终可传承"
- 稻盛和夫："人生的意义在于提升心性，磨炼灵魂"

**类型10-心理学/哲学概念：**
- 马斯洛需求层次："人在满足基本需求后，会追求更高层次的自我实现"
- 沉没成本效应："已经付出的成本不应该影响未来的决策"
- 幸存者偏差："我们只看到成功者，却忽略了更多失败者"
- 登门槛效应："先接受小要求，更容易接受大要求"
- 鸟笼效应："人们会为了一个不需要的东西，再添置更多不需要的东西"
- 刺猬法则："保持适当距离，才能既温暖又不伤害"
- 蘑菇定律："新人都要经历一段被忽视的时期，熬过去就好了"
- 飞轮效应："万事开头难，但一旦转动起来，就会越来越轻松"
- 延迟满足："能够忍耐当下的诱惑，才能获得更大的回报"
- 峰终定律："人们对一段经历的评价，取决于高峰和结尾时的感受"

**引用规则（必须严格执行！）：**
- {article_count}篇文章必须从{article_count}个不同的大类中选择引用
- 每篇文章至少引用2处，且来自不同来源
- 禁止连续两次生成使用相同的引用
- 引用要自然融入文章，不能生硬堆砌
- 引用后要有自己的解读和延伸，不能只是引用完就结束

## 内容结构（确保每篇字数在{word_count}~{int(word_count) + 100}字）

1. **开头（约400字）**：独特钩子 + 第二人称描写读者处境 + 情感共鸣
2. **处境展开（约500字）**：用"你"描写读者的经历、感受、委屈（禁止讲具体故事）
3. **道理阐述（约400字）**：引经据典 + 深度分析 + 人生洞察
4. **收尾引流（约350字）**：总结升华 + 给出希望 + 引流话术

## 标题要求（15个标题必须全新原创！根据文案内容生成！）

**标题核心原则（最重要！）：**
- **必须根据当前文案内容原创标题**，禁止使用固定模板
- **标题要用第二人称"你"或共情表达**：让读者觉得说的就是自己
- **标题要夸赞/肯定/共情读者**：认同、理解、赞美、懂得读者
- **标题要短**：控制在10-20字
- **标题要正面积极**：给读者温暖和力量
- **让读者看了开心、被认可、被理解、被看见**
- **禁止用悬念、反问、疑问句**
- **15个标题必须完全不同，每次生成都必须全新，禁止和之前生成的标题雷同**

**爆款标题公式（6大类30种公式，必须混合使用，增加多样性！）：**

**【A类：第二人称夸赞型】**
公式1：你+品质+肯定词
- 你的善良，终会被这世界温柔以待
公式2：你这种人+赞美结果
- 你这种人，活该被人尊重
公式3：像你这样的人+肯定
- 像你这样的人，真的不多了
公式4：你吃过的苦/扛过的事+好结果
- 你吃过的苦，都会变成光照亮前路
公式5：你值得+美好事物
- 你值得被好好爱着

**【B类：共情理解型】**
公式6：有一种人+描述+那就是你
- 有一种人，默默付出从不邀功，说的就是你
公式7：懂你的人+都知道你的好
- 懂你的人，都知道你有多不容易
公式8：不是所有人都能+做到这些
- 不是所有人都能像你这样扛下来
公式9：这世上+最[形容词]的人+就是你这样的
- 这世上最靠谱的人，就是你这样的
公式10：能做到这些的人+真的很少
- 能把委屈咽下去还笑着的人，真的很少

**【C类：场景画面型】**
公式11：当所有人都[负面行为]时+你却[正面行为]
- 当所有人都在抱怨时，你却在默默努力
公式12：别人[行为A]+你却[行为B]
- 别人在享受时，你在扛着
公式13：那些[时间/场景]+你一个人[行为]
- 那些深夜里，你一个人扛过了多少
公式14：[场景描述]+你从不说苦
- 再难的日子，你从不说苦
公式15：[困境]+你硬是[结果]
- 没人帮忙，你硬是撑过来了

**【D类：对话温暖型】**
公式16：想告诉你+[温暖的话]
- 想告诉你，你已经很棒了
公式17：别怕+[鼓励的话]
- 别怕，你值得被好好对待
公式18：辛苦了+[肯定的话]
- 辛苦了，你的付出都被看见了
公式19：嘿+[亲切的话]
- 嘿，你比自己想象的更厉害
公式20：记住+[祝福的话]
- 记住，好事正在向你走来

**【E类：转折反差型】**
公式21：你不争不抢+却[好结果]
- 你不争不抢，却赢得了所有人的尊重
公式22：你从不炫耀+但[事实]
- 你从不炫耀，但实力有目共睹
公式23：表面[A]+其实[B正面]
- 表面平平无奇，其实你早已不凡
公式24：看起来[A]+其实[B正面]
- 看起来普普通通，其实你藏着大智慧
公式25：不声不响+却[好结果]
- 不声不响，却活成了别人羡慕的样子

**【F类：结果导向型】**
公式26：你现在的[付出]+以后都会[回报]
- 你现在的努力，以后都会变成惊喜
公式27：那些[经历]+都会成为[正面结果]
- 那些默默坚持的日子，都会变成礼物
公式28：时间会证明+[肯定的结果]
- 时间会证明，你的选择是对的
公式29：总有一天+[美好结果]
- 总有一天，你会感谢现在努力的自己
公式30：属于你的[好事]+[正在/一定会]来
- 属于你的幸福，正在路上

**【G类：金句哲理型】（不用"你"，用普世道理）**
公式31：[人生道理]+[正面洞察]
- 善良的人，日子都不会太差
公式32：[现象描述]+才是[本质]
- 不声不响扛下所有的人，才是真正的强者
公式33：真正的[品质]+是[具体表现]
- 真正的智慧，是懂得善待自己
公式34：[名人/古语]+说得好：[金句]
- 老话说得好：厚德载物，好事自来
公式35：世上最[形容词]的事+就是[描述]
- 世上最幸福的事，就是被人懂得

**【H类：故事悬念型】（引发好奇，正面结局）**
公式36：[人物]+[经历]+[好结果]
- 那个默默坚持的人，后来都怎样了
公式37：[数字]+年后+[好变化]
- 十年后再看，当初的坚持是对的
公式38：[某类人]+最后都[好结果]
- 心善的人，最后都过得不错
公式39：[时间点]+终于等来+[好事]
- 熬过低谷的人，终于等来了好消息
公式40：[坚持]+之后+[收获]
- 默默努力之后，终于迎来了转机

**【I类：现象洞察型】（客观描述，正面共鸣）**
公式41：越是[特质]的人+越[好现象]
- 越是善良的人，越过得顺心
公式42：[某类人]+往往[好特点]
- 不爱说话的人，往往最靠谱
公式43：有一种[品质]+叫做[正面描述]
- 有一种幸福，叫做问心无愧
公式44：[行为]+的人+[正面评价]
- 默默付出的人，总会被看见
公式45：这世上+[正面现象]
- 这世上，好人终有好报

**【J类：情感共鸣型】（第三人称，正面现象）**
公式46：有些人+[描述]+[好结果]
- 有些人不争不抢，却什么都有了
公式47：最美的+[情感描述]
- 最美的遇见，是有人懂你的不容易
公式48：[付出]+终会换来+[好结果]
- 真心付出的人，终会遇到珍惜你的人
公式49：[某种人]+活得最[好状态]
- 心里有光的人，活得最自在
公式50：宁可+[好选择]+也要+[好结果]
- 宁可慢一点，也要稳稳地幸福

**【K类：夸赞赞美型】（直接夸赞，混合人称）**
公式51：能做到[某事]的人+都不简单
- 能把苦咽下去还笑着的人，都不简单
公式52：[品质]+是最稀缺的财富
- 真诚，是这个时代最稀缺的财富
公式53：像这样的人+[赞美结果]
- 像这样的人，走到哪里都会发光
公式54：[某种人]+才是真正的强者
- 不声不响扛下所有的人，才是真正的强者
公式55：这份[品质]+让人敬佩
- 这份隐忍和坚持，让人敬佩
公式56：世上最[形容词]的人+就是[描述]
- 世上最了不起的人，就是熬过苦难还心存善良的人
公式57：[某种品质]+比什么都珍贵
- 一颗干净的心，比什么都珍贵

**【L类：肯定认可型】（肯定选择和努力，混合人称）**
公式58：这条路+没有走错
- 这条路没有走错，只是需要时间证明
公式59：所有的付出+都是值得的
- 所有默默的付出，都是值得的
公式60：已经做得很好了+别再苛责自己
- 已经做得很好了，别再苛责自己
公式61：坚持[某事]+是对的
- 坚持做自己，永远是对的
公式62：善良+终会有回报
- 善良的人，终会有回报
公式63：不是不够好+是[真相]
- 不是不够好，是有些人不懂珍惜
公式64：走过的路+[肯定评价]
- 走过的每一步，都算数

**【M类：期许祝福型】（美好祝愿，混合人称）**
公式65：愿[某种人]+[美好祝愿]
- 愿善良的人，都被温柔以待
公式66：好事+正在路上
- 好事正在路上，别着急
公式67：值得+[美好事物]
- 值得拥有世间所有的美好
公式68：往后余生+[祝福]
- 往后余生，不再委屈自己
公式69：属于[某种人]的幸福+一定会来
- 属于老实人的幸福，一定会来
公式70：再等等+[好结果]
- 再等等，柳暗花明就在前面
公式71：好日子+还在后头
- 好日子还在后头，别灰心
公式72：愿所有[付出]+都能[好结果]
- 愿所有的坚持，都能开花结果

**【N类：扬眉吐气型】（第三人称衬托，让读者爽）**
公式73：那些[负面行为]的人+现在[反转结果]
- 那些笑话你的人，现在都闭嘴了
公式74：曾经[看不起你的人]+开始[后悔/羡慕]
- 曾经看不起你的人，开始后悔了
公式75：他们怕的就是+你[变强的表现]
- 他们怕的就是你不再忍让
公式76：那些否定你的人+[反转结果]
- 那些否定你的人，现在都在羡慕你
公式77：让[某些人]+看看+[你的厉害]
- 让那些小瞧你的人，好好看看
公式78：你的[改变]+让他们[反应]
- 你的强大，让他们开始害怕了
公式79：[时间]后+那些人+[反转]
- 三年后，那些嘲笑你的人都沉默了
公式80：他们以为你[负面]+没想到你[正面]
- 他们以为你会认输，没想到你越挫越勇

**标题情感方向（必须让读者感到）：**
- ✅ 被看见："原来有人懂我"
- ✅ 被认可："我的付出是值得的"
- ✅ 被夸赞："原来我还挺厉害的"
- ✅ 被祝福："好事要来了"
- ✅ 被理解："终于有人理解我的苦"
- ✅ 被治愈："心里暖暖的"
- ✅ 被肯定："我没有白努力"
- ✅ 扬眉吐气："那些看不起我的人要后悔了"

**标题禁忌：**
- ❌ 不要用疑问句、反问句
- ❌ 不要制造悬念
- ❌ 不要用"为什么""怎么""吗"等疑问词
- ❌ 不要打击、质疑读者
- ❌ 不要用负面情绪词汇
- ❌ 禁止用封建迷信词汇（天选、宇宙、仙师、菩萨、佛祖等）
- ❌ 禁止照抄示例，必须根据文案内容原创

**标题生成规则：**
- 根据每篇文案的具体内容，使用上面的公式框架原创5个标题
- **每篇文章的5个标题必须来自不同的类别（A-N类共14类）**，增加多样性
- **必须混合使用：至少2个用"你"的标题 + 至少2个不用"你"的标题**
- {article_count}篇文章共{article_count * 5}个标题，必须全部不同，全部原创
- **禁止连续使用同一类型的公式**，必须交叉混合
- G/H/I/J类是不用"你"的标题，K/L/M类是混合人称的夸赞肯定期许类标题，N类是第三人称衬托让读者爽的标题
- 全部是肯定、夸赞、共情、祝福、洞察、扬眉吐气的语气
- 标题要与文案内容紧密相关，不能泛泛而谈
- **每次生成的标题必须与之前生成的标题有明显区别，避免雷同**
- 禁止出现任何重复、相似、雷同的标题
- 禁止照抄示例，必须根据文案内容原创

## 敏感词规避（以下词汇绝对不能出现！）

- 玄学类：仙师、徒儿、为师、天机、命盘、磁场、气运、福报、姻缘、符咒、渡劫、运势、运气、贵人、小人、业障、因果报应、天命、命数、劫难、气场、能量场
- 迷信类：算命、占卜、风水、转运、开光、法事、神仙、菩萨、鬼神、法术、阴阳、五行、宇宙能量
- 人设类：为师、贫道、本仙师、师傅、徒儿、道友、有缘人
- 承诺类：改运、转运、翻身、暴富

## 引流类型：{flow_type}
{flow_instruction}

## 输出格式
═══════════════════════════════════════
【第一篇】
═══════════════════════════════════════

【标题】
标题内容1
标题内容2
标题内容3
标题内容4
标题内容5

---

正文内容...（必须控制在{word_count}~{int(word_count) + 100}字）

═══════════════════════════════════════
【第二篇】
═══════════════════════════════════════
...（以此类推，共{article_count}篇）

注意：标题部分只输出5行干净的标题文字，每行一个标题，不要带"标题1："等前缀，方便直接复制使用。

请直接输出仿写结果，不要有任何说明性文字。"""

        return prompt

    def get_flow_instruction(self, flow_type, yinliu_content, product_name, product_material, article_count=3):
        """获取引流类型的具体指令"""
        if flow_type == "置顶引流":
            instruction = f"""## 结尾引流方式：置顶引流

**【最重要！结尾多样化要求】：**
- **{article_count}篇文案的结尾必须完全不同！禁止复制粘贴！**
- 用户提供的话术只是**参考风格和方向**，不是让你直接照抄
- 每篇结尾都要根据该篇文案的具体内容，原创一个全新的结尾
- {article_count}篇结尾要使用不同的切入角度、不同的铺垫方式、不同的表达
- 结尾要与该篇正文内容自然衔接，不能生硬套用模板

**引流的核心目标：让读者不去看就睡不着觉、吃不下饭、觉得亏了100万！**

**引流铺垫要求：**
- 不能只是简单说"去看主页置顶"，必须说清楚**为什么要去看**
- 要制造强烈的好奇心和紧迫感
- 要让读者觉得：不去看就错过了改变人生的机会
- 引流话术要与前文内容自然衔接，有理有据

**置顶引流的钩子类型（{article_count}篇分别用不同类型）：**
1. **悬念型**：解决方法/核心秘诀在置顶视频里
2. **专属型**：专门为"像你这样的人"准备的内容
3. **价值型**：最核心的思路和方法在置顶
4. **紧迫型**：别等到事情更糟了才后悔
5. **救赎型**：那里有你一直在找的答案
6. **共鸣型**：我也曾经历过，经验都在置顶
7. **秘诀型**：关键点太重要了，专门录了视频

**用户提供的参考话术（学习风格，不要照抄！）：**"""
            if yinliu_content:
                instruction += f"\n{yinliu_content}"

        elif flow_type == "橱窗引流":
            instruction = f"""## 结尾引流方式：橱窗引流

**【最重要！结尾多样化要求】：**
- **{article_count}篇文案的结尾必须完全不同！禁止复制粘贴！**
- 用户提供的话术只是**参考风格和方向**，不是让你直接照抄
- 每篇结尾都要根据该篇文案的具体内容，原创一个全新的结尾
- {article_count}篇结尾要使用不同的切入角度、不同的铺垫方式、不同的表达
- 结尾要与该篇正文内容自然衔接，不能生硬套用模板

**橱窗引流的钩子设计（{article_count}篇分别用不同类型）：**
1. **助力钩子**：那里有能帮助你的好物/工具
2. **改变钩子**：想要改变现状，需要一些助力
3. **犒赏钩子**：你值得对自己好一点
4. **能量钩子**：一件对的物品能带来力量
5. **祝福钩子**：送你一份温暖的祝福

**结尾话术参考（仅供参考风格，必须原创改写！）：**

1. **助力型风格：**
参考方向：强调外在助力帮助稳住心神、找回状态

2. **犒赏型风格：**
参考方向：强调对自己好一点、犒赏辛苦的自己

3. **改变型风格：**
参考方向：强调行动起来、改变的第一步

4. **能量型风格：**
参考方向：强调好物带来力量、找回状态

5. **祝福型风格：**
参考方向：强调你值得最好的、送上祝福

**用户提供的参考话术（学习风格，不要照抄！）：**"""
            if yinliu_content:
                instruction += f"\n{yinliu_content}"

        elif flow_type == "带货引流":
            instruction = f"""## 结尾引流方式：带货引流

**【最重要！结尾多样化要求】：**
- **{article_count}篇文案的结尾必须完全不同！禁止复制粘贴！**
- 用户提供的话术只是**参考风格和方向**，不是让你直接照抄
- 每篇结尾都要根据该篇文案的具体内容，原创一个全新的结尾
- {article_count}篇结尾要使用不同的切入角度、不同的铺垫方式、不同的表达
- 结尾要与该篇正文内容自然衔接，不能生硬套用模板

**商品名称：{product_name}**
**产品素材：{product_material}**

**带货引流要求：**
- 引流话术要与前文内容自然衔接
- 要突出商品能解决读者的问题/满足读者的需求
- 要制造紧迫感和稀缺感
- 话术要真诚，不能太硬广

**用户提供的参考话术（学习风格，不要照抄！）：**"""
            if yinliu_content:
                instruction += f"\n{yinliu_content}"
            else:
                instruction += f"\n这款{product_name}真的很不错，点我头像进橱窗了解一下。"

        else:  # 纯夸赞不引流
            instruction = f"""## 结尾方式：纯夸赞不引流

**【最重要！结尾多样化要求】：**
- **{article_count}篇文案的结尾必须完全不同！禁止复制粘贴！**
- 每篇结尾都要根据该篇文案的具体内容，原创一个全新的结尾
- {article_count}篇结尾要使用不同的切入角度、不同的表达方式
- 结尾要与该篇正文内容自然衔接

**要求：**
- 纯夸赞读者，给予温暖和力量
- 不需要任何引流话术
- 结尾要让读者感到被理解、被肯定、被温暖
- 给读者希望和方向，让他们觉得"还有出路"
- 用温暖有力的话收尾，让读者感动"""

        return instruction

    def save_document(self, content, output_path, index):
        """保存为Word文档"""
        try:
            from docx import Document
            from docx.shared import Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = Document()

            # 设置标题
            title = doc.add_heading('百家号仿写文案', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 添加内容
            lines = content.split('\n')
            for line in lines:
                line = line.strip()
                if not line:
                    continue

                if line.startswith('═') or line.startswith('---'):
                    doc.add_paragraph('─' * 40)
                elif line.startswith('【第') and '篇】' in line:
                    p = doc.add_paragraph()
                    run = p.add_run(line)
                    run.bold = True
                    run.font.size = Pt(14)
                elif line.startswith('【标题'):
                    p = doc.add_paragraph(line)
                    p.runs[0].bold = True
                else:
                    doc.add_paragraph(line)

            # 生成文件名
            timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
            filename = f"{timestamp}_{index}.docx"
            filepath = os.path.join(output_path, filename)

            doc.save(filepath)
            self.log(f"文档已保存: {filename}")
            return filepath

        except ImportError:
            self.log("错误：请安装python-docx库 (pip install python-docx)")
            # 备用方案：保存为txt
            timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
            filename = f"{timestamp}_{index}.txt"
            filepath = os.path.join(output_path, filename)
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write(content)
            self.log(f"已保存为TXT: {filename}")
            return filepath

        except Exception as e:
            self.log(f"保存文档失败: {str(e)}")
            return None

    def append_reference_to_library(self, article, flow_type):
        """追加参考文案到素材库Excel"""
        try:
            import openpyxl
            from openpyxl import Workbook

            # 确保目录存在
            if not os.path.exists(MATERIAL_LIBRARY_DIR):
                os.makedirs(MATERIAL_LIBRARY_DIR)

            # 打开或创建Excel文件
            if os.path.exists(MATERIAL_LIBRARY_FILE):
                wb = openpyxl.load_workbook(MATERIAL_LIBRARY_FILE)
                ws = wb.active
            else:
                wb = Workbook()
                ws = wb.active
                ws.title = "爆款素材库"
                # 添加表头
                ws.append(["引流类型", "日期", "正文", "字数"])

            # 当前日期
            today = datetime.now().strftime("%Y-%m-%d")

            # 统计中文字数
            char_count = len(re.findall(r'[\u4e00-\u9fff]', article))

            # 追加
            ws.append([flow_type, today, article, char_count])

            # 保存
            wb.save(MATERIAL_LIBRARY_FILE)
            self.log(f"已追加参考文案到素材库（{char_count}字）")

        except Exception as e:
            self.log(f"追加素材库失败: {str(e)}")

    def finish_task(self):
        """完成任务，恢复UI状态"""
        self.is_running = False
        self.start_btn.config(state=tk.NORMAL)
        self.stop_btn.config(state=tk.DISABLED)
        # 如果有生成过的文案，启用重新生成按钮
        if self.last_articles:
            self.regenerate_btn.config(state=tk.NORMAL)
        self.update_status("处理完成")

        # 询问是否打开输出文件夹
        if messagebox.askyesno("完成", "文案生成完成！是否打开输出文件夹？"):
            self.open_output_folder()

    def open_output_folder(self):
        """打开输出文件夹"""
        output_path = self.output_path.get().replace('/', '\\')
        if os.path.exists(output_path):
            os.startfile(output_path)
        else:
            messagebox.showerror("错误", f"文件夹不存在: {output_path}")

    def start_synth_voice(self):
        """开始批量合成语音"""
        # 从页面获取文案目录
        txt_dir = self.voice_input_path.get().replace('/', '\\')

        if not os.path.exists(txt_dir):
            messagebox.showerror("错误", f"文案目录不存在: {txt_dir}")
            return

        # 获取所有txt文件
        txt_files = [f for f in os.listdir(txt_dir) if f.endswith('.txt')]
        if not txt_files:
            messagebox.showerror("错误", "文案目录中没有TXT文件")
            return

        # 获取音色和输出目录
        voice_type = self.voice_type.get()
        save_dir = self.voice_output_path.get().replace('/', '\\')

        # 确认
        if not messagebox.askyesno("确认", f"找到 {len(txt_files)} 个TXT文件\n音色: {voice_type}\n是否开始批量合成语音？"):
            return

        # 在新线程中执行
        self.synth_voice_btn.config(state=tk.DISABLED)
        thread = threading.Thread(target=self.synth_voice_task, args=(txt_dir, txt_files, voice_type, save_dir))
        thread.daemon = True
        thread.start()

    def synth_voice_task(self, txt_dir, txt_files, voice_type, save_dir):
        """批量合成语音任务"""
        import glob as glob_module

        # 配置
        BITBROWSER_API = "http://127.0.0.1:54345"
        BROWSER_ID = "fd66587b053346ddb01a3892cea21ceb"
        CHROMEDRIVER_PATH = r"C:\Users\Administrator\AppData\Roaming\BitBrowser\chromedriver\140\chromedriver.exe"
        DOWNLOAD_DIR = r"C:\Users\Administrator\Downloads\11"
        SAVE_DIR = save_dir
        STEP_DELAY = 2

        def get_all_files(directory):
            files = set()
            if not os.path.exists(directory):
                return files
            for f in os.listdir(directory):
                full_path = os.path.join(directory, f)
                if os.path.isfile(full_path) and not f.endswith('.crdownload') and not f.endswith('.tmp'):
                    files.add(full_path)
            return files

        def wait_for_new_file(directory, before_files, timeout=30):
            for _ in range(timeout):
                current_files = get_all_files(directory)
                new_files = current_files - before_files
                if new_files:
                    new_file = list(new_files)[0]
                    time.sleep(0.5)
                    return new_file
                time.sleep(1)
            return None

        def rename_file(old_path, new_name):
            if not old_path or not os.path.exists(old_path):
                return None
            directory = os.path.dirname(old_path)
            extension = os.path.splitext(old_path)[1]
            new_path = os.path.join(directory, new_name + extension)
            if os.path.exists(new_path):
                os.remove(new_path)
            os.rename(old_path, new_path)
            return new_path

        try:
            from selenium import webdriver
            from selenium.webdriver.chrome.service import Service
            from selenium.webdriver.chrome.options import Options

            self.log("=" * 50)
            self.log("开始批量合成语音")
            self.log("=" * 50)

            # 打开浏览器
            self.log("正在打开比特浏览器...")
            url = f"{BITBROWSER_API}/browser/open"
            data = {"id": BROWSER_ID}
            resp = requests.post(url, json=data)
            result = resp.json()

            if not result.get("success"):
                self.log(f"打开浏览器失败: {result}")
                return

            ws_url = result["data"]["ws"]
            port = ws_url.split(":")[2].split("/")[0]
            self.log(f"浏览器已打开，端口: {port}")

            time.sleep(STEP_DELAY)

            # 连接浏览器
            options = Options()
            options.add_experimental_option("debuggerAddress", f"127.0.0.1:{port}")
            service = Service(executable_path=CHROMEDRIVER_PATH)
            driver = webdriver.Chrome(service=service, options=options)

            # 确保在配音神器页面
            if "peiyinshenqi" not in driver.current_url:
                self.log("跳转到配音神器页面...")
                driver.get("https://peiyinshenqi.com/tts/index")
                time.sleep(3)

            # 处理每个TXT文件
            for idx, txt_file in enumerate(txt_files):
                filename = os.path.splitext(txt_file)[0]  # 去掉.txt扩展名
                txt_path = os.path.join(txt_dir, txt_file)

                self.log(f"\n[{idx+1}/{len(txt_files)}] 处理: {filename}")

                # 检查是否已存在配音文件
                existing_files = glob_module.glob(os.path.join(SAVE_DIR, f"{filename}.*"))
                if existing_files:
                    self.log(f"  跳过（已存在）: {filename}")
                    continue

                # 读取文案内容
                with open(txt_path, 'r', encoding='utf-8') as f:
                    text = f.read().strip()

                if not text:
                    self.log(f"  跳过（空文件）: {filename}")
                    continue

                # === 步骤1: 清空 ===
                js_clear = '''
                var links = document.querySelectorAll('.el-link--inner');
                for(var i=0; i<links.length; i++){
                    if(links[i].innerText.includes('清空')){
                        links[i].click();
                        return true;
                    }
                }
                return false;
                '''
                driver.execute_script(js_clear)
                time.sleep(1)

                # 处理弹窗
                js_confirm = '''
                var btns = document.querySelectorAll('.el-message-box__btns button');
                for(var i=0; i<btns.length; i++){
                    if(btns[i].innerText.includes('清除') || btns[i].innerText.includes('确定')){
                        btns[i].click();
                        return true;
                    }
                }
                return false;
                '''
                driver.execute_script(js_confirm)
                time.sleep(STEP_DELAY)

                # === 步骤2: 输入文案 ===
                js_click_editor = '''
                var editor = document.querySelector('.editor[contenteditable="true"]');
                if(editor){ editor.click(); editor.focus(); return true; }
                return false;
                '''
                driver.execute_script(js_click_editor)
                time.sleep(1)

                # 转义文案中的特殊字符
                escaped_text = text.replace('\\', '\\\\').replace('"', '\\"').replace('\n', '\\n').replace('\r', '')
                js_input = f'''
                var editor = document.querySelector('.editor[contenteditable="true"]');
                if(editor){{
                    editor.innerText = "{escaped_text}";
                    editor.dispatchEvent(new Event('input', {{bubbles: true}}));
                    return true;
                }}
                return false;
                '''
                driver.execute_script(js_input)
                time.sleep(STEP_DELAY)

                # === 步骤3: 选择音色 ===
                # 根据页面选择的音色来点击对应的元素
                js_voice = f'''
                var items = document.querySelectorAll('.voice-name');
                for(var i=0; i<items.length; i++){{
                    if(items[i].innerText.includes('{voice_type}')){{
                        var parent = items[i].closest('.sub-item-txt') || items[i].parentElement;
                        if(parent){{ parent.click(); return '已选择: ' + items[i].innerText; }}
                    }}
                }}
                return '未找到音色: {voice_type}';
                '''
                result = driver.execute_script(js_voice)
                self.log(f"  音色: {result}")
                time.sleep(STEP_DELAY)

                # === 步骤4: 点击合成配音 ===
                js_synthesis = '''
                var btns = document.querySelectorAll('.el-button--primary');
                for(var i=0; i<btns.length; i++){
                    if(btns[i].innerText.includes('合成配音')){
                        btns[i].click();
                        return true;
                    }
                }
                return false;
                '''
                driver.execute_script(js_synthesis)
                time.sleep(STEP_DELAY)

                # === 步骤7: 点击开始合成 ===
                js_confirm_synthesis = '''
                var btns = document.querySelectorAll('.el-button--primary');
                for(var i=0; i<btns.length; i++){
                    if(btns[i].innerText.includes('开始合成')){
                        btns[i].click();
                        return true;
                    }
                }
                return false;
                '''
                driver.execute_script(js_confirm_synthesis)
                time.sleep(STEP_DELAY)

                # === 步骤8: 等待合成完成 ===
                self.log("  等待合成完成...")
                max_wait = 180
                for i in range(max_wait):
                    js_check_loading = '''
                    var mask = document.querySelector('.el-loading-mask.is-fullscreen');
                    if(mask){
                        var style = window.getComputedStyle(mask);
                        if(style.display !== 'none'){
                            var text = mask.querySelector('.el-loading-text');
                            return text ? text.innerText : 'loading';
                        }
                    }
                    return null;
                    '''
                    loading_status = driver.execute_script(js_check_loading)
                    if loading_status is None:
                        self.log(f"  合成完成 ({i+1}秒)")
                        break
                    time.sleep(1)
                else:
                    self.log("  合成超时，跳过")
                    continue

                time.sleep(STEP_DELAY)

                # === 步骤9: 下载配音 ===
                before_files = get_all_files(DOWNLOAD_DIR)

                js_download = '''
                var btns = document.querySelectorAll('.el-button');
                for(var i=0; i<btns.length; i++){
                    if(btns[i].innerText.includes('下载配音')){
                        btns[i].click();
                        return true;
                    }
                }
                return false;
                '''
                driver.execute_script(js_download)

                new_file = wait_for_new_file(DOWNLOAD_DIR, before_files, timeout=30)
                mp3_file = None
                if new_file:
                    mp3_file = rename_file(new_file, filename)
                    self.log(f"  配音已下载: {os.path.basename(mp3_file)}")

                time.sleep(STEP_DELAY)

                # === 步骤10: 下载字幕 ===
                time.sleep(1)
                before_files = get_all_files(DOWNLOAD_DIR)

                js_download_subtitle = '''
                var previewBtns = document.querySelectorAll('button[classs="preview-btn"]');
                for(var i=0; i<previewBtns.length; i++){
                    if(previewBtns[i].innerText.includes('下载字幕')){
                        previewBtns[i].click();
                        return true;
                    }
                }
                var icons = document.querySelectorAll('.el-icon-chat-dot-square');
                for(var i=0; i<icons.length; i++){
                    var btn = icons[i].closest('button');
                    if(btn){ btn.click(); return true; }
                }
                return false;
                '''
                driver.execute_script(js_download_subtitle)

                # 等待字幕解析
                time.sleep(1)
                for i in range(60):
                    js_check_loading = '''
                    var mask = document.querySelector('.el-loading-mask.is-fullscreen');
                    if(mask){
                        var style = window.getComputedStyle(mask);
                        if(style.display !== 'none'){ return 'loading'; }
                    }
                    return null;
                    '''
                    if driver.execute_script(js_check_loading) is None:
                        break
                    time.sleep(1)

                new_file = wait_for_new_file(DOWNLOAD_DIR, before_files, timeout=30)
                srt_file = None
                if new_file:
                    srt_file = rename_file(new_file, filename)
                    self.log(f"  字幕已下载: {os.path.basename(srt_file)}")

                # === 步骤11: 转移文件 ===
                import shutil

                # 判断是否下载成功（至少配音文件要存在）
                download_success = mp3_file and os.path.exists(mp3_file)

                if download_success:
                    if not os.path.exists(SAVE_DIR):
                        os.makedirs(SAVE_DIR)

                    dst = os.path.join(SAVE_DIR, os.path.basename(mp3_file))
                    if os.path.exists(dst):
                        os.remove(dst)
                    shutil.move(mp3_file, dst)
                    self.log(f"  配音已转移")

                    if srt_file and os.path.exists(srt_file):
                        dst = os.path.join(SAVE_DIR, os.path.basename(srt_file))
                        if os.path.exists(dst):
                            os.remove(dst)
                        shutil.move(srt_file, dst)
                        self.log(f"  字幕已转移")

                    # === 步骤12: 移动TXT到回收站（只有下载成功才移动） ===
                    TXT_RECYCLE_DIR = txt_dir.replace("视频文案", "回收站")
                    if TXT_RECYCLE_DIR == txt_dir:
                        TXT_RECYCLE_DIR = os.path.join(os.path.dirname(txt_dir), "回收站", os.path.basename(txt_dir))
                    if not os.path.exists(TXT_RECYCLE_DIR):
                        os.makedirs(TXT_RECYCLE_DIR)
                    txt_dst = os.path.join(TXT_RECYCLE_DIR, txt_file)
                    if os.path.exists(txt_dst):
                        os.remove(txt_dst)
                    shutil.move(txt_path, txt_dst)
                    self.log(f"  TXT已移到回收站")

                    self.log(f"  完成: {filename}")
                else:
                    self.log(f"  下载失败，跳过: {filename}")

            self.log("\n" + "=" * 50)
            self.log("批量合成语音完成！")
            self.log("=" * 50)

            self.root.after(0, lambda: messagebox.showinfo("完成", "批量合成语音完成！"))

        except Exception as e:
            self.log(f"错误: {str(e)}")
            import traceback
            self.log(traceback.format_exc())
        finally:
            self.root.after(0, lambda: self.synth_voice_btn.config(state=tk.NORMAL))


def main():
    root = tk.Tk()
    app = FangxieApp(root)
    root.mainloop()


if __name__ == "__main__":
    main()
