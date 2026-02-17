# -*- coding: utf-8 -*-
"""
百家号文案仿写工具 - 可视化界面
功能：读取参考文案，根据选择的引流类型生成仿写文案
支持流式/非流式调用，主模型+备用模型切换
"""

import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext
import os
import json
import threading
from datetime import datetime
import requests
import re
import time
import random

# 开头钩子库 - 用于随机选择
HOOK_LIBRARY = {
    "A类-轻松幽默型": [
        "我发现一个规律：越是老实人，越容易被安排加班。你说气不气？",
        "昨天算了一笔账，这些年帮别人花的时间，够我学会三门外语了。",
        "我妈说我最大的优点是善良，我爸说这也是我最大的缺点。好家伙，亲爹。",
        "朋友说我是'便利店型人格'——24小时营业，随叫随到，还不涨价。",
        "有人说我脾气好，我笑了笑没说话。其实不是脾气好，是懒得计较。",
        "我终于明白，为什么'好人卡'发得最多——因为好人最好打发。",
        "同事问我：'你怎么从来不生气？'我说：'生气要花力气，我选择省着点用。'",
        "我的人生信条曾经是'吃亏是福'，直到我发现福没来，亏倒是吃了不少。",
        "别人都在研究怎么赚钱，我在研究怎么不被人当免费劳动力。",
        "我这人有个毛病，别人一说'就你能帮我'，我就跟中了蛊似的。",
    ],
    "B类-温暖治愈型": [
        "嘿，今天想跟你聊点轻松的，关于那些默默付出却不求回报的人。",
        "你知道吗，这世上有一种人，他们的好，是藏在细节里的。",
        "我一直觉得，善良的人身上有光，只是有时候这光被辜负了。",
        "如果你正在看这段话，我想告诉你：你的好，有人看得见。",
        "今天不讲大道理，就想跟你说说心里话。",
        "有些人，值得被这个世界温柔以待，比如正在看这段话的你。",
        "我见过很多人，但像你这样的，真的不多。",
        "你有没有被人夸过'你人真好'？今天我想认真聊聊这件事。",
        "这段话，送给每一个在生活里默默扛着的人。",
        "我相信，看到这里的你，一定是个心里有温度的人。",
    ],
    "C类-反转惊喜型": [
        "我曾经以为自己是个'老好人'，后来发现，我是个'聪明的好人'。",
        "都说老实人吃亏，但我认识一个老实人，现在过得比谁都好。",
        "你以为善良是软弱？不，善良是一种选择，而且是强者的选择。",
        "有人说心软的人没出息，我偏不信这个邪。",
        "我见过最厉害的人，恰恰是最善良的那个。",
        "别人都说我太老实会吃亏，结果呢？我还真没亏。",
        "都说好人没好报，但我今天要讲一个好人有好报的故事。",
        "我以前觉得'人善被人欺'是真理，直到我遇见了一个人。",
        "谁说善良的人就要受委屈？我第一个不服。",
        "老实人的春天，其实一直都在，只是很多人没发现。",
    ],
    "D类-对话引入型": [
        "他说完这句话，我愣在原地半天没回过神：'你就是太好说话了。'",
        "有人问我：'你这辈子最后悔的事是什么？'我想了想，说了两个字。",
        "算了，不争了——你是不是也经常这样跟自己说？",
        "我爷爷临终前拉着我的手说：'记住，吃亏的人，老天爷都记着账呢。'",
        "'你怎么这么傻？'这句话，你听过多少次了？",
        "我妈常说一句话：'人善被人欺，马善被人骑。'我以前不信，现在信了。",
        "'别太老实了，会吃亏的。'说这话的人，后来怎么样了？",
        "朋友跟我说：'你知道你最大的问题是什么吗？就是太把别人当回事。'",
        "'谢谢你'——这三个字，你等了多久才听到？",
        "有人当面问我：'你是不是傻？人家都欺负到头上了你还忍？'",
    ],
    "E类-生活观察型": [
        "菜市场大妈的一句话，让我愣了半天：'姑娘，你这么好说话，不怕被人欺负啊？'",
        "堵在路上，收音机里突然传来一句话，我眼眶一下就红了。",
        "排队的时候，前面两个人的对话让我心里一惊。",
        "超市结账时，前面那个人的一个举动，让我看清了人性。",
        "参加同学聚会，有个人的变化让我震惊了。",
        "小区门口，两个大妈聊天，聊着聊着说出了一个真相。",
        "去参加婚礼，新郎的一句话让全场安静了。",
        "年夜饭桌上，我爸突然放下筷子，说了一句话。",
        "早上买早餐，老板娘的一句话让我想了一整天。",
        "坐出租车，司机师傅跟我聊了一路，最后一句话让我沉默了。",
    ],
    "F类-数字锚定型": [
        "认识老王二十年了，他教会我一件事：别对谁都掏心掏肺。",
        "被人欺负了三年，我终于想通了一个道理。",
        "用了十年时间，我才学会一个字：不。",
        "三次被人背叛之后，我悟了。",
        "五十岁之后，我才明白什么叫'人走茶凉'。",
        "帮了他七年，他一句谢谢都没说过。",
        "借出去的三万块，要了五年都没要回来。",
        "在这个单位干了八年，我终于明白了一个道理。",
    ],
    "G类-金句破题型": [
        "最傻的事，就是跟烂人讲道理。",
        "你知道什么人最可怕吗？不是坏人，是那些笑着捅你刀子的人。",
        "千万别做老好人，我吃过这个亏，现在告诉你。",
        "你越忍让，别人越得寸进尺，这是我用十年换来的教训。",
        "老实人不是没脾气，是把脾气都咽进了肚子里。",
        "你的善良，要带点锋芒。",
        "不是所有的忍让都叫大度，有时候那叫窝囊。",
        "这世上最傻的事，就是把真心给了不值得的人。",
        "有些人，你帮他一百次，他记不住；你拒绝他一次，他记你一辈子。",
    ],
    "H类-悬念钩子型": [
        "有件事我憋了很久，今天必须说出来。",
        "你可能不信，但接下来我说的都是真事。",
        "我要告诉你一个很多人不愿意承认的真相。",
        "接下来这段话，可能会让你不舒服，但我还是要说。",
        "有个规律，我观察了很多年才看透。",
        "今天说的这些话，可能会得罪人，但我不在乎。",
        "有些话，我本来不想说，但看到你，我忍不住了。",
        "接下来的话，你可能不爱听，但句句都是真的。",
        "有件事，我一直没跟任何人说过，今天破例。",
    ],
    "I类-共鸣代入型": [
        "那种被人当众下面子的感觉，我太懂了。",
        "被最信任的人捅刀子，那种滋味，经历过的人都懂。",
        "明明没做错什么，却总被人针对，这种事我也遇到过。",
        "有些委屈，说出来都没人信。",
        "你是不是也有过这种感觉：付出最多的人，往往最不被珍惜。",
        "那种心寒的感觉，我懂。就像一盆冷水从头浇到脚。",
        "有一种苦，叫做'打碎了牙往肚子里咽'。",
        "那种被人利用完就扔掉的感觉，我经历过。",
    ],
    "J类-人物故事型": [
        "我有个朋友，前两天跟我说了一件事，我听完沉默了很久。",
        "我爸这辈子只教过我一个道理，我到现在才真正理解。",
        "单位有个人，大家都不待见他，后来我才知道原因。",
        "我们小区有个大爷，天天在楼下坐着，有一天他跟我说了一番话。",
        "我表姐的经历，让我彻底看清了人心。",
        "我有个同学，当年是班里最老实的人，你猜他现在怎么样了？",
        "我舅舅年轻时吃过一个大亏，他把这个教训告诉了我。",
        "我认识一个人，他的经历让我相信：好人终有好报。",
    ],
    "K类-自我剖析型": [
        "说出来不怕你笑话，我以前也是个傻子。",
        "回头看看这些年，我最后悔的一件事是太心软。",
        "如果能重来，我绝对不会再做老好人。",
        "我吃过的亏，今天全告诉你，希望你别再走我的老路。",
        "我这辈子最大的毛病，就是太把别人当回事。",
        "我曾经也是个'老好人'，后来我学聪明了。",
        "我年轻时犯过一个错，现在想起来还后悔。",
        "我以前总觉得吃亏是福，现在不这么想了。",
    ],
    "L类-转折反差型": [
        "以前我不信这个道理，直到自己栽了跟头。",
        "年轻的时候觉得这话是废话，现在才知道是真理。",
        "曾经有人跟我说过一句话，我没当回事，后来我后悔了。",
        "我一直以为自己做得对，直到那件事发生。",
        "以前别人说我太老实，我还不服气，现在我服了。",
        "我曾经以为善良是优点，后来才知道，善良过头就是缺点。",
        "年轻时我不懂，现在我懂了，可惜晚了。",
        "我以前总是忍，以为忍一忍就过去了，结果呢？",
    ],
}

def get_random_hooks():
    """随机选择3个不同类型的开头，确保至少1个是轻松/温暖/反转类型"""
    light_types = ["A类-轻松幽默型", "B类-温暖治愈型", "C类-反转惊喜型"]
    other_types = [k for k in HOOK_LIBRARY.keys() if k not in light_types]

    # 确保至少选1个轻松类型
    selected_types = [random.choice(light_types)]
    # 从剩余类型中再选2个
    remaining = [t for t in HOOK_LIBRARY.keys() if t != selected_types[0]]
    selected_types.extend(random.sample(remaining, 2))

    # 打乱顺序
    random.shuffle(selected_types)

    # 从每个类型中随机选一个示例
    result = []
    for t in selected_types:
        hook = random.choice(HOOK_LIBRARY[t])
        result.append({"type": t, "example": hook})

    return result

# 配置文件路径
CONFIG_FILE = os.path.join(os.path.dirname(__file__), "fangxie_config.json")

# 默认配置
DEFAULT_CONFIG = {
    "use_stream": True,  # 是否使用流式调用
    # 流式调用 - 主模型配置
    "stream_main_url": "https://api.aifuwu.icu/v1",
    "stream_main_key": "sk-hc6yUaXg89eK5UgUii10DPWmdaJZdPXqPbPcKSRbmWgxeeDK",
    "stream_main_model": "gemini-3-pro-preview",
    "stream_main_max_tokens": 16000,
    # 流式调用 - 备用模型配置
    "stream_backup_url": "https://yunyi.rdzhvip.com/v1",
    "stream_backup_key": "A68YGT53-41HW-8GT7-353U-MJVDNY8NE5KX",
    "stream_backup_model": "claude-opus-4-5-20251101",
    "stream_backup_max_tokens": 16000,
    # 非流式调用 - 主模型配置
    "non_stream_main_url": "https://yunyi.rdzhvip.com/v1",
    "non_stream_main_key": "A68YGT53-41HW-8GT7-353U-MJVDNY8NE5KX",
    "non_stream_main_model": "claude-opus-4-5-20251101",
    "non_stream_main_max_tokens": 16000,
    # 非流式调用 - 备用模型配置
    "non_stream_backup_url": "https://api.aifuwu.icu/v1",
    "non_stream_backup_key": "sk-hc6yUaXg89eK5UgUii10DPWmdaJZdPXqPbPcKSRbmWgxeeDK",
    "non_stream_backup_model": "gemini-3-pro-preview",
    "non_stream_backup_max_tokens": 16000,
    # 引流话术库（按类型分开存储）
    "yinliu_templates": {
        "置顶引流": [
            "关于怎么彻底走出这个困局，我在主页置顶视频里讲得很透。那里有一套方法，是我这些年摸爬滚打总结出来的，不适合在这里公开讲。你要是真想改变，点我头像，去看置顶第一条，看完你就明白该怎么做了。",
            "现在你就点开我的头像，进我主页，去看置顶的那个视频，我在那里给你留了一套破局的方法。听得懂那便是你的收获，听不懂说明时机还没到。准备好了吗？这一局该你赢了。",
            "从今天开始，学着为自己活一次... 答应我，现在点击我的头像进主页看看，置顶前两条视频，那里有我想对你说的心里话，这些话我只说给你听，因为你扛了太久，总该有人懂你的不易。",
            "如果你也想从这种困境里走出来，想活得通透一点、轻松一点，点我头像，去看主页置顶的视频。那里有你一直在找的答案，也有你需要的那份力量。我在那里等你。"
        ],
        "橱窗引流": [
            "现在你只需要做一件事，点开我的头像，进入主页橱窗，不要带着太多顾虑去挑拣，也不要问哪一件最好，你静静的看。若它让你心有所动，那就是适合你的，切莫错过。",
            "选一件，是给自己一份小小的犒赏；选两件，是为自己的生活增添一点温暖；若你愿意，选三件，就是给自己一个完整的礼物，让生活多一些美好。",
            "去吧，朋友，点开头像进橱窗，去找回那个本该发光的自己。我在那头等着，看你越来越好，自在如风。",
            "选一件，第一眼入心的就是你接下来全力冲锋的底气；选2件，便是对自己一路隐忍坚持的犒赏；选3件，更代表你不只顾着当下赶路，还在为长远未来布局筹谋。"
        ],
        "带货引流": [
            "轻轻一点左下角，它就能到你家。我知道你会犹豫，怕没用，怕白费钱，但我劝你大胆试这一次。这不是消费，是投资，投资你的安稳生活，投资你的人生底气。点击左下角把它请回家...",
            "好东西从不会长久停留，立刻点击左下角，让这份美好为你的生活增添一点温暖和力量...",
            "如今这几件东西我把它放到了左下角的通道里，别再犹豫了，去左下角把它请回家吧。这不是在帮别人，是在帮你自己... 对自己好一点，你值得。",
            "现在请顺应内心的指引，点击下方通道，把这款好物带回家。你为别人操心了大半辈子，是时候对自己好一点了。"
        ]
    }
}

def load_config():
    """加载配置"""
    if os.path.exists(CONFIG_FILE):
        try:
            with open(CONFIG_FILE, 'r', encoding='utf-8') as f:
                config = json.load(f)
                # 补充缺失的配置项
                for key in DEFAULT_CONFIG:
                    if key not in config:
                        config[key] = DEFAULT_CONFIG[key]
                return config
        except:
            pass
    return DEFAULT_CONFIG.copy()

def save_config(config):
    """保存配置"""
    with open(CONFIG_FILE, 'w', encoding='utf-8') as f:
        json.dump(config, f, ensure_ascii=False, indent=2)

class FangxieApp:
    def __init__(self, root):
        self.root = root
        self.root.title("百家号文案仿写工具")
        self.root.geometry("1000x900")
        self.root.resizable(True, True)

        # 加载配置
        self.config = load_config()

        # 素材库路径（固定）
        self.material_path = r"D:\A百家号带货文案库"

        # 引流类型映射
        self.flow_types = {
            "置顶引流": "置顶视频引流素材.txt",
            "橱窗引流": "橱窗引流素材.txt",
            "带货引流": "带货引流素材.txt",
            "纯夸赞不引流": None
        }

        self.create_widgets()

    def create_widgets(self):
        # 创建主框架（带滚动条）
        canvas = tk.Canvas(self.root)
        scrollbar = ttk.Scrollbar(self.root, orient="vertical", command=canvas.yview)
        scrollable_frame = ttk.Frame(canvas)

        scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )

        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)

        # 绑定鼠标滚轮
        def _on_mousewheel(event):
            canvas.yview_scroll(int(-1*(event.delta/120)), "units")
        canvas.bind_all("<MouseWheel>", _on_mousewheel)

        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")

        main_frame = ttk.Frame(scrollable_frame, padding="10")
        main_frame.pack(fill=tk.BOTH, expand=True)

        # 标题
        title_label = ttk.Label(main_frame, text="百家号文案仿写工具", font=("微软雅黑", 16, "bold"))
        title_label.pack(pady=(0, 15))

        # === API配置区域（可折叠） ===
        api_header_frame = ttk.Frame(main_frame)
        api_header_frame.pack(fill=tk.X, pady=5)

        # 流式/非流式开关 + 展开/收起按钮
        self.use_stream = tk.BooleanVar(value=self.config.get("use_stream", True))
        stream_check = ttk.Checkbutton(api_header_frame, text="使用流式调用（推荐）",
                                        variable=self.use_stream, command=self.on_stream_change)
        stream_check.pack(side=tk.LEFT)

        ttk.Label(api_header_frame, text="  |  ", foreground="gray").pack(side=tk.LEFT)

        self.api_expanded = tk.BooleanVar(value=False)
        self.toggle_api_btn = ttk.Button(api_header_frame, text="展开API配置 ▼",
                                          command=self.toggle_api_config, width=15)
        self.toggle_api_btn.pack(side=tk.LEFT, padx=5)

        ttk.Button(api_header_frame, text="保存配置", command=self.save_api_config, width=10).pack(side=tk.LEFT, padx=5)
        ttk.Button(api_header_frame, text="重置默认", command=self.reset_api_config, width=10).pack(side=tk.LEFT)

        # API配置详情区域（默认隐藏）
        self.api_detail_frame = ttk.LabelFrame(main_frame, text="API配置详情", padding="10")

        # 创建Notebook用于切换流式/非流式配置
        self.api_notebook = ttk.Notebook(self.api_detail_frame)
        self.api_notebook.pack(fill=tk.X, expand=True)

        # === 流式调用配置页 ===
        stream_page = ttk.Frame(self.api_notebook, padding="5")
        self.api_notebook.add(stream_page, text="流式调用配置")

        # 流式-主模型
        stream_main_frame = ttk.LabelFrame(stream_page, text="主模型", padding="5")
        stream_main_frame.pack(fill=tk.X, pady=5)
        self._create_model_config(stream_main_frame, "stream_main")

        # 流式-备用模型
        stream_backup_frame = ttk.LabelFrame(stream_page, text="备用模型（主模型失败后切换）", padding="5")
        stream_backup_frame.pack(fill=tk.X, pady=5)
        self._create_model_config(stream_backup_frame, "stream_backup")

        # === 非流式调用配置页 ===
        non_stream_page = ttk.Frame(self.api_notebook, padding="5")
        self.api_notebook.add(non_stream_page, text="非流式调用配置")

        # 非流式-主模型
        non_stream_main_frame = ttk.LabelFrame(non_stream_page, text="主模型", padding="5")
        non_stream_main_frame.pack(fill=tk.X, pady=5)
        self._create_model_config(non_stream_main_frame, "non_stream_main")

        # 非流式-备用模型
        non_stream_backup_frame = ttk.LabelFrame(non_stream_page, text="备用模型（主模型失败后切换）", padding="5")
        non_stream_backup_frame.pack(fill=tk.X, pady=5)
        self._create_model_config(non_stream_backup_frame, "non_stream_backup")

        # === 输入路径 ===
        input_frame = ttk.LabelFrame(main_frame, text="参考文案路径", padding="10")
        input_frame.pack(fill=tk.X, pady=5)

        self.input_path = tk.StringVar(value=r"d:\A百家号带货文案库\仿写文案.txt")
        ttk.Entry(input_frame, textvariable=self.input_path, width=70).pack(side=tk.LEFT, fill=tk.X, expand=True)
        ttk.Button(input_frame, text="选择文件", command=self.select_input_file, width=10).pack(side=tk.LEFT, padx=5)
        ttk.Button(input_frame, text="选择文件夹", command=self.select_input_folder, width=10).pack(side=tk.LEFT)

        # === 输出路径 ===
        output_frame = ttk.LabelFrame(main_frame, text="输出保存路径", padding="10")
        output_frame.pack(fill=tk.X, pady=5)

        self.output_path = tk.StringVar(value=r"D:\A百家号带货视频\带货文案")
        ttk.Entry(output_frame, textvariable=self.output_path, width=70).pack(side=tk.LEFT, fill=tk.X, expand=True)
        ttk.Button(output_frame, text="选择文件夹", command=self.select_output_folder, width=10).pack(side=tk.LEFT, padx=5)

        # === 引流类型选择 ===
        flow_frame = ttk.LabelFrame(main_frame, text="引流类型", padding="10")
        flow_frame.pack(fill=tk.X, pady=5)

        self.flow_type = tk.StringVar(value="置顶引流")
        for ft in self.flow_types.keys():
            ttk.Radiobutton(flow_frame, text=ft, variable=self.flow_type, value=ft,
                           command=self.on_flow_type_change).pack(side=tk.LEFT, padx=10)

        # === 引流话术区域 ===
        self.yinliu_frame = ttk.LabelFrame(main_frame, text="引流话术（可选）", padding="10")
        self.yinliu_frame.pack(fill=tk.X, pady=5)

        # 话术下拉框
        yinliu_combo_frame = ttk.Frame(self.yinliu_frame)
        yinliu_combo_frame.pack(fill=tk.X, pady=(0, 5))
        ttk.Label(yinliu_combo_frame, text="已保存话术:").pack(side=tk.LEFT)
        self.yinliu_combo = ttk.Combobox(yinliu_combo_frame, width=50, state="readonly")
        self.yinliu_combo.pack(side=tk.LEFT, padx=5)
        self.yinliu_combo.bind("<<ComboboxSelected>>", self.on_yinliu_select)

        # 话术操作按钮
        yinliu_btn_frame = ttk.Frame(self.yinliu_frame)
        yinliu_btn_frame.pack(fill=tk.X, pady=(0, 5))
        ttk.Button(yinliu_btn_frame, text="保存当前话术", command=self.save_yinliu_template, width=12).pack(side=tk.LEFT, padx=2)
        ttk.Button(yinliu_btn_frame, text="删除选中", command=self.delete_yinliu_template, width=10).pack(side=tk.LEFT, padx=2)
        ttk.Button(yinliu_btn_frame, text="清空输入", command=self.clear_yinliu_text, width=10).pack(side=tk.LEFT, padx=2)

        # 话术输入框
        self.yinliu_text = scrolledtext.ScrolledText(self.yinliu_frame, height=4, width=80)
        self.yinliu_text.pack(fill=tk.X)

        # 初始化话术下拉框
        self.update_yinliu_combo()

        # === 带货信息区域（默认隐藏） ===
        self.daihuo_frame = ttk.LabelFrame(main_frame, text="带货商品信息", padding="10")

        # 商品名称
        name_frame = ttk.Frame(self.daihuo_frame)
        name_frame.pack(fill=tk.X, pady=2)
        ttk.Label(name_frame, text="商品名称:", width=10).pack(side=tk.LEFT)
        self.product_name = tk.StringVar()
        ttk.Entry(name_frame, textvariable=self.product_name, width=50).pack(side=tk.LEFT, fill=tk.X, expand=True)

        # 产品素材
        ttk.Label(self.daihuo_frame, text="产品素材/介绍:").pack(anchor=tk.W, pady=(5, 2))
        self.product_material = scrolledtext.ScrolledText(self.daihuo_frame, height=4, width=80)
        self.product_material.pack(fill=tk.X)

        # === 操作按钮 ===
        self.btn_frame_container = ttk.Frame(main_frame)
        self.btn_frame_container.pack(fill=tk.X, pady=10)

        btn_frame = ttk.Frame(self.btn_frame_container)
        btn_frame.pack()

        self.start_btn = ttk.Button(btn_frame, text="开始生成", command=self.start_generate, width=15)
        self.start_btn.pack(side=tk.LEFT, padx=10)

        self.stop_btn = ttk.Button(btn_frame, text="停止", command=self.stop_generate, width=10, state=tk.DISABLED)
        self.stop_btn.pack(side=tk.LEFT, padx=10)

        self.regenerate_btn = ttk.Button(btn_frame, text="重新生成", command=self.regenerate, width=12, state=tk.DISABLED)
        self.regenerate_btn.pack(side=tk.LEFT, padx=10)

        ttk.Button(btn_frame, text="打开输出文件夹", command=self.open_output_folder, width=15).pack(side=tk.LEFT, padx=10)

        # === 重新生成建议 ===
        self.suggestion_frame = ttk.LabelFrame(main_frame, text="重新生成建议（可选）", padding="10")
        self.suggestion_frame.pack(fill=tk.X, pady=5)

        ttk.Label(self.suggestion_frame, text="输入你的修改建议，重新生成时会参考：").pack(anchor=tk.W)
        self.suggestion_text = scrolledtext.ScrolledText(self.suggestion_frame, height=3, width=80)
        self.suggestion_text.pack(fill=tk.X)
        self.suggestion_text.insert("1.0", "例如：开头太普通，换一个更有冲击力的；结尾引流太生硬，要更自然一些")

        # === 进度条 ===
        progress_frame = ttk.Frame(main_frame)
        progress_frame.pack(fill=tk.X, pady=5)

        self.progress_var = tk.DoubleVar(value=0)
        self.progress_bar = ttk.Progressbar(progress_frame, variable=self.progress_var, maximum=100)
        self.progress_bar.pack(fill=tk.X)

        self.status_label = ttk.Label(progress_frame, text="就绪")
        self.status_label.pack(pady=5)

        # === 日志区域 ===
        log_frame = ttk.LabelFrame(main_frame, text="运行日志", padding="10")
        log_frame.pack(fill=tk.BOTH, expand=True, pady=5)

        self.log_text = scrolledtext.ScrolledText(log_frame, height=12, width=80)
        self.log_text.pack(fill=tk.BOTH, expand=True)

        # 初始化运行状态
        self.is_running = False
        # 保存最后一次生成的参考文案信息（用于重新生成）
        self.last_articles = []
        self.last_flow_type = ""
        self.last_yinliu_content = ""
        self.last_product_name = ""
        self.last_product_material = ""

    def _create_model_config(self, parent, prefix):
        """创建模型配置UI组件"""
        # URL
        url_frame = ttk.Frame(parent)
        url_frame.pack(fill=tk.X, pady=2)
        ttk.Label(url_frame, text="URL:", width=12).pack(side=tk.LEFT)
        url_var = tk.StringVar(value=self.config.get(f"{prefix}_url", DEFAULT_CONFIG.get(f"{prefix}_url", "")))
        setattr(self, f"{prefix}_url", url_var)
        ttk.Entry(url_frame, textvariable=url_var, width=60).pack(side=tk.LEFT, fill=tk.X, expand=True)

        # Key
        key_frame = ttk.Frame(parent)
        key_frame.pack(fill=tk.X, pady=2)
        ttk.Label(key_frame, text="API Key:", width=12).pack(side=tk.LEFT)
        key_var = tk.StringVar(value=self.config.get(f"{prefix}_key", DEFAULT_CONFIG.get(f"{prefix}_key", "")))
        setattr(self, f"{prefix}_key", key_var)
        ttk.Entry(key_frame, textvariable=key_var, width=60).pack(side=tk.LEFT, fill=tk.X, expand=True)

        # Model + Max Tokens
        model_frame = ttk.Frame(parent)
        model_frame.pack(fill=tk.X, pady=2)
        ttk.Label(model_frame, text="模型:", width=12).pack(side=tk.LEFT)
        model_var = tk.StringVar(value=self.config.get(f"{prefix}_model", DEFAULT_CONFIG.get(f"{prefix}_model", "")))
        setattr(self, f"{prefix}_model", model_var)
        ttk.Entry(model_frame, textvariable=model_var, width=35).pack(side=tk.LEFT)
        ttk.Label(model_frame, text="  Max Tokens:").pack(side=tk.LEFT)
        tokens_var = tk.StringVar(value=str(self.config.get(f"{prefix}_max_tokens", DEFAULT_CONFIG.get(f"{prefix}_max_tokens", 16000))))
        setattr(self, f"{prefix}_max_tokens", tokens_var)
        ttk.Entry(model_frame, textvariable=tokens_var, width=10).pack(side=tk.LEFT)

    def toggle_api_config(self):
        """切换API配置区域的显示/隐藏"""
        if self.api_expanded.get():
            self.api_detail_frame.pack_forget()
            self.toggle_api_btn.config(text="展开API配置 ▼")
            self.api_expanded.set(False)
        else:
            # 在api_header_frame之后插入
            self.api_detail_frame.pack(fill=tk.X, pady=5, after=self.toggle_api_btn.master)
            self.toggle_api_btn.config(text="收起API配置 ▲")
            self.api_expanded.set(True)

    def on_stream_change(self):
        """流式开关改变"""
        pass  # 配置会在保存时更新

    def save_api_config(self):
        """保存API配置"""
        try:
            self.config["use_stream"] = self.use_stream.get()
            # 流式配置
            self.config["stream_main_url"] = self.stream_main_url.get().strip()
            self.config["stream_main_key"] = self.stream_main_key.get().strip()
            self.config["stream_main_model"] = self.stream_main_model.get().strip()
            self.config["stream_main_max_tokens"] = int(self.stream_main_max_tokens.get())
            self.config["stream_backup_url"] = self.stream_backup_url.get().strip()
            self.config["stream_backup_key"] = self.stream_backup_key.get().strip()
            self.config["stream_backup_model"] = self.stream_backup_model.get().strip()
            self.config["stream_backup_max_tokens"] = int(self.stream_backup_max_tokens.get())
            # 非流式配置
            self.config["non_stream_main_url"] = self.non_stream_main_url.get().strip()
            self.config["non_stream_main_key"] = self.non_stream_main_key.get().strip()
            self.config["non_stream_main_model"] = self.non_stream_main_model.get().strip()
            self.config["non_stream_main_max_tokens"] = int(self.non_stream_main_max_tokens.get())
            self.config["non_stream_backup_url"] = self.non_stream_backup_url.get().strip()
            self.config["non_stream_backup_key"] = self.non_stream_backup_key.get().strip()
            self.config["non_stream_backup_model"] = self.non_stream_backup_model.get().strip()
            self.config["non_stream_backup_max_tokens"] = int(self.non_stream_backup_max_tokens.get())
            save_config(self.config)
            messagebox.showinfo("成功", "API配置已保存")
        except Exception as e:
            messagebox.showerror("错误", f"保存配置失败：{e}")

    def reset_api_config(self):
        """重置为默认配置"""
        if messagebox.askyesno("确认", "确定要重置为默认配置吗？"):
            self.config = DEFAULT_CONFIG.copy()
            self.use_stream.set(self.config["use_stream"])
            # 流式配置
            self.stream_main_url.set(self.config["stream_main_url"])
            self.stream_main_key.set(self.config["stream_main_key"])
            self.stream_main_model.set(self.config["stream_main_model"])
            self.stream_main_max_tokens.set(str(self.config["stream_main_max_tokens"]))
            self.stream_backup_url.set(self.config["stream_backup_url"])
            self.stream_backup_key.set(self.config["stream_backup_key"])
            self.stream_backup_model.set(self.config["stream_backup_model"])
            self.stream_backup_max_tokens.set(str(self.config["stream_backup_max_tokens"]))
            # 非流式配置
            self.non_stream_main_url.set(self.config["non_stream_main_url"])
            self.non_stream_main_key.set(self.config["non_stream_main_key"])
            self.non_stream_main_model.set(self.config["non_stream_main_model"])
            self.non_stream_main_max_tokens.set(str(self.config["non_stream_main_max_tokens"]))
            self.non_stream_backup_url.set(self.config["non_stream_backup_url"])
            self.non_stream_backup_key.set(self.config["non_stream_backup_key"])
            self.non_stream_backup_model.set(self.config["non_stream_backup_model"])
            self.non_stream_backup_max_tokens.set(str(self.config["non_stream_backup_max_tokens"]))
            save_config(self.config)
            messagebox.showinfo("成功", "已重置为默认配置")

    def on_flow_type_change(self, event=None):
        """引流类型改变时的处理"""
        flow_type = self.flow_type.get()

        if flow_type == "带货引流":
            self.daihuo_frame.pack(fill=tk.X, pady=5, before=self.btn_frame_container)
        else:
            self.daihuo_frame.pack_forget()

        if flow_type == "纯夸赞不引流":
            self.yinliu_frame.pack_forget()
        else:
            self.yinliu_frame.pack(fill=tk.X, pady=5, before=self.daihuo_frame if flow_type == "带货引流" else self.btn_frame_container)

        # 切换引流类型时更新话术下拉框
        self.update_yinliu_combo()

    def update_yinliu_combo(self):
        """更新话术下拉框"""
        flow_type = self.flow_type.get()
        if flow_type == "纯夸赞不引流":
            return

        # 获取当前类型的话术列表
        templates = self.config.get("yinliu_templates", {}).get(flow_type, [])

        # 构建下拉框选项（显示前20个字）
        options = ["-- 不使用已保存话术 --"]
        for i, tpl in enumerate(templates):
            preview = tpl[:20].replace('\n', ' ') + "..." if len(tpl) > 20 else tpl.replace('\n', ' ')
            options.append(f"{i+1}. {preview}")

        self.yinliu_combo['values'] = options
        self.yinliu_combo.current(0)

    def on_yinliu_select(self, event=None):
        """选择话术时填充到文本框"""
        flow_type = self.flow_type.get()
        selection = self.yinliu_combo.current()

        if selection <= 0:  # 选择了"不使用"
            return

        # 获取对应的话术内容
        templates = self.config.get("yinliu_templates", {}).get(flow_type, [])
        if selection - 1 < len(templates):
            template = templates[selection - 1]
            # 填充到文本框
            self.yinliu_text.delete("1.0", tk.END)
            self.yinliu_text.insert("1.0", template)

    def save_yinliu_template(self):
        """保存当前话术到列表"""
        flow_type = self.flow_type.get()
        if flow_type == "纯夸赞不引流":
            messagebox.showwarning("提示", "纯夸赞模式不需要引流话术")
            return

        content = self.yinliu_text.get("1.0", tk.END).strip()
        if not content:
            messagebox.showwarning("提示", "请先输入话术内容")
            return

        # 确保配置中有话术库
        if "yinliu_templates" not in self.config:
            self.config["yinliu_templates"] = {"置顶引流": [], "橱窗引流": [], "带货引流": []}
        if flow_type not in self.config["yinliu_templates"]:
            self.config["yinliu_templates"][flow_type] = []

        # 检查是否已存在
        templates = self.config["yinliu_templates"][flow_type]
        if content in templates:
            messagebox.showinfo("提示", "该话术已存在")
            return

        # 添加到列表
        templates.append(content)
        save_config(self.config)

        # 更新下拉框
        self.update_yinliu_combo()
        messagebox.showinfo("成功", f"话术已保存到【{flow_type}】列表")

    def delete_yinliu_template(self):
        """删除选中的话术"""
        flow_type = self.flow_type.get()
        selection = self.yinliu_combo.current()

        if selection <= 0:
            messagebox.showwarning("提示", "请先选择要删除的话术")
            return

        templates = self.config.get("yinliu_templates", {}).get(flow_type, [])
        if selection - 1 < len(templates):
            if messagebox.askyesno("确认", "确定要删除选中的话术吗？"):
                del templates[selection - 1]
                save_config(self.config)
                self.update_yinliu_combo()
                self.yinliu_text.delete("1.0", tk.END)
                messagebox.showinfo("成功", "话术已删除")

    def clear_yinliu_text(self):
        """清空话术输入框"""
        self.yinliu_text.delete("1.0", tk.END)
        self.yinliu_combo.current(0)

    def select_input_file(self):
        file_path = filedialog.askopenfilename(
            title="选择参考文案文件",
            filetypes=[("文本文件", "*.txt"), ("所有文件", "*.*")]
        )
        if file_path:
            self.input_path.set(file_path)

    def select_input_folder(self):
        folder_path = filedialog.askdirectory(title="选择参考文案文件夹")
        if folder_path:
            self.input_path.set(folder_path)

    def select_output_folder(self):
        folder_path = filedialog.askdirectory(title="选择输出保存文件夹")
        if folder_path:
            self.output_path.set(folder_path)

    def log(self, message):
        """添加日志"""
        timestamp = datetime.now().strftime("%H:%M:%S")
        log_line = f"[{timestamp}] {message}\n"
        self.log_text.insert(tk.END, log_line)
        self.log_text.see(tk.END)
        self.root.update()

        try:
            log_file = os.path.join(self.output_path.get() or ".", "debug_log.txt")
            with open(log_file, 'a', encoding='utf-8') as f:
                f.write(log_line)
        except:
            pass

    def update_status(self, message):
        """更新状态"""
        self.status_label.config(text=message)
        self.root.update()

    def update_progress(self, value):
        """更新进度条"""
        self.progress_var.set(value)
        self.root.update()

    def start_generate(self):
        """开始生成"""
        input_path = self.input_path.get()
        output_path = self.output_path.get()

        if not input_path:
            messagebox.showerror("错误", "请选择参考文案路径")
            return

        if not os.path.exists(input_path):
            messagebox.showerror("错误", f"参考文案路径不存在：{input_path}")
            return

        if not output_path:
            messagebox.showerror("错误", "请选择输出保存路径")
            return

        if self.flow_type.get() == "带货引流":
            if not self.product_name.get().strip():
                messagebox.showerror("错误", "请填写带货商品名称")
                return
            if not self.product_material.get("1.0", tk.END).strip():
                messagebox.showerror("错误", "请填写产品素材/介绍")
                return

        os.makedirs(output_path, exist_ok=True)

        self.start_btn.config(state=tk.DISABLED)
        self.stop_btn.config(state=tk.NORMAL)
        self.is_running = True

        self.log_text.delete(1.0, tk.END)

        thread = threading.Thread(target=self.generate_task)
        thread.daemon = True
        thread.start()

    def stop_generate(self):
        """停止生成"""
        self.is_running = False
        self.log("用户停止了生成任务")
        self.update_status("已停止")
        self.start_btn.config(state=tk.NORMAL)
        self.stop_btn.config(state=tk.DISABLED)

    def regenerate(self):
        """重新生成 - 使用上次的参考文案和当前设置"""
        if not self.last_articles:
            messagebox.showwarning("提示", "没有可重新生成的文案，请先执行一次生成")
            return

        output_path = self.output_path.get()
        if not output_path:
            messagebox.showerror("错误", "请选择输出保存路径")
            return

        os.makedirs(output_path, exist_ok=True)

        self.start_btn.config(state=tk.DISABLED)
        self.stop_btn.config(state=tk.NORMAL)
        self.regenerate_btn.config(state=tk.DISABLED)
        self.is_running = True

        self.log_text.delete(1.0, tk.END)
        self.log("开始重新生成...")

        # 获取用户建议
        suggestion = self.suggestion_text.get("1.0", tk.END).strip()
        # 过滤掉默认提示文字
        if suggestion.startswith("例如："):
            suggestion = ""

        thread = threading.Thread(target=self.regenerate_task, args=(suggestion,))
        thread.daemon = True
        thread.start()

    def regenerate_task(self, user_suggestion):
        """重新生成任务"""
        try:
            output_path = self.output_path.get()

            if user_suggestion:
                self.log(f"用户建议：{user_suggestion}")

            total = len(self.last_articles)
            for art_idx, article in enumerate(self.last_articles):
                if not self.is_running:
                    break

                self.log(f"\n--- 重新生成第 {art_idx+1} 篇 ---")
                self.update_status(f"正在重新生成第 {art_idx+1} 篇...")
                self.update_progress((art_idx / total) * 100)

                # 生成仿写文案（带用户建议）
                result = self.generate_document(
                    article, self.last_flow_type, self.last_yinliu_content,
                    self.last_product_name, self.last_product_material,
                    user_suggestion
                )

                if result:
                    self.save_document(result, output_path, art_idx + 1)
                else:
                    self.log(f"第 {art_idx+1} 篇重新生成失败")

            self.update_progress(100)
            self.log("\n" + "="*50)
            self.log("重新生成完成！")
            self.finish_task()

        except Exception as e:
            self.log(f"重新生成出错: {str(e)}")
            import traceback
            self.log(traceback.format_exc())
            self.finish_task()

    def generate_task(self):
        """生成任务主函数"""
        try:
            input_path = self.input_path.get()
            output_path = self.output_path.get()
            flow_type = self.flow_type.get()

            # 读取参考文案
            self.log("正在读取参考文案...")
            if os.path.isfile(input_path):
                with open(input_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                files_content = [(os.path.basename(input_path), content)]
            else:
                files_content = []
                for fname in os.listdir(input_path):
                    if fname.endswith('.txt'):
                        fpath = os.path.join(input_path, fname)
                        with open(fpath, 'r', encoding='utf-8') as f:
                            files_content.append((fname, f.read()))

            if not files_content:
                self.log("错误：没有找到任何文案文件")
                self.finish_task()
                return

            # 读取引流素材
            yinliu_content = self.yinliu_text.get("1.0", tk.END).strip()

            # 带货信息
            product_name = ""
            product_material = ""
            if flow_type == "带货引流":
                product_name = self.product_name.get().strip()
                product_material = self.product_material.get("1.0", tk.END).strip()

            total_files = len(files_content)
            for idx, (fname, content) in enumerate(files_content):
                if not self.is_running:
                    break

                self.log(f"\n{'='*50}")
                self.log(f"处理文件 [{idx+1}/{total_files}]: {fname}")
                self.update_progress((idx / total_files) * 100)

                # 分割参考文案
                articles = self.parse_articles(content)
                self.log(f"识别到 {len(articles)} 篇参考文案")

                # 保存最后一次生成的信息（用于重新生成）
                self.last_articles = articles
                self.last_flow_type = flow_type
                self.last_yinliu_content = yinliu_content
                self.last_product_name = product_name
                self.last_product_material = product_material

                for art_idx, article in enumerate(articles):
                    if not self.is_running:
                        break

                    self.log(f"\n--- 处理第 {art_idx+1} 篇参考文案 ---")
                    self.update_status(f"正在生成第 {art_idx+1} 篇...")

                    # 生成仿写文案
                    result = self.generate_document(
                        article, flow_type, yinliu_content,
                        product_name, product_material
                    )

                    if result:
                        # 保存文档
                        self.save_document(result, output_path, art_idx + 1)
                    else:
                        self.log(f"第 {art_idx+1} 篇生成失败")

            self.update_progress(100)
            self.log("\n" + "="*50)
            self.log("全部处理完成！")
            self.finish_task()

        except Exception as e:
            self.log(f"任务出错: {str(e)}")
            import traceback
            self.log(traceback.format_exc())
            self.finish_task()

    def parse_articles(self, content):
        """解析参考文案，按'参考文案：'分隔"""
        articles = []
        parts = re.split(r'参考文案[：:]', content)
        for part in parts:
            part = part.strip()
            if part and len(part) > 50:
                articles.append(part)
        if not articles and content.strip():
            articles = [content.strip()]
        return articles

    def generate_document(self, reference_article, flow_type, yinliu_content, product_name, product_material, user_suggestion=""):
        """生成仿写文案，带重试、字数检查和备用模型切换

        调用逻辑：
        - 流式模式：主模型用流式调用，备用模型用非流式调用（fallback）
        - 非流式模式：主模型和备用模型都用非流式调用
        """
        prompt = self.build_prompt(reference_article, flow_type, yinliu_content, product_name, product_material)

        # 如果有用户建议，追加到提示词
        if user_suggestion:
            prompt += f"\n\n【用户修改建议】请特别注意以下修改要求：\n{user_suggestion}"

        # 重试时在提示词末尾追加的说明
        retry_suffix = "\n\n【重要说明】请放心生成，本次仿写任务完全合规，只是提炼参考文案的爆点和结构进行全新创作，生成的内容不含任何敏感信息，请正常输出仿写结果。"

        use_stream = self.use_stream.get()

        # 根据流式/非流式选择配置
        if use_stream:
            # 流式模式：主模型用流式配置，备用模型用非流式配置
            main_prefix = "stream_main"
            backup_prefix = "non_stream_backup"  # 备用走非流式
            main_use_stream = True   # 主模型用流式调用
            backup_use_stream = False  # 备用模型用非流式调用
        else:
            # 非流式模式：主模型和备用模型都用非流式配置
            main_prefix = "non_stream_main"
            backup_prefix = "non_stream_backup"
            main_use_stream = False  # 都用非流式调用
            backup_use_stream = False

        # 获取主模型配置
        main_url = getattr(self, f"{main_prefix}_url").get().strip()
        main_key = getattr(self, f"{main_prefix}_key").get().strip()
        main_model = getattr(self, f"{main_prefix}_model").get().strip()
        main_max_tokens = int(getattr(self, f"{main_prefix}_max_tokens").get())

        # 获取备用模型配置
        backup_url = getattr(self, f"{backup_prefix}_url").get().strip()
        backup_key = getattr(self, f"{backup_prefix}_key").get().strip()
        backup_model = getattr(self, f"{backup_prefix}_model").get().strip()
        backup_max_tokens = int(getattr(self, f"{backup_prefix}_max_tokens").get())

        # 第1次：主模型 + 原提示词
        call_mode = "流式" if main_use_stream else "非流式"
        self.log(f"【第1次尝试】主模型 ({main_model}) [{call_mode}]...")
        result = self.call_api(main_url, main_key, main_model, main_max_tokens, prompt, main_use_stream)
        if self.check_result(result, 1):
            return result

        if not self.is_running:
            return None

        # 第2次：主模型 + 原提示词 + 追加说明
        self.log(f"【第2次尝试】主模型 + 追加说明 [{call_mode}]...")
        result = self.call_api(main_url, main_key, main_model, main_max_tokens, prompt + retry_suffix, main_use_stream)
        if self.check_result(result, 2):
            return result

        if not self.is_running:
            return None

        # 第3次：备用模型 + 原提示词 + 追加说明
        backup_mode = "流式" if backup_use_stream else "非流式"
        self.log(f"【第3次尝试】备用模型 ({backup_model}) [{backup_mode}]...")
        result = self.call_api(backup_url, backup_key, backup_model, backup_max_tokens, prompt + retry_suffix, backup_use_stream)
        if self.check_result(result, 3):
            return result

        if not self.is_running:
            return None

        # 第4次：备用模型 + 原提示词 + 追加说明
        self.log(f"【第4次尝试】备用模型（最后一次）[{backup_mode}]...")
        result = self.call_api(backup_url, backup_key, backup_model, backup_max_tokens, prompt + retry_suffix, backup_use_stream)
        if self.check_result(result, 4):
            return result

        self.log("4次尝试均失败")
        return None

    def check_result(self, result, attempt_num):
        """检查结果：有效性 + 字数"""
        if not result or not self.is_valid_result(result):
            self.log(f"第{attempt_num}次：API返回无效或被拒绝")
            return False

        char_count = self.count_chinese_chars(result)
        self.log(f"第{attempt_num}次：生成字数 {char_count}")

        if char_count < 1500:
            self.log(f"字数不足1500，需要重试...")
            return False
        else:
            self.log(f"字数合格（{char_count}字），生成成功！")
            return True

    def count_chinese_chars(self, text):
        """统计中文字符数量"""
        if not text:
            return 0
        chinese_chars = re.findall(r'[\u4e00-\u9fff]', text)
        return len(chinese_chars)

    def is_valid_result(self, result):
        """检查结果是否有效"""
        if not result:
            return False
        # 检查是否是拒绝回答
        reject_keywords = ["抱歉", "无法", "不能", "拒绝", "违反", "政策", "sorry", "cannot", "can't"]
        result_lower = result.lower()
        for kw in reject_keywords:
            if kw in result_lower and len(result) < 500:
                return False
        return len(result) > 300

    def call_api(self, base_url, api_key, model, max_tokens, prompt, use_stream):
        """调用API（流式或非流式）"""
        if use_stream:
            return self.call_llm_stream(base_url, api_key, model, max_tokens, prompt)
        else:
            return self.call_llm_non_stream(base_url, api_key, model, max_tokens, prompt)

    def call_llm_stream(self, base_url, api_key, model, max_tokens, prompt):
        """流式API调用"""
        url = f"{base_url.rstrip('/')}/chat/completions"
        headers = {
            "Content-Type": "application/json",
            "Authorization": f"Bearer {api_key}"
        }
        data = {
            "model": model,
            "messages": [
                {"role": "system", "content": "你是一个专业的文案写作专家，擅长仿写百家号引流文案。"},
                {"role": "user", "content": prompt}
            ],
            "max_tokens": max_tokens,
            "stream": True
        }
        if "thinking" not in model.lower():
            data["temperature"] = 0.7

        self.log(f"[流式API] 请求: {url}, 模型: {model}")

        try:
            response = requests.post(url, headers=headers, json=data, timeout=300, stream=True)

            if response.status_code != 200:
                error_text = response.text[:200] if response.text else "空响应"
                self.log(f"[流式API] 失败: HTTP {response.status_code}: {error_text}")
                return None

            full_content = ""
            for line in response.iter_lines():
                if not self.is_running:
                    return None
                if line:
                    line_text = line.decode('utf-8')
                    if line_text.startswith('data: '):
                        json_str = line_text[6:]
                        if json_str.strip() == '[DONE]':
                            break
                        try:
                            chunk = json.loads(json_str)
                            if 'choices' in chunk and len(chunk['choices']) > 0:
                                delta = chunk['choices'][0].get('delta', )
                                content = delta.get('content', '')
                                if content:
                                    full_content += content
                        except json.JSONDecodeError:
                            continue

            if full_content:
                self.log(f"[流式API] 成功获取 {len(full_content)} 字符")
                return full_content
            else:
                self.log("[流式API] 返回空内容")
                return None

        except Exception as e:
            self.log(f"[流式API] 异常: {type(e).__name__}: {e}")
            return None

    def call_llm_non_stream(self, base_url, api_key, model, max_tokens, prompt):
        """非流式API调用"""
        url = f"{base_url.rstrip('/')}/chat/completions"
        headers = {
            "Content-Type": "application/json",
            "Authorization": f"Bearer {api_key}"
        }
        data = {
            "model": model,
            "messages": [
                {"role": "system", "content": "你是一个专业的文案写作专家，擅长仿写百家号引流文案。"},
                {"role": "user", "content": prompt}
            ],
            "max_tokens": max_tokens
        }
        if "thinking" not in model.lower():
            data["temperature"] = 0.7

        self.log(f"[非流式API] 请求: {url}, 模型: {model}")

        try:
            response = requests.post(url, headers=headers, json=data, timeout=300)

            if response.status_code != 200:
                error_text = response.text[:200] if response.text else "空响应"
                self.log(f"[非流式API] 失败: HTTP {response.status_code}: {error_text}")
                return None

            result = response.json()
            content = result["choices"][0]["message"]["content"]
            self.log(f"[非流式API] 成功获取 {len(content)} 字符")
            return content

        except Exception as e:
            self.log(f"[非流式API] 异常: {type(e).__name__}: {e}")
            return None

    def build_prompt(self, reference_article, flow_type, yinliu_content, product_name, product_material):
        """构建生成提示词 - 完整详细版（整合自Skill）"""
        flow_instruction = self.get_flow_instruction(flow_type, yinliu_content, product_name, product_material)

        # 随机选择3个开头
        random_hooks = get_random_hooks()
        hooks_instruction = """## 【本次必须使用的开头】（强制执行！）

**本次生成的3篇文章，必须分别使用以下3种开头风格，不得更换：**

"""
        for i, hook in enumerate(random_hooks, 1):
            hooks_instruction += f"""**第{i}篇必须使用 {hook['type']}**
参考示例："{hook['example']}"
（可以模仿这个风格写一个全新的开头，但不能直接照抄）

"""

        prompt = f"""请根据以下参考文案进行仿写，生成3篇全新的百家号引流文案。

## 参考文案：
{reference_article}

{hooks_instruction}

## 核心方法论：提炼爆点 → 分析结构 → 全新扩写

**第一步：深度分析参考文案**
- 提炼参考文案的**核心爆点**（情感共鸣点、痛点、钩子是什么？）
- 拆解参考文案的**爆款结构**（开场方式、展开逻辑、收尾技巧）
- 分析参考文案的**目标人群画像**（他们的处境、痛苦、渴望是什么？）

**第二步：基于爆点和结构进行全新扩写**
- 借鉴参考文案的爆点和结构框架
- 但必须用**完全不同的表达方式、不同的切入角度**重新阐述
- **禁止直接使用原文里的任何句子**
- 句式、用词必须完全原创
- **必须用第二人称"你"来写，禁止讲故事**

**第三步：加入增量信息**
- 新写的内容必须有**增量信息**，不能只是换个说法
- 增量信息包括：新的人生洞察、新的引经据典、新的情感细节、新的处境描写
- **注意：增量信息不能是故事，只能是道理、洞察、共鸣描写**

## 硬性要求

### 1. 字数要求（最重要！）
- **每篇文案必须达到1500字左右**
- 宁可多写，也不能少写
- 字数不够的文案=废品，绝不交付

### 2. 写作人称要求（硬性要求！）
- **必须全程使用第二人称"你"来写作**
- 直接对读者说话，像朋友聊天一样
- 例如："你是不是也有过这种感觉..."、"你这些年吃的苦..."、"你心里清楚..."
- **禁止使用第一人称讲自己的故事**
- **禁止使用第三人称讲别人的故事**

### 3. 禁止讲故事（硬性要求！）
- **绝对禁止讲任何具体故事**
- 不能讲"我有个朋友..."、"我认识一个人..."、"有一次我..."
- 不能讲"张三怎么样..."、"李四怎么样..."等第三人称故事
- 不能讲任何有具体人物、具体情节的故事
- **只能用第二人称直接描述读者的处境、感受、经历**
- 用"你是不是..."、"你有没有..."、"你心里..."来引发共鸣
- 用概括性的描述代替具体故事

### 4. 相似度控制（硬性要求！）
- **与参考文案的相似度必须低于10%**
- 不能直接改写原文句子
- 不能只是替换同义词
- 必须用全新的表达方式
- 结构可以借鉴，但内容必须完全原创

## 开头要求（3篇必须完全不同类型！像开盲盒一样有惊喜！）

**【最重要】强制多样化机制：**
- **每次生成必须使用全新的开头，绝对禁止重复**
- **3篇文章的开头必须来自3个不同的大类**
- **同一个开头句式只能用一次**
- **必须包含至少1篇轻松/温暖/有趣风格的开头**
- **禁止3篇全是沉重压抑的风格**

**永久禁止的开头（已经用烂了！）：**
- ❌ "我劝你别太善良"
- ❌ "有一种人，你越对他好，他越瞧不起你"
- ❌ "凭什么受伤的总是老实人？"
- ❌ "四十岁之后，我才明白一个道理"
- ❌ "那天饭局上，有人说了一句话"
- ❌ "你累了"、"深夜睡不着"、"夜深人静"
- ❌ "我知道你是什么样的人"
- ❌ "我问你一个问题"
- ❌ "你有没有发现"、"你有没有想过"
- ❌ "人这辈子"、"人啊"开头的句式
- ❌ 任何以"有一种..."开头的句式

**开头钩子库（150+种，3篇必须从不同类型中选择，必须有1篇轻松风格）：**

**A类-轻松幽默型（必选！3篇中至少用1个）：**
- "我发现一个规律：越是老实人，越容易被安排加班。你说气不气？"
- "昨天算了一笔账，这些年帮别人花的时间，够我学会三门外语了。"
- "我妈说我最大的优点是善良，我爸说这也是我最大的缺点。好家伙，亲爹。"
- "朋友说我是'便利店型人格'——24小时营业，随叫随到，还不涨价。"
- "有人说我脾气好，我笑了笑没说话。其实不是脾气好，是懒得计较。"
- "我终于明白，为什么'好人卡'发得最多——因为好人最好打发。"
- "同事问我：'你怎么从来不生气？'我说：'生气要花力气，我选择省着点用。'"
- "我的人生信条曾经是'吃亏是福'，直到我发现福没来，亏倒是吃了不少。"
- "别人都在研究怎么赚钱，我在研究怎么不被人当免费劳动力。"
- "我这人有个毛病，别人一说'就你能帮我'，我就跟中了蛊似的。"

**B类-温暖治愈型（推荐！让读者感到被理解）：**
- "嘿，今天想跟你聊点轻松的，关于那些默默付出却不求回报的人。"
- "你知道吗，这世上有一种人，他们的好，是藏在细节里的。"
- "我一直觉得，善良的人身上有光，只是有时候这光被辜负了。"
- "如果你正在看这段话，我想告诉你：你的好，有人看得见。"
- "今天不讲大道理，就想跟你说说心里话。"
- "有些人，值得被这个世界温柔以待，比如正在看这段话的你。"
- "我见过很多人，但像你这样的，真的不多。"
- "你有没有被人夸过'你人真好'？今天我想认真聊聊这件事。"
- "这段话，送给每一个在生活里默默扛着的人。"
- "我相信，看到这里的你，一定是个心里有温度的人。"

**C类-反转惊喜型（有趣！先抑后扬）：**
- "我曾经以为自己是个'老好人'，后来发现，我是个'聪明的好人'。"
- "都说老实人吃亏，但我认识一个老实人，现在过得比谁都好。"
- "你以为善良是软弱？不，善良是一种选择，而且是强者的选择。"
- "有人说心软的人没出息，我偏不信这个邪。"
- "我见过最厉害的人，恰恰是最善良的那个。"
- "别人都说我太老实会吃亏，结果呢？我还真没亏。"
- "都说好人没好报，但我今天要讲一个好人有好报的故事。"
- "我以前觉得'人善被人欺'是真理，直到我遇见了一个人。"
- "谁说善良的人就要受委屈？我第一个不服。"
- "老实人的春天，其实一直都在，只是很多人没发现。"

**D类-对话引入型：**
- "他说完这句话，我愣在原地半天没回过神：'你就是太好说话了。'"
- "有人问我：'你这辈子最后悔的事是什么？'我想了想，说了两个字。"
- "算了，不争了——你是不是也经常这样跟自己说？"
- "我爷爷临终前拉着我的手说：'记住，吃亏的人，老天爷都记着账呢。'"
- "'你怎么这么傻？'这句话，你听过多少次了？"
- "我妈常说一句话：'人善被人欺，马善被人骑。'我以前不信，现在信了。"
- "'别太老实了，会吃亏的。'说这话的人，后来怎么样了？"
- "朋友跟我说：'你知道你最大的问题是什么吗？就是太把别人当回事。'"
- "'谢谢你'——这三个字，你等了多久才听到？"
- "有人当面问我：'你是不是傻？人家都欺负到头上了你还忍？'"

**E类-生活观察型（接地气！）：**
- "菜市场大妈的一句话，让我愣了半天：'姑娘，你这么好说话，不怕被人欺负啊？'"
- "堵在路上，收音机里突然传来一句话，我眼眶一下就红了。"
- "排队的时候，前面两个人的对话让我心里一惊。"
- "超市结账时，前面那个人的一个举动，让我看清了人性。"
- "参加同学聚会，有个人的变化让我震惊了。"
- "小区门口，两个大妈聊天，聊着聊着说出了一个真相。"
- "去参加婚礼，新郎的一句话让全场安静了。"
- "年夜饭桌上，我爸突然放下筷子，说了一句话。"
- "早上买早餐，老板娘的一句话让我想了一整天。"
- "坐出租车，司机师傅跟我聊了一路，最后一句话让我沉默了。"

**F类-数字锚定型：**
- "认识老王二十年了，他教会我一件事：别对谁都掏心掏肺。"
- "被人欺负了三年，我终于想通了一个道理。"
- "用了十年时间，我才学会一个字：不。"
- "三次被人背叛之后，我悟了。"
- "五十岁之后，我才明白什么叫'人走茶凉'。"
- "帮了他七年，他一句谢谢都没说过。"
- "借出去的三万块，要了五年都没要回来。"
- "在这个单位干了八年，我终于明白了一个道理。"

**G类-金句破题型：**
- "最傻的事，就是跟烂人讲道理。"
- "你知道什么人最可怕吗？不是坏人，是那些笑着捅你刀子的人。"
- "千万别做老好人，我吃过这个亏，现在告诉你。"
- "你越忍让，别人越得寸进尺，这是我用十年换来的教训。"
- "老实人不是没脾气，是把脾气都咽进了肚子里。"
- "你的善良，要带点锋芒。"
- "不是所有的忍让都叫大度，有时候那叫窝囊。"
- "这世上最傻的事，就是把真心给了不值得的人。"
- "有些人，你帮他一百次，他记不住；你拒绝他一次，他记你一辈子。"

**H类-悬念钩子型：**
- "有件事我憋了很久，今天必须说出来。"
- "你可能不信，但接下来我说的都是真事。"
- "我要告诉你一个很多人不愿意承认的真相。"
- "接下来这段话，可能会让你不舒服，但我还是要说。"
- "有个规律，我观察了很多年才看透。"
- "今天说的这些话，可能会得罪人，但我不在乎。"
- "有些话，我本来不想说，但看到你，我忍不住了。"
- "接下来的话，你可能不爱听，但句句都是真的。"
- "有件事，我一直没跟任何人说过，今天破例。"

**I类-共鸣代入型：**
- "那种被人当众下面子的感觉，我太懂了。"
- "被最信任的人捅刀子，那种滋味，经历过的人都懂。"
- "明明没做错什么，却总被人针对，这种事我也遇到过。"
- "有些委屈，说出来都没人信。"
- "你是不是也有过这种感觉：付出最多的人，往往最不被珍惜。"
- "那种心寒的感觉，我懂。就像一盆冷水从头浇到脚。"
- "有一种苦，叫做'打碎了牙往肚子里咽'。"
- "那种被人利用完就扔掉的感觉，我经历过。"

**J类-人物故事型：**
- "我有个朋友，前两天跟我说了一件事，我听完沉默了很久。"
- "我爸这辈子只教过我一个道理，我到现在才真正理解。"
- "单位有个人，大家都不待见他，后来我才知道原因。"
- "我们小区有个大爷，天天在楼下坐着，有一天他跟我说了一番话。"
- "我表姐的经历，让我彻底看清了人心。"
- "我有个同学，当年是班里最老实的人，你猜他现在怎么样了？"
- "我舅舅年轻时吃过一个大亏，他把这个教训告诉了我。"
- "我认识一个人，他的经历让我相信：好人终有好报。"

**K类-自我剖析型：**
- "说出来不怕你笑话，我以前也是个傻子。"
- "回头看看这些年，我最后悔的一件事是太心软。"
- "如果能重来，我绝对不会再做老好人。"
- "我吃过的亏，今天全告诉你，希望你别再走我的老路。"
- "我这辈子最大的毛病，就是太把别人当回事。"
- "我曾经也是个'老好人'，后来我学聪明了。"
- "我年轻时犯过一个错，现在想起来还后悔。"
- "我以前总觉得吃亏是福，现在不这么想了。"

**L类-转折反差型：**
- "以前我不信这个道理，直到自己栽了跟头。"
- "年轻的时候觉得这话是废话，现在才知道是真理。"
- "曾经有人跟我说过一句话，我没当回事，后来我后悔了。"
- "我一直以为自己做得对，直到那件事发生。"
- "以前别人说我太老实，我还不服气，现在我服了。"
- "我曾经以为善良是优点，后来才知道，善良过头就是缺点。"
- "年轻时我不懂，现在我懂了，可惜晚了。"
- "我以前总是忍，以为忍一忍就过去了，结果呢？"

## 引经据典要求（3篇必须从不同类别中选择！）

**禁止高频使用的引用（太常见了，用了就重写！）：**
- ❌ 《道德经》"上善若水"
- ❌ 《增广贤文》"谁人背后无人说"
- ❌ 曾国藩的任何名言（用太多了）
- ❌ 杨绛"人生最曼妙的风景"
- ❌ 王阳明"此心光明"
- ❌ 莫言的任何名言（用太多了）
- ❌ 螃蟹效应、破窗效应（用太多了）

**引用类型库（10大类，每篇必须从不同类别中选择！）：**

**类型1-儒家经典：**
- 《论语》："君子坦荡荡，小人长戚戚"
- 《论语》："己所不欲，勿施于人"
- 《孟子》："穷则独善其身，达则兼济天下"
- 《孟子》："生于忧患，死于安乐"
- 《大学》："知止而后有定，定而后能静"
- 《中庸》："君子素其位而行，不愿乎其外"

**类型2-道家智慧：**
- 《道德经》："知人者智，自知者明"
- 《道德经》："大直若屈，大巧若拙"
- 《道德经》："祸兮福之所倚，福兮祸之所伏"
- 《庄子》："井蛙不可以语于海者，拘于虚也"
- 《庄子》："相濡以沫，不如相忘于江湖"
- 《列子》："天地无全功，圣人无全能"

**类型3-处世格言：**
- 《菜根谭》："宠辱不惊，闲看庭前花开花落"
- 《菜根谭》："路径窄处，留一步与人行"
- 《围炉夜话》："十分不耐烦，乃为人大病"
- 《小窗幽记》："宠辱不惊，看庭前花开花落"
- 《幽梦影》："少年人须有老成之识见，老成人须有少年之襟怀"
- 《呻吟语》："轻信轻发，听言之大戒也"

**类型4-史书智慧：**
- 《史记》："桃李不言，下自成蹊"
- 《史记》："燕雀安知鸿鹄之志"
- 《资治通鉴》："兼听则明，偏信则暗"
- 《战国策》："狡兔三窟，仅得免其死耳"
- 《左传》："多行不义必自毙"
- 《汉书》："水至清则无鱼，人至察则无徒"

**类型5-诗词名句：**
- 苏轼："竹杖芒鞋轻胜马，谁怕？一蓑烟雨任平生"
- 苏轼："人有悲欢离合，月有阴晴圆缺"
- 辛弃疾："众里寻他千百度，蓦然回首，那人却在灯火阑珊处"
- 陆游："山重水复疑无路，柳暗花明又一村"
- 李白："长风破浪会有时，直挂云帆济沧海"
- 杜甫："会当凌绝顶，一览众山小"
- 王维："行到水穷处，坐看云起时"
- 白居易："试玉要烧三日满，辨材须待七年期"

**类型6-明清名人：**
- 王阳明《传习录》："破山中贼易，破心中贼难"
- 王阳明："你未看此花时，此花与汝同归于寂"
- 张居正："谋之在众，断之在独"
- 于谦："粉身碎骨浑不怕，要留清白在人间"
- 左宗棠："好便宜者，不可与之交财"
- 林则徐："海纳百川，有容乃大"
- 纪晓岚："事能知足心常惬，人到无求品自高"

**类型7-民国大家：**
- 鲁迅："真的猛士，敢于直面惨淡的人生"
- 鲁迅："不在沉默中爆发，就在沉默中灭亡"
- 胡适："做学问要在不疑处有疑，待人要在有疑处不疑"
- 梁启超："患难困苦，是磨炼人格之最高学校"
- 林语堂："人生不过如此，且行且珍惜"
- 丰子恺："不乱于心，不困于情，不畏将来，不念过往"

**类型8-当代作家：**
- 杨绛《走到人生边上》："一个人经过不同程度的锻炼，就获得不同程度的修养"
- 季羡林："人生在世，不如意事十之八九，常想一二，不思八九"
- 余华《活着》："人是为活着本身而活着，而不是为了活着之外的任何事物所活着"
- 史铁生："命定的局限尽可永在，不屈的挑战却不可须臾或缺"
- 路遥《平凡的世界》："生活不能等待别人来安排，要自己去争取和奋斗"
- 汪曾祺："人间送小温，不要太多，不要太少"
- 贾平凹："人的一生，苦也罢，乐也罢，最重要的是心间的一泓清泉里不能没有月辉"
- 三毛："一个人至少拥有一个梦想，有一个理由去坚强"

**类型9-当代名人金句：**
- 白岩松："一个人的价值，不是看他拥有多少，而是看他能给予多少"
- 罗翔："一个知识越贫乏的人，越是拥有一种莫名奇怪的勇气和自豪感"
- 董宇辉："人生就是一场修行，修的是一颗心"
- 任正非："烧不死的鸟就是凤凰"
- 曹德旺："人这一辈子，吃亏就是福"
- 褚时健："人生总有起落，精神终可传承"
- 稻盛和夫："人生的意义在于提升心性，磨炼灵魂"

**类型10-心理学/哲学概念：**
- 马斯洛需求层次："人在满足基本需求后，会追求更高层次的自我实现"
- 沉没成本效应："已经付出的成本不应该影响未来的决策"
- 幸存者偏差："我们只看到成功者，却忽略了更多失败者"
- 登门槛效应："先接受小要求，更容易接受大要求"
- 鸟笼效应："人们会为了一个不需要的东西，再添置更多不需要的东西"
- 刺猬法则："保持适当距离，才能既温暖又不伤害"
- 蘑菇定律："新人都要经历一段被忽视的时期，熬过去就好了"
- 飞轮效应："万事开头难，但一旦转动起来，就会越来越轻松"
- 延迟满足："能够忍耐当下的诱惑，才能获得更大的回报"
- 峰终定律："人们对一段经历的评价，取决于高峰和结尾时的感受"

**引用规则（必须严格执行！）：**
- 3篇文章必须从3个不同的大类中选择引用
- 每篇文章至少引用2处，且来自不同来源
- 禁止连续两次生成使用相同的引用
- 引用要自然融入文章，不能生硬堆砌
- 引用后要有自己的解读和延伸，不能只是引用完就结束

## 内容结构（确保字数达标1500字左右）

1. **开头（约400字）**：独特钩子 + 第二人称描写读者处境 + 情感共鸣
2. **处境展开（约550字）**：用"你"描写读者的经历、感受、委屈（禁止讲具体故事）
3. **道理阐述（约400字）**：引经据典 + 深度分析 + 人生洞察
4. **收尾引流（约250字）**：总结升华 + 给出希望 + 引流话术

## 标题要求（15个标题必须完全不同！）

**标题核心原则（最重要！）：**
- **标题必须用第二人称"你"**：直接对读者说话，让读者觉得说的就是自己
- **标题要夸赞/肯定读者**：认同、理解、赞美读者
- **标题要短**：控制在10-20字
- **标题要正面积极**：给读者温暖和力量
- **禁止用悬念、反问、疑问句**：不要制造悬念，不要反问读者
- **15个标题必须完全不同**：不能有任何重复、相似、雷同的表达

**标题类型库（50种，每篇5个标题从不同类型选）：**

**【夸赞品质类】**
1. "你这种人，活该被人尊重"
2. "像你这样的人，太难得了"
3. "你身上有种光，藏不住的"
4. "你的格局，一般人比不了"
5. "你这样的人，走到哪都发光"

**【认同付出类】**
6. "你的善良，终会被看见"
7. "你的付出，老天都记着"
8. "你做的一切，都不会白费"
9. "你吃过的苦，都算数"
10. "你扛过的事，都是资本"

**【肯定能力类】**
11. "能扛事的人，都不简单"
12. "你比自己想象的更厉害"
13. "能熬过来的人，都有大本事"
14. "你的坚持，终会有回报"
15. "你的隐忍，是一种智慧"

**【祝福期许类】**
16. "你吃过的苦，都会变成光"
17. "好日子，在后头等着你"
18. "你值得这世上最好的一切"
19. "属于你的，正在路上"
20. "你的好运，马上就到"

**【格局境界类】**
21. "懂得忍耐的人，都有大格局"
22. "不争不抢的人，往往赢到最后"
23. "看得开的人，才能走得远"
24. "你的沉默，是最高级的智慧"
25. "你的大度，别人学不来"

**【因果回报类】**
26. "厚道人，迟早会有好报"
27. "心软的人，最有福气"
28. "善良的人，老天不会亏待"
29. "你积的德，都会还给你"
30. "好人有好报，你就是证明"

**【理解共鸣类】**
31. "你的委屈，有人懂"
32. "你不是一个人在扛"
33. "你的苦，我都知道"
34. "你的不容易，我看在眼里"
35. "你受的累，都值得"

**【身份认同类】**
36. "老实人，才是真正的聪明人"
37. "心软的人，骨子里最硬"
38. "不爱说话的人，心里最有数"
39. "像你这样的人，最让人心疼"
40. "你这种人，最值得深交"

**【温暖治愈类】**
41. "余生不长，你该对自己好点了"
42. "你已经很棒了，别太苛责自己"
43. "你值得被好好爱着"
44. "辛苦了，给自己一个拥抱"
45. "你的存在，本身就是一种力量"

**【转折肯定类】**
46. "曾经看不起你的人，现在都后悔了"
47. "熬过去，你就赢了"
48. "你的时代，终于要来了"
49. "属于你的精彩，才刚刚开始"
50. "你的春天，不会缺席"

**标题禁忌：**
- ❌ 不要用疑问句、反问句
- ❌ 不要制造悬念
- ❌ 不要用"为什么""怎么""吗"等疑问词
- ❌ 不要打击、质疑读者
- ❌ 不要用负面情绪词汇

**标题生成规则：**
- 每篇文章的5个标题必须从不同类型中选择
- 3篇文章共15个标题，风格要有变化
- 全部使用第二人称"你"
- 全部是肯定、夸赞、祝福的语气
- 禁止出现任何重复、相似、雷同的标题

## 敏感词规避（以下词汇绝对不能出现！）

- 玄学类：仙师、徒儿、为师、天机、命盘、磁场、气运、福报、姻缘、符咒、渡劫、运势、运气、贵人、小人、业障、因果报应、天命、命数、劫难、气场、能量场
- 迷信类：算命、占卜、风水、转运、开光、法事、神仙、菩萨、鬼神、法术、阴阳、五行、宇宙能量
- 人设类：为师、贫道、本仙师、师傅、徒儿、道友、有缘人
- 承诺类：改运、转运、翻身、暴富

## 引流类型：{flow_type}
{flow_instruction}

## 输出格式
═══════════════════════════════════════
【第一篇】
═══════════════════════════════════════

【标题1】XXX
【标题2】XXX
【标题3】XXX
【标题4】XXX
【标题5】XXX

---

正文内容...（必须1500字以上）

═══════════════════════════════════════
【第二篇】
═══════════════════════════════════════
...（以此类推，共3篇）

请直接输出仿写结果，不要有任何说明性文字。"""

        return prompt

    def get_flow_instruction(self, flow_type, yinliu_content, product_name, product_material):
        """获取引流类型的具体指令"""
        if flow_type == "置顶引流":
            instruction = """## 结尾引流方式：置顶引流

**引流的核心目标：让读者不去看就睡不着觉、吃不下饭、觉得亏了100万！**

**引流铺垫要求（最重要！）：**
- 不能只是简单说"去看主页置顶"，必须说清楚**为什么要去看**
- 要制造强烈的好奇心和紧迫感
- 要让读者觉得：不去看就错过了改变人生的机会
- 要让读者觉得：那里有专门为他准备的答案/方法/秘诀
- 引流话术要与前文内容自然衔接，有理有据

**置顶引流的钩子设计：**
1. **悬念钩子**：前文讲了问题，但解决方法/核心秘诀在置顶视频里
2. **稀缺钩子**：这些内容太重要了，不适合公开讲，只放在置顶
3. **专属钩子**：那里有专门为"像你这样的人"准备的内容
4. **紧迫钩子**：错过这个机会，可能要再等很久

**置顶引流话术库（每次选择不同的，必须有铺垫！）：**

1. **悬念型铺垫：**
"关于怎么彻底走出这个困局，我在主页置顶视频里讲得很透。那里有一套方法，是我这些年摸爬滚打总结出来的，不适合在这里公开讲。你要是真想改变，点我头像，去看置顶第一条，看完你就明白该怎么做了。"

2. **专属型铺垫：**
"我知道你现在最需要的是什么。我专门录了一期视频放在主页置顶，就是给像你这样心里有苦却不知道怎么办的人看的。点我头像，去看第一条，那里的话，是专门说给你听的。"

3. **价值型铺垫：**
"这个道理我想了很多年才想通，三言两语说不完。我把最核心的部分放在了主页置顶视频里，那里有完整的思路和方法。你要是不想再走弯路，点我头像，去看置顶那条，看完你会感谢自己今天的选择。"

4. **紧迫型铺垫：**
"有些话我不方便在这里说太多，但我不忍心看你继续这样熬下去。我在主页置顶放了一期视频，里面的内容可能会颠覆你的认知。点我头像，现在就去看，别等到事情更糟了才后悔。"

5. **救赎型铺垫：**
"如果你也想从这种困境里走出来，想活得通透一点、轻松一点，点我头像，去看主页置顶的视频。那里有你一直在找的答案，也有你需要的那份力量。"

6. **共鸣型铺垫：**
"我知道你现在的感受，因为我也曾经历过。后来我想通了一些事，才慢慢走出来的。这些经验我都放在了主页置顶视频里，点我头像去看，希望能帮你少走一些弯路。"

7. **秘诀型铺垫：**
"关于怎么改变这种局面，有一个关键点我没在这里讲。因为这个方法太重要了，我专门录了视频放在主页置顶。点我头像，去看第一条，看懂了，你的人生可能就不一样了。"

8. **故事续集型铺垫：**
"这个故事还没讲完，最关键的转折点我放在了主页置顶视频里。你要是想知道后来怎么样了，点我头像去看，那里有你想要的答案。"

**用户自定义话术：**"""
            if yinliu_content:
                instruction += f"\n{yinliu_content}"

        elif flow_type == "橱窗引流":
            instruction = """## 结尾引流方式：橱窗引流

**橱窗引流的钩子设计：**
1. **助力钩子**：那里有能帮助你的好物/工具
2. **改变钩子**：想要改变现状，需要一些助力
3. **犒赏钩子**：你值得对自己好一点

**橱窗引流话术库（必须有铺垫，自然过渡！）：**

1. **助力型铺垫：**
"走到今天，你已经很不容易了。有时候，我们需要一些外在的助力，来帮自己稳住心神、找回状态。我在主页橱窗里放了一些我精心挑选的好物，都是能帮你调整状态、提升心境的。点我头像，进主页橱窗看看，选一件让你心动的，就当是给自己的一份礼物。"

2. **犒赏型铺垫：**
"你为别人操心了大半辈子，是时候对自己好一点了。我在主页橱窗里准备了一些能让你静心、养神的好物。点我头像，进橱窗看看，挑一件喜欢的，犒赏一下这些年辛苦的自己。"

3. **改变型铺垫：**
"想要改变现状，光靠想是不够的，得有行动。我在主页橱窗里放了一些能帮你提升状态、改善心境的好物，每一件都是我精心挑选的。点我头像，进橱窗看看，选一件带走，就是你改变的第一步。"

4. **能量型铺垫：**
"有时候，一件对的物品，能给你带来意想不到的力量。我在主页橱窗里放了一些我自己也在用的好物，能帮你静心、提神、找回状态。点我头像，进橱窗看看，第一眼让你心动的那件，就是适合你的。"

5. **祝福型铺垫：**
"最后送你一句话：你值得这世间最好的一切。我在主页橱窗里准备了一些好物，希望能给你带来一点温暖和力量。点我头像，进橱窗看看，带走一件，就当是我送你的祝福。"

**用户自定义话术：**"""
            if yinliu_content:
                instruction += f"\n{yinliu_content}"

        elif flow_type == "带货引流":
            instruction = f"""## 结尾引流方式：带货引流

**商品名称：{product_name}**
**产品素材：{product_material}**

**带货引流要求：**
- 引流话术要与前文内容自然衔接
- 要突出商品能解决读者的问题/满足读者的需求
- 要制造紧迫感和稀缺感
- 话术要真诚，不能太硬广

**用户自定义话术：**"""
            if yinliu_content:
                instruction += f"\n{yinliu_content}"
            else:
                instruction += f"\n这款{product_name}真的很不错，点我头像进橱窗了解一下。"

        else:  # 纯夸赞不引流
            instruction = """## 结尾方式：纯夸赞不引流

**要求：**
- 纯夸赞读者，给予温暖和力量
- 不需要任何引流话术
- 结尾要让读者感到被理解、被肯定、被温暖
- 给读者希望和方向，让他们觉得"还有出路"
- 用温暖有力的话收尾，让读者感动"""

        return instruction

    def save_document(self, content, output_path, index):
        """保存为Word文档"""
        try:
            from docx import Document
            from docx.shared import Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = Document()

            # 设置标题
            title = doc.add_heading('百家号仿写文案', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 添加内容
            lines = content.split('\n')
            for line in lines:
                line = line.strip()
                if not line:
                    continue

                if line.startswith('═') or line.startswith('---'):
                    doc.add_paragraph('─' * 40)
                elif line.startswith('【第') and '篇】' in line:
                    p = doc.add_paragraph()
                    run = p.add_run(line)
                    run.bold = True
                    run.font.size = Pt(14)
                elif line.startswith('【标题'):
                    p = doc.add_paragraph(line)
                    p.runs[0].bold = True
                else:
                    doc.add_paragraph(line)

            # 生成文件名
            timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
            filename = f"{timestamp}_{index}.docx"
            filepath = os.path.join(output_path, filename)

            doc.save(filepath)
            self.log(f"文档已保存: {filename}")
            return filepath

        except ImportError:
            self.log("错误：请安装python-docx库 (pip install python-docx)")
            # 备用方案：保存为txt
            timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
            filename = f"{timestamp}_{index}.txt"
            filepath = os.path.join(output_path, filename)
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write(content)
            self.log(f"已保存为TXT: {filename}")
            return filepath

        except Exception as e:
            self.log(f"保存文档失败: {str(e)}")
            return None

    def finish_task(self):
        """完成任务，恢复UI状态"""
        self.is_running = False
        self.start_btn.config(state=tk.NORMAL)
        self.stop_btn.config(state=tk.DISABLED)
        # 如果有生成过的文案，启用重新生成按钮
        if self.last_articles:
            self.regenerate_btn.config(state=tk.NORMAL)
        self.update_status("处理完成")

        # 询问是否打开输出文件夹
        if messagebox.askyesno("完成", "文案生成完成！是否打开输出文件夹？"):
            self.open_output_folder()

    def open_output_folder(self):
        """打开输出文件夹"""
        output_path = self.output_path.get().replace('/', '\\')
        if os.path.exists(output_path):
            import subprocess
            subprocess.Popen(f'explorer "{output_path}"', shell=True)
        else:
            messagebox.showerror("错误", f"文件夹不存在: {output_path}")


def main():
    root = tk.Tk()
    app = FangxieApp(root)
    root.mainloop()


if __name__ == "__main__":
    main()
