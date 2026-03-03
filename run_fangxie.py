# -*- coding: utf-8 -*-
import anthropic
from docx import Document
from datetime import datetime
import os
import re

client = anthropic.Anthropic(
    base_url="https://yunyi.rdzhvip.com/claude",
    api_key="A68YGT53-41HW-8GT7-353U-MJVDNY8NE5KX"
)

prompt = """你是专业情感文案写手。现在直接生成3篇带货文案，每篇1200-1500字。

【商品】黄铜貔貅摆件（一对）
- 公貔貅主招财，脚踩绣球；母貔貅主守财，脚踩如意
- 成对摆放寓意"招财守财，财富双收"
- 黄铜铸造，真材实料，纳米镀层防水防污
- 适合客厅、办公桌、收银台摆放

【写作要求】
1. 第二人称"你"贯穿全文，像朋友聊天
2. 3篇开头必须完全不同
3. 每篇5个标题，共15个标题不能重复
4. 结尾自然引导到貔貅产品，点击小黄车购买
5. 每篇引用不同经典（道德经/菜根谭/庄子等）

【3篇主题】
第1篇：从"被忽视的付出"切入，肯定读者默默承受的一切
第2篇：从"中年觉醒"切入，人到中年才懂善待自己
第3篇：从"独处智慧"切入，独处是内心富足的表现

【禁止词汇】命运、运势、福报、因果、天命、气场、能量场、风水、阴阳、五行、改运、转运、暴富

【输出格式】严格按以下格式，直接开始：

【第一篇】

【标题1】xxx
【标题2】xxx
【标题3】xxx
【标题4】xxx
【标题5】xxx

---

正文内容（1200-1500字）

═══════════════════════════════════════════════════

【第二篇】

【标题1】xxx
【标题2】xxx
【标题3】xxx
【标题4】xxx
【标题5】xxx

---

正文内容（1200-1500字）

═══════════════════════════════════════════════════

【第三篇】

【标题1】xxx
【标题2】xxx
【标题3】xxx
【标题4】xxx
【标题5】xxx

---

正文内容（1200-1500字）

现在直接从【第一篇】开始输出："""

print("正在生成文案...")

response = client.messages.create(
    model="claude-opus-4-5-20251101",
    max_tokens=16000,
    messages=[{"role": "user", "content": prompt}]
)

content = response.content[0].text
print(f"生成内容：{len(content)}字")
print("="*50)
print(content[:2000])
print("="*50)

# 解析
articles = []
parts = re.split(r'【第[一二三]篇】', content)
for part in parts:
    part = part.strip()
    if len(part) < 500:
        continue
    titles = re.findall(r'【标题[1-5]】\s*(.+?)(?=\n|$)', part)
    body = re.sub(r'【标题[1-5]】.+?\n', '', part)
    body = re.sub(r'^-{3,}$', '', body, flags=re.MULTILINE)
    body = re.sub(r'^═+$', '', body, flags=re.MULTILINE)
    body = body.strip()
    if len(body) >= 500:
        articles.append({'titles': titles or [f"文案{len(articles)+1}"], 'content': body})
        print(f"第{len(articles)}篇：{len(body)}字，{len(titles)}个标题")

# 保存
if articles:
    os.makedirs(r"D:\A百家号带货视频\带货文案", exist_ok=True)
    ts = datetime.now().strftime('%Y%m%d%H%M%S')
    fp = rf"D:\A百家号带货视频\带货文案\{ts}_1.docx"
    doc = Document()
    doc.add_paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph("═" * 40)
    for i, a in enumerate(articles, 1):
        doc.add_paragraph(f"【第{i}篇】", style='Heading 1')
        for j, t in enumerate(a['titles'][:5], 1):
            doc.add_paragraph(f"【标题{j}】{t}")
        doc.add_paragraph("---")
        doc.add_paragraph(a['content'])
        doc.add_paragraph(f"（字数：{len(a['content'])}）")
        doc.add_paragraph("═" * 40)
    doc.save(fp)
    print(f"\n已保存：{fp}")
else:
    print("解析失败，原文：")
    print(content[:3000])
